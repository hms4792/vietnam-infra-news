"""
report_regional.py — 지역성 프로젝트형 보고서 생성기 v4.0
===========================================================
대상 플랜:
  VN-TRAN-2055, VN-URB-METRO-2030, VN-MEKONG-DELTA-2030,
  VN-RED-RIVER-2030, VN-IP-NORTH-2030, VN-WW-2030,
  VN-SWM-NATIONAL-2030 (혼합)

보고서 구조:
  표지
  Section 1. 마스터플랜 개요 + 공간 지도
  Section 2. 핵심 사업 목록 (전체 조감표)
  Section 3. 하부 프로젝트별 상세 (★) — 개요 + 지도 + 추진경위 + 최신/역사기사
  Section 4. 지역(Province)별 활동 현황
  Section 5. 한국 기업 기회
  Appendix : 전체 관련 기사
"""
import os, json
from datetime import datetime
from pathlib import Path
from collections import defaultdict

from report_lib import (
    new_doc, add_heading, add_hr, add_note, add_bullet, add_label_value,
    styled_table, kpi_box, info_box, project_card, build_cover,
    render_articles, render_history_timeline, render_province_section,
    render_korean_opportunity, build_appendix, insert_map,
    load_history, load_kpi_db, save_report,
    select_articles, _date_key, PALETTE, PALETTE_HEX, _rgb, cell_bg,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path("/home/work/claw")

# ─────────────────────────────────────────────────────────
# 플랜별 정적 메타데이터
# ─────────────────────────────────────────────────────────
PLAN_CONFIG = {
    "VN-TRAN-2055": {
        "name_ko": "국가 교통인프라 마스터플랜 2055",
        "name_en": "Vietnam National Transport Infrastructure Master Plan to 2055",
        "legal_basis": "Decision 1454/QD-TTg (July 1, 2021)",
        "lead_agency": "Ministry of Transport (MOT)",
        "investment": "약 $243B (2021-2030 $130B + 2030-2055 $113B)",
        "period": "2021 – 2055",
        "theme": ("orange", "EA580C", "FFF7ED"),
        "report_type": "🗺️ 지역성 프로젝트 보고서",
        "overview": (
            "베트남 2055년 교통 인프라 비전: 현대 고속도로 9,014km, 철도 4,000km, "
            "국제공항 30개, 심수항 세계 4위권 목표. "
            "2030 목표: 고속도로 5,000km, 고속철도 기본 선로 완공, "
            "Long Thanh 국제공항 1단계 운영, Lach Huyen 심수항 2단계 완공."
        ),
        "key_milestones": [
            "2021.07: Decision 1454/QD-TTg 마스터플랜 확정",
            "2023: 고속도로 남북축 핵심 구간 착공",
            "2025: Long Thanh 공항 1단계 준공 목표 (지연 → 2026)",
            "2030: 고속도로 5,000km 달성 목표",
        ],
        "projects_by_province": True,
        "korean_opportunities": [
            ("고속도로 EPC", "HIGH",
             "남북고속도로·링로드·지방 고속도로 — 현대건설·GS건설·DL이앤씨 수주 기회"),
            ("공항 설계·시공", "HIGH",
             "Long Thanh 1단계($16B), 다낭 T3, 노이바이 T3 — 인천공항공사 기술자문 + 한국 건설사"),
            ("철도 차량·시스템", "HIGH",
             "고속철도·도시철도 차량 — 현대로템 + LS Electric 컨소시엄"),
            ("항만 장비·물류", "MEDIUM",
             "Lach Huyen 2단계, Lien Chieu 신항 — 한국 항만 장비(크레인) + 물류 거점"),
            ("교통 ITS·스마트도로", "MEDIUM",
             "고속도로 스마트 교통 시스템 — 한국 ITS 솔루션 (SK, KT 컨소시엄)"),
            ("ODA 활용", "HIGH",
             "EDCF(수출입은행) 철도·항만 ODA 패키지 — 한국 기업 참여 조건부"),
        ],
    },

    "VN-URB-METRO-2030": {
        "name_ko": "도시철도·메트로 개발계획 2030",
        "name_en": "Urban Rail Transit Master Plan 2030",
        "legal_basis": "Decision 1210/QD-TTg (2023) + Hanoi Decision 519/QD-TTg (2016) + HCMC Resolution",
        "lead_agency": "Ministry of Construction (MOC) + 하노이/호치민 인민위원회",
        "investment": "$60B+ (하노이 $30B + 호치민 $25B + 기타 도시)",
        "period": "2023 – 2035",
        "theme": ("red", "DC2626", "FEF2F2"),
        "report_type": "🗺️ 지역성 프로젝트 보고서",
        "overview": (
            "하노이 15개 노선/616km, 호치민 8개 노선/220km 목표. "
            "2025 현재: 하노이 2A선(13km) + 3호선(부분) 운영. 호치민 1호선 건설중. "
            "전국 6개 도시 메트로 계획 (하노이·호치민·다낭·하이퐁·껀터·낀호아)."
        ),
        "key_milestones": [
            "2021.11: 하노이 Line 2A (Cat Linh – Ha Dong) 정식 개통",
            "2024: 하노이 Line 3 Nhon – Hanoi Station 부분 개통",
            "2024: 호치민 Line 1 준공 목표 (지연 → 2025~2026)",
            "2035: 하노이 주요 노선 완공 목표",
        ],
        "projects_by_province": True,
        "korean_opportunities": [
            ("차량 제작·납품", "HIGH",
             "현대로템 — 하노이 Line 1·3·5·6, 호치민 Line 2·3·4 입찰 기회"),
            ("신호·전장 시스템", "HIGH",
             "LS Electric·현대엘리베이터 — 변전소·전기시스템 일괄 공급"),
            ("건설 EPC", "HIGH",
             "현대건설·DL이앤씨·SK에코플랜트 — 지하구간 EPC 패키지"),
            ("운영·유지관리 시스템", "MEDIUM",
             "SR(수서고속철도) 운영 노하우 + 한국 유지관리 솔루션 수출"),
            ("역사 디자인·상업화", "MEDIUM",
             "롯데건설 역사 개발 + 한국 리테일 브랜드 입점"),
            ("ODA·EDCF", "HIGH",
             "EDCF 철도 패키지 — 하노이 Line 5 ODA 협의 진행 중"),
        ],
    },

    "VN-MEKONG-DELTA-2030": {
        "name_ko": "메콩델타 지역개발 마스터플랜 2021-2030",
        "name_en": "Mekong Delta Regional Development Master Plan 2021–2030, Vision 2050",
        "legal_basis": "Decision 287/QD-TTg (2022) + Decision 616/QD-TTg (2026 조정)",
        "lead_agency": "Ministry of Planning & Investment (MPI) + 13개 성시 인민위원회",
        "investment": "$65B (2021-2030) | 고속도로 830km 포함",
        "period": "2021 – 2030 / Vision 2050",
        "theme": ("teal", "0D9488", "F0FDFA"),
        "report_type": "🗺️ 지역성 프로젝트 보고서",
        "overview": (
            "베트남 13개 성시, 면적 4만km², 인구 1,700만 메콩 삼각주 종합개발계획. "
            "기후변화 대응(해수면 상승, 염수 침투), 농업현대화, 교통 인프라 구축이 핵심. "
            "2030년 고속도로 830km 목표 — 현재 180km 운영."
        ),
        "key_milestones": [
            "2022: Decision 287/QD-TTg 메콩델타 마스터플랜 확정",
            "2026: Decision 616/QD-TTg 조정계획 발표",
            "2026: 메콩델타 고속도로 5개 구간 동시 착공",
            "2030: 고속도로 830km, 국제공항 2개(껀터·퍼꾸억) 운영 목표",
        ],
        "projects_by_province": True,
        "korean_opportunities": [
            ("고속도로 건설", "HIGH",
             "메콩델타 고속도로 830km — 한국 건설사 패키지 수주 (연약지반 공법)"),
            ("기후변화 인프라", "HIGH",
             "해수 차단 제방·수문·범람 방지 — K-water, 한국수자원공사 기술"),
            ("농업현대화", "MEDIUM",
             "스마트팜·수직농장·콜드체인 — 한국 농업기술기업 진출"),
            ("수산업 인프라", "MEDIUM",
             "새우·메기 가공시설·항만 냉동창고 — 한국 식품기업 투자"),
            ("재생에너지", "MEDIUM",
             "메콩델타 수상태양광 잠재력 1,200MW+ — 한국 RE 기업"),
        ],
    },

    "VN-RED-RIVER-2030": {
        "name_ko": "홍강 델타 지역 개발 마스터플랜 2030",
        "name_en": "Red River Delta Regional Development Master Plan 2021–2030, Vision 2045",
        "legal_basis": "Decision 368/QD-TTg (2022) + Decision 612/QD-TTg (2026 조정)",
        "lead_agency": "Ministry of Planning & Investment (MPI) + 11개 성시",
        "investment": "$150B+ (2021-2030)",
        "period": "2021 – 2030 / Vision 2045",
        "theme": ("red", "BE123C", "FFF1F2"),
        "report_type": "🗺️ 지역성 프로젝트 보고서",
        "overview": (
            "베트남 경제 핵심 지역 — 하노이·하이퐁·꽝닌 삼각축. "
            "11개 성시, 면적 2.1만km², 인구 2,300만. "
            "2030년 고속도로 1,300km, 산업단지 100개, FDI 누적 $180B 목표."
        ),
        "key_milestones": [
            "2022: Decision 368/QD-TTg 홍강 델타 마스터플랜 확정",
            "2026: Decision 612/QD-TTg 조정계획 발표",
            "2025: 링로드4 착공 (하노이-박닌-흥이엔 112km)",
            "2030: 고속도로 1,300km, 연간 FDI $15B 목표",
        ],
        "projects_by_province": True,
        "korean_opportunities": [
            ("산업단지 개발", "HIGH",
             "VSIP·KIA·한국형 산업단지 — LH공사·포스코인터내셔널 추가 부지 확보"),
            ("링로드4 건설", "HIGH",
             "하노이 외곽순환 112km/$8.5B — 현대건설·GS건설 컨소시엄"),
            ("스마트산업단지", "HIGH",
             "폐수처리·환경모니터링·에너지관리 패키지 — 한국 환경기업"),
            ("반도체·전자 공급망", "MEDIUM",
             "삼성(박닌·타이응웬) 밸류체인 — 한국 소부장 기업 동반 진출"),
            ("물류 거점", "MEDIUM",
             "하이퐁항 배후단지 — CJ대한통운·한진 복합물류센터"),
        ],
    },

    "VN-IP-NORTH-2030": {
        "name_ko": "북부 산업단지 개발 계획 2030",
        "name_en": "Northern Industrial Park Development Plan 2030",
        "legal_basis": "Decision 1107/QD-TTg (Industrial Park Master Plan 2030, Vision 2050)",
        "lead_agency": "Ministry of Planning & Investment (MPI) + DEZA",
        "investment": "$30B+ (산업단지 인프라 + FDI 유치)",
        "period": "2021 – 2030",
        "theme": ("violet", "7C3AED", "F5F3FF"),
        "report_type": "🗺️ 지역성 프로젝트 보고서",
        "overview": (
            "하노이·박닌·하이퐁·타이응웬·빈푹 5개 핵심 성시 중심 산업단지 클러스터. "
            "삼성·LG·폭스콘 등 글로벌 전자기업 밀집. "
            "2030년 산업단지 300개, 입주율 90%, 스마트 산업단지 전환 50% 목표."
        ),
        "key_milestones": [
            "2021: Decision 1107/QD-TTg 산업단지 마스터플랜 확정",
            "2024: VSIP Hung Yen 3단계 착공",
            "2025: 북부 산업단지 입주율 평균 85% 달성",
            "2030: 스마트 산업단지 50% 전환 목표",
        ],
        "projects_by_province": True,
        "korean_opportunities": [
            ("산업단지 개발·운영", "HIGH",
             "LH공사·포스코인터내셔널·한국토지신탁 — 신규 산업단지 개발 합작"),
            ("전자·반도체 소부장", "HIGH",
             "삼성(박닌·타이응웬) 협력사 — 공장 증설 + 신규 입주"),
            ("스마트팩토리 솔루션", "HIGH",
             "현대·LS·포스코 스마트팩토리 시스템 — 산업단지 디지털화"),
            ("폐수처리·환경", "MEDIUM",
             "산업단지 공동 폐수처리장 — 한국 환경엔지니어링사"),
            ("에너지 공급", "MEDIUM",
             "산업단지 자체 발전·재생에너지 — 한국 RE·ESS 기업"),
        ],
    },

    "VN-WW-2030": {
        "name_ko": "국가 폐수처리 마스터플랜 2021-2030",
        "name_en": "National Wastewater Treatment Master Plan 2021–2030",
        "legal_basis": "Decision 1393/QD-TTg (2021) + Decision 1659/QD-TTg (2021, 도시위생)",
        "lead_agency": "Ministry of Construction (MOC) + MONRE",
        "investment": "$12B (2021-2030) | ODA 비중 40%",
        "period": "2021 – 2030",
        "theme": ("sky", "0EA5E9", "F0F9FF"),
        "report_type": "🗺️ 지역성 프로젝트 보고서",
        "overview": (
            "2020년 기준 베트남 도시 폐수처리율 15%. "
            "2025년 20~30%, 2030년 70% 달성 목표. "
            "2030년 처리용량 300-400만 ㎥/일 목표. "
            "주요 프로젝트: 하노이 7개 처리장 120만㎥/일, 호치민 9개 210만㎥/일."
        ),
        "key_milestones": [
            "2021: Decision 1393/QD-TTg 마스터플랜 확정",
            "2024: 하노이 Yen So 2단계 (40만㎥/일) 착공",
            "2025: 호치민 Thu Duc 3단계 (22.5만㎥/일) 완공 목표",
            "2030: 전국 처리율 70% 달성 목표",
        ],
        "projects_by_province": True,
        "korean_opportunities": [
            ("WWTP EPC", "HIGH",
             "하노이 Yen So·Bac Thang Long, 호치민 Nhieu Loc — 한국 환경건설사 EPC"),
            ("막분리(MBR) 기술", "HIGH",
             "코오롱인더스트리·도레이·휴비스 MBR 모듈 — 기술이전 + 현지 생산"),
            ("슬러지 처리·에너지화", "MEDIUM",
             "슬러지 소화조·바이오가스 — 한국 기술 + 탄소크레딧 연계"),
            ("스마트 모니터링", "MEDIUM",
             "폐수처리장 원격모니터링·AI 운영 — 한국 환경IT 기업"),
            ("ODA 패키지", "HIGH",
             "EDCF + KOICA — 하노이·하이퐁 폐수처리 ODA 패키지 활용"),
        ],
    },

    "VN-SWM-NATIONAL-2030": {
        "name_ko": "전국 고형폐기물 통합관리 국가전략 2025/2050",
        "name_en": "National Integrated Solid Waste Management Strategy 2025/2050",
        "legal_basis": "Decision 491/QD-TTg (2018) + Decision 1658/QD-TTg (2021 녹색성장)",
        "lead_agency": "Ministry of Natural Resources & Environment (MONRE) + MOC",
        "investment": "$10B+ (2021-2030) — WtE 중심",
        "period": "2018 – 2025 / Vision 2050",
        "theme": ("green", "15803D", "F0FDF4"),
        "report_type": "🗺️ 지역성 프로젝트 보고서",
        "overview": (
            "2025 목표: 도시 수거율 90%, 농촌 70%. 처리 기술 다변화(매립→소각·WtE). "
            "2030 목표: 매립지 비율 30% 이하, WtE 처리 50%. "
            "핵심 프로젝트: 하노이 4,000TPD WtE, 호치민 2,000TPD WtE, 다낭 650TPD."
        ),
        "key_milestones": [
            "2018: Decision 491/QD-TTg 국가전략 확정",
            "2022: 하노이 WtE 4,000TPD 사업자 선정 공고",
            "2024: 호치민 WtE 2,000TPD EPC 입찰",
            "2025: 도시 수거율 90% 달성 목표",
        ],
        "projects_by_province": True,
        "korean_opportunities": [
            ("WtE 소각발전", "HIGH",
             "하노이 4,000TPD, 호치민 2,000TPD — 코오롱글로벌·한국지역난방공사 컨소시엄"),
            ("재활용 기술", "HIGH",
             "분리수거·자동분류 설비 — 한국 재활용 기술 수출"),
            ("의료폐기물 처리", "MEDIUM",
             "병원 의료폐기물 소각 시스템 — 한국 의료폐기물 전문기업"),
            ("스마트 수거 시스템", "MEDIUM",
             "IoT 기반 폐기물 수거 경로 최적화 — 한국 스마트시티 솔루션"),
            ("ODA·환경펀드", "MEDIUM",
             "GCF(녹색기후기금) + EDCF — 한국 환경기업 참여 연계"),
        ],
    },
}

# ─────────────────────────────────────────────────────────
# 핵심 함수: 지역성 보고서 생성
# ─────────────────────────────────────────────────────────
def generate_regional_report(plan_id: str, output_dir: str) -> str:
    if plan_id not in PLAN_CONFIG:
        raise ValueError(f"PLAN_CONFIG에 {plan_id}가 없습니다.")

    cfg = PLAN_CONFIG[plan_id]
    theme_color, theme_hex, theme_light = cfg["theme"]
    now = datetime.now()
    week = now.isocalendar()[1]
    week_str = f"W{week:02d}/{now.year}"

    # 기사 로드
    buckets = load_history([plan_id])
    arts = sorted(buckets.get(plan_id, []), key=_date_key, reverse=True)
    rec_arts, year_label = select_articles(arts)

    # KPI DB 로드
    kpi_db = load_kpi_db()

    doc = new_doc()

    # ── 표지 ──────────────────────────────────────────────
    build_cover(
        doc, plan_id,
        cfg["name_ko"], cfg["name_en"],
        cfg["legal_basis"], cfg["lead_agency"],
        cfg["investment"], cfg["period"],
        len(arts), week_str, cfg["report_type"],
        theme_hex
    )

    # ── Section 1: 마스터플랜 개요 ────────────────────────
    add_heading(doc, "1. 마스터플랜 개요", 1, theme_color, 13)
    add_hr(doc, theme_hex)

    # 개요 설명 박스
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    ot = doc.add_table(rows=1, cols=1); ot.style = "Table Grid"
    oc = ot.rows[0].cells[0]; cell_bg(oc, theme_light)
    op = oc.paragraphs[0]
    op.add_run(cfg["overview"]).font.size = Pt(10)
    doc.add_paragraph()

    # 추진 경위 타임라인
    add_heading(doc, "1-1. 추진 경위", 2, theme_color)
    for ms in cfg.get("key_milestones", []):
        add_bullet(doc, ms, size=10)
    doc.add_paragraph()

    # 공간 지도
    add_heading(doc, "1-2. 공간 범위 지도", 2, theme_color)
    insert_map(doc, plan_id, caption=f"{cfg['name_ko']} — 주요 사업 위치 (OSM)")
    doc.add_paragraph()

    # ── Section 2: 핵심 사업 목록 ─────────────────────────
    add_heading(doc, "2. 핵심 사업 목록 (전체 조감)", 1, theme_color, 13)
    add_hr(doc, theme_hex)

    # KPI DB에서 프로젝트 수집
    projects_all = []
    if plan_id in kpi_db:
        prov_data = kpi_db[plan_id].get("provinces", {})
        for prov_name, pdata in prov_data.items():
            for proj in pdata.get("projects", []):
                proj_copy = dict(proj)
                proj_copy["_province"] = prov_name
                projects_all.append(proj_copy)

    if projects_all:
        rows = []
        for i, proj in enumerate(projects_all, 1):
            inv = proj.get("investment") or proj.get("investment_usd") or "TBD"
            area = f"{proj.get('area_ha','')}ha" if proj.get("area_ha") else ""
            length = f"{proj.get('length_km','')}km" if proj.get("length_km") else ""
            scale = " / ".join(filter(None, [area, length, str(inv)]))[:40]
            rows.append([
                str(i),
                proj.get("_province", "-")[:15],
                proj.get("name", "-")[:45],
                scale[:35],
                proj.get("status", "계획")[:20],
                proj.get("deadline", "-")[:12],
            ])
        styled_table(doc,
            ["No", "지역", "프로젝트명", "규모/투자", "현황", "완료목표"],
            rows,
            col_widths=[0.3, 1.1, 2.8, 2.0, 1.2, 0.9],
            header_bg=theme_color, font_size=9)
    else:
        add_note(doc, "ℹ️ KPI DB에 등록된 세부 프로젝트가 없습니다. province_project_kpi.json 확장 후 재생성 필요.", "gray")

    doc.add_page_break()

    # ── Section 3: 하부 프로젝트별 상세 ──────────────────
    add_heading(doc, "3. 하부 프로젝트별 상세 추적", 1, theme_color, 13)
    add_hr(doc, theme_hex)
    add_note(doc,
        "각 프로젝트의 추진 현황을 역사 기사(2019-2025) → 최신 기사(2026) 흐름으로 추적합니다.",
        color_key="gray")
    doc.add_paragraph()

    if plan_id in kpi_db:
        prov_data = kpi_db[plan_id].get("provinces", {})
        proj_idx = 0
        for prov_name, pdata in prov_data.items():
            projects = pdata.get("projects", [])
            if not projects: continue

            # 성(Province) 헤더
            add_heading(doc, f"📍 {prov_name}", 2, theme_color)

            # KPI 목표 (있으면)
            kpi_2030 = pdata.get("kpi_2030", {})
            if kpi_2030:
                kpi_items = [(k.replace("_"," ").title(), v, "") for k,v in list(kpi_2030.items())[:4]]
                kpi_box(doc, kpi_items, bg_hex=theme_light, border_hex=theme_hex,
                        title=f"{prov_name} — KPI 2030 목표")

            for proj in projects:
                proj_idx += 1
                # key_projects가 string인 경우 dict로 변환
                if isinstance(proj, str):
                    proj = {"name": proj}
                # 프로젝트 카드
                project_card(doc, proj, theme_hex, theme_color, idx=proj_idx)

                # 해당 프로젝트 관련 기사 필터링
                proj_name_lower = proj.get("name","").lower()
                proj_keywords = [w for w in proj_name_lower.split() if len(w) > 3]
                prov_lower = prov_name.lower()

                proj_arts = []
                for art in arts:
                    text = (
                        (art.get("title","") or "") + " " +
                        (art.get("summary_en","") or "") + " " +
                        (art.get("content","") or "")[:300]
                    ).lower()
                    # 프로젝트 키워드 or 성 이름 포함
                    if (any(kw in text for kw in proj_keywords) or
                        prov_lower in text):
                        proj_arts.append(art)

                proj_arts = sorted(set(id(a) for a in proj_arts) and proj_arts,
                                   key=_date_key, reverse=True)
                # 중복 제거
                seen = set(); uniq = []
                for a in proj_arts:
                    url = a.get("url","") or str(id(a))
                    if url not in seen: seen.add(url); uniq.append(a)
                proj_arts = uniq

                # 최신 기사 (2026)
                recent_proj, yr_lbl = select_articles(proj_arts)
                if recent_proj:
                    add_heading(doc, f"  📰 최신 뉴스 ({yr_lbl})", 3, theme_color)
                    render_articles(doc, recent_proj, theme_color, max_n=3)

                # 역사 기사 타임라인 (2025 이전)
                hist_arts = [a for a in proj_arts if _date_key(a) < "2026-01-01"]
                if hist_arts:
                    add_heading(doc, "  📚 역사 기사 추적 (2019–2025)", 3, theme_color)
                    render_history_timeline(doc, hist_arts, theme_color, max_n=3)

                if not proj_arts:
                    add_note(doc,
                        "  ℹ️ 이 프로젝트 관련 기사가 아직 부족합니다. "
                        "키워드 매칭은 다음 주간 수집부터 강화됩니다.",
                        "gray")

                add_hr(doc, "E5E7EB"); doc.add_paragraph()
    else:
        # KPI DB 없는 경우 — 기사만으로 섹션 구성
        add_note(doc, f"ℹ️ {plan_id} KPI DB가 아직 없습니다. 기사 기반 섹션만 표시합니다.", "gray")
        add_heading(doc, "최신 뉴스", 2, theme_color)
        render_articles(doc, rec_arts, theme_color, max_n=8)
        add_heading(doc, "역사 기사 타임라인", 2, theme_color)
        render_history_timeline(doc, arts, theme_color)

    doc.add_page_break()

    # ── Section 4: 지역별 활동 현황 ──────────────────────
    add_heading(doc, "4. 지역(Province)별 활동 현황", 1, theme_color, 13)
    add_hr(doc, theme_hex)
    render_province_section(doc, arts, plan_id, kpi_db, theme_hex)
    doc.add_page_break()

    # ── Section 5: 한국 기업 기회 ────────────────────────
    add_heading(doc, "5. 한국 기업 기회 분석", 1, theme_color, 13)
    add_hr(doc, theme_hex)
    render_korean_opportunity(doc, cfg["korean_opportunities"], theme_hex)
    doc.add_page_break()

    # ── Appendix ─────────────────────────────────────────
    build_appendix(doc, arts, theme_hex)

    # ── 저장 ─────────────────────────────────────────────
    safe_id = plan_id.replace("-", "_")
    fname = f"MI_REPORT_{safe_id}_{week_str.replace('/','_')}_{now.strftime('%Y%m%d')}.docx"
    fpath, sz = save_report(doc, output_dir, fname)
    return fpath


# ─────────────────────────────────────────────────────────
# 일괄 생성
# ─────────────────────────────────────────────────────────
def generate_all_regional(output_dir: str) -> list:
    results = []
    for plan_id in PLAN_CONFIG:
        print(f"  📄 {plan_id} ...", end=" ", flush=True)
        try:
            fpath = generate_regional_report(plan_id, output_dir)
            sz = os.path.getsize(fpath) // 1024
            print(f"✅ {sz}KB  [{fpath.split('/')[-1]}]")
            results.append((plan_id, fpath, sz))
        except Exception as e:
            print(f"❌ {e}")
            results.append((plan_id, None, 0))
    return results


if __name__ == "__main__":
    import sys
    sys.path.insert(0, "/home/work/claw/scripts")
    out = "/home/work/claw/outputs/reports/MI_Reports_v4"
    Path(out).mkdir(parents=True, exist_ok=True)

    plan = sys.argv[1] if len(sys.argv) > 1 else "VN-TRAN-2055"
    print(f"지역성 보고서 생성: {plan}")
    fpath = generate_regional_report(plan, out)
    print(f"✅ {fpath}  ({os.path.getsize(fpath)//1024}KB)")

"""
Water 통합 MI 보고서 생성기 — Vietnam Infrastructure Hub
=========================================================
PDP8 방식 통합 구조: 3개 하부 트래킹 → 단일 보고서
  표지  : 베트남 국가 수자원·상수도 정책 통합 개요
  Sec 1 : 수자원 마스터플랜 2021-2030/2050 (Decision 1622)
  Sec 2 : 도시 상수도 인프라 계획 2025/2035 (MOC)
  Sec 3 : 농촌 상수도·위생 국가전략 2030/2045 (Decision 1978)
  Sec 4 : Province 활동도 (3개 하부 교차 분석)
  Sec 5 : 한국 기업 기회 분석
  Appendix: 전체 기사 목록
"""
import os, sys, json, re
from datetime import datetime
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR   = Path("/home/work/claw")
OUTPUT_DIR = Path(os.getenv("PIPELINE_OUTPUT_DIR", "/home/work/claw/outputs"))

# Water 색상 팔레트
COLOR_MAIN    = (0x06, 0xB6, 0xD4)   # cyan     — 표지
COLOR_RESRC   = (0x06, 0x96, 0xC4)   # dark cyan — 수자원
COLOR_URBAN   = (0x22, 0x78, 0xA4)   # blue     — 도시 상수도
COLOR_RURAL   = (0x16, 0xA3, 0x4A)   # green    — 농촌 상수도
COLOR_PROV    = (0x7C, 0x3A, 0xED)   # violet   — Province
COLOR_KOR     = (0xDC, 0x26, 0x26)   # red      — 한국 기회

HEX = {
    "main":  "06B6D4", "resrc": "0696C4", "urban": "2278A4",
    "rural": "16A34A", "prov":  "7C3AED", "kor":   "DC2626",
    "light": "E0F2FE",
}

SUB_TRACKS = [
    {
        "id":      "VN-WAT-RESOURCES",
        "label":   "국가 수자원 마스터플랜 2021-2030/2050",
        "emoji":   "🌊",
        "color":   COLOR_RESRC,
        "hex":     HEX["resrc"],
        "legal":   "Decision 1622/QD-TTg, December 27, 2022",
        "lead":    "MONRE (환경자원부)",
        "nature":  "수량·수질·유역 관리 및 국가 수안보 계획 — 공급 인프라 건설 계획 아님",
        "kpi_2030": "총 수요 1,225억 m³/년 충족 | 도시 95-100% 청정수 | 농촌 65% | 수자원 개발 90% 통제 | 손실률 10%",
        "kpi_2050": "세계 수준 수안보 국가 그룹 진입 | 총 수요 1,317억 m³/년 | 모든 상황 선제 대응",
        "key_projects": [
            "13개 주요 하천유역 통합 관리 계획 수립",
            "홍강 하류 오염 하천 복원 (Nhue-Day-Bac Hung Hai)",
            "메콩델타 담수 저장·염수차단 인프라 강화",
            "중남부 가뭄 지역 유역간 용수 이송 최적화",
            "국가 실시간 수자원 모니터링·예보 디지털 시스템",
            "11개 주요 유역 저수지 연계운영 AI 최적화",
        ],
    },
    {
        "id":      "VN-WAT-URBAN",
        "label":   "도시 상수도 인프라 계획 2025/2035",
        "emoji":   "🏙️",
        "color":   COLOR_URBAN,
        "hex":     HEX["urban"],
        "legal":   "Decision 1929/QD-TTg (2009) + Decision 1566/QD-TTg (2016) + 2026-2035 신규 프로그램 준비 중",
        "lead":    "MOC (건설부)",
        "nature":  "도시·산업단지 상수도 인프라 — 정수장·배관망·안전급수·누수저감",
        "kpi_2025_achieved": "도시 상수도 커버리지 ~95% | 평균 손실률 ~15% | 31/63개 성시 안전급수계획 수립",
        "kpi_2030": "집중급수 시스템 90% 안전급수계획 수립·이행 | 손실률 <15%(집중) / <20%(분산) | 100% 비상계획",
        "kpi_2035": "집중급수 100% 안전급수계획 | 분산 80% | 손실률 <13%(집중) / <18%(분산)",
        "kpi_2050": "모든 도시·산업단지 안정적 고품질 상수도 100% 공급",
        "key_projects": [
            "노후 배관 교체 — 하노이·HCM시티 우선 (손실률 15%→13% 달성)",
            "스마트 수도 시스템 — IoT 누수감지·압력관리 전국 확대",
            "하노이: Red River 지표수 취수 강화 + 지하수 의존도 감소",
            "HCM시티: Thu Duc 수처리장 확장 + Dong Nai 원수원 보호",
            "메콩델타 도시: 염수침입 대응 백업 상수원 확보 (Can Tho, Ca Mau)",
            "법적 기반: 수도공급·배수법 제정 (MOC 추진)",
        ],
    },
    {
        "id":      "VN-WAT-RURAL",
        "label":   "농촌 상수도·위생 국가전략 2030/2045",
        "emoji":   "🌾",
        "color":   COLOR_RURAL,
        "hex":     HEX["rural"],
        "legal":   "Decision 1978/QD-TTg, November 24, 2021",
        "lead":    "MARD (농업농촌개발부) + UNICEF·World Bank 지원",
        "nature":  "6,200만+ 농촌 인구 상수도·위생 — 소수민족·오지 지역 중점",
        "baseline_2021": "청정수 접근 51% | 위생시설 75% | 수원 대장균 오염 44%",
        "kpi_2030": "안전 관리 청정수·위생 접근 보장 | 가축폐수 100% 처리",
        "kpi_2045": "농촌 인구 100% 안전 청정수 | 농촌 거주지 50% 하수수집 | 생활하수 30% 처리",
        "priority_regions": [
            "중부 고원 소수민족 (Gia Lai, Kon Tum, Dak Lak)",
            "북부 산악 소수민족 (Lai Chau, Dien Bien, Ha Giang)",
            "중남부 만성가뭄 (Ninh Thuan, Binh Thuan)",
            "메콩델타 염수침입 (Ca Mau, Kien Giang, Bac Lieu)",
            "외딴 섬 (Phu Quoc, Con Dao, Ly Son)",
        ],
        "key_projects": [
            "소수민족 오지 분산급수 시스템 설치 확대",
            "기후탄력적 농촌 급수 (내염성 시스템, 빗물 집수)",
            "농촌 위생 시설 보급 (SDG 6 목표)",
            "UNICEF/World Bank ODA 활용 WASH 프로그램",
        ],
    },
]


# ── 유틸 ─────────────────────────────────────────────────────
def _rgb(t): return RGBColor(*t)
def _hex_bg(cell, h):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),h)
    p.append(s)
def _hr(doc, color="CCCCCC"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'), color)
    pb.append(bot); pPr.append(pb)
def _h(doc, text, lvl, color=COLOR_MAIN, size=None):
    h = doc.add_heading(text, lvl)
    for r in h.runs:
        r.font.color.rgb = _rgb(color)
        if size: r.font.size = Pt(size)
    return h
def _date_key(a):
    d = str(a.get('published_date','') or '')
    try:
        if re.match(r'\d{4}-\d{2}-\d{2}', d): return d[:10]
        return datetime.strptime(d.strip(), '%b %d, %Y').strftime('%Y-%m-%d')
    except: return d[:10] if len(d) >= 10 else '1900-01-01'
def _is_vi(t):
    return any(c in (t or '') for c in ['ă','ơ','ư','đ','ấ','ề','ộ','ừ','ạ','ọ','ổ','ị','ế'])
def _title(a):
    t = a.get('title','')
    if _is_vi(t):
        en = a.get('summary_en','')
        if en and not _is_vi(en): return f"[VI] {en[:80]}"
        ko = a.get('summary_ko','')
        if ko and not _is_vi(ko): return f"[VI] {ko[:70]}"
    return t[:85]
def _summary(a):
    ko = a.get('summary_ko','')
    en = a.get('summary_en','')
    if ko and not _is_vi(ko): return '🇰🇷', ko[:200]
    if en and not _is_vi(en): return '🇺🇸', en[:200]
    return None, None
def _select_articles(arts):
    s = sorted(arts, key=_date_key, reverse=True)
    for yr in ['2026','2025']:
        y = [a for a in s if yr in str(a.get('published_date',''))]
        if y: return y, yr
    return s[:20], "전체"
def _relevance(a, tid):
    must = {
        "VN-WAT-RESOURCES": ["water resources","river basin","water security","flood control",
                              "drought","saltwater intrusion","dam","reservoir","irrigation",
                              "groundwater","tài nguyên nước","lưu vực sông","hạn hán",
                              "xâm nhập mặn","lũ lụt","đập","hồ chứa"],
        "VN-WAT-URBAN":     ["water supply plant","water treatment plant","clean water supply",
                              "water pipeline","water network","water loss","safe water",
                              "nhà máy nước","cấp nước đô thị","nước sạch đô thị","đường ống cấp nước"],
        "VN-WAT-RURAL":     ["rural water supply","rural sanitation","rural clean water",
                              "commune water","ethnic minority water",
                              "cấp nước nông thôn","vệ sinh nông thôn","nước sạch nông thôn"],
    }
    text = ((a.get('title','') or '') + ' ' +
            (a.get('summary_en','') or '') + ' ' +
            (a.get('summary_ko','') or '') + ' ' +
            (a.get('content','') or '')[:300]).lower()
    return any(kw in text for kw in must.get(tid, []))

def _add_articles(doc, arts, color, max_n=5):
    for i, art in enumerate(arts[:max_n], 1):
        t = _title(art); flag, sm = _summary(art); date = _date_key(art)
        src = (art.get('source','') or '')[:30]
        ap = doc.add_paragraph(); ap.paragraph_format.left_indent = Inches(0.2)
        nr = ap.add_run(f"[{i}] "); nr.bold=True; nr.font.size=Pt(10); nr.font.color.rgb=_rgb(color)
        ap.add_run(t).bold=True; ap.runs[-1].font.size=Pt(10)
        mp = doc.add_paragraph(); mp.paragraph_format.left_indent = Inches(0.4)
        mp.add_run(f"📅 {date}  |  📰 {src}").font.size=Pt(8.5)
        if flag and sm:
            sp = doc.add_paragraph(); sp.paragraph_format.left_indent = Inches(0.4)
            sp.add_run(f"{flag} ").font.size=Pt(9.5); sp.add_run(sm).font.size=Pt(9.5)
            en = art.get('summary_en','')
            if flag == '🇰🇷' and en and not _is_vi(en):
                ep = doc.add_paragraph(); ep.paragraph_format.left_indent = Inches(0.4)
                ep.add_run("🇺🇸 ").font.size=Pt(9.5); ep.add_run(en[:180]).font.size=Pt(9.5)
        doc.add_paragraph()


def generate_water_report(output_dir: str) -> str:
    with open(BASE_DIR/"config/history_db.json", encoding='utf-8') as f:
        hdb = json.load(f)
    with open(BASE_DIR/"config/water_structure.json", encoding='utf-8') as f:
        wstruct = json.load(f)

    now  = datetime.now(); week = now.isocalendar()[1]

    # 트래킹별 기사 버킷
    buckets = defaultdict(list)
    for art in hdb['articles'].values():
        for pid in (art.get('matched_plans') or []):
            if pid.startswith('VN-WAT-'):
                buckets[pid].append(art)
    total_arts = sum(len(v) for v in buckets.values())

    doc = Document()
    for sec in doc.sections:
        sec.top_margin=Inches(1.0); sec.bottom_margin=Inches(1.0)
        sec.left_margin=Inches(1.2); sec.right_margin=Inches(1.2)

    # ══ COVER ══
    cov = doc.add_table(rows=1,cols=1); cov.style="Table Grid"
    c = cov.rows[0].cells[0]; _hex_bg(c,"0F172A")
    p = c.paragraphs[0]
    r = p.add_run("VIETNAM INFRASTRUCTURE INTELLIGENCE HUB")
    r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    tp = doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2 = tp.add_run("베트남 국가 수자원·상수도 통합 정책")
    r2.bold=True; r2.font.size=Pt(20); r2.font.color.rgb=_rgb(COLOR_MAIN)

    sp = doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run("Vietnam National Water Sector — Integrated Master Plans (수자원 + 도시상수도 + 농촌급수)").font.size=Pt(10)
    doc.add_paragraph()

    it = doc.add_table(rows=5,cols=2); it.style="Table Grid"
    for i,(k,v) in enumerate([
        ("수자원 마스터플랜",   "Decision 1622/QD-TTg (2022.12.27) — MONRE"),
        ("도시 상수도 계획",   "Decision 1929(2009) + 1566(2016) + 2026-2035 신규 — MOC"),
        ("농촌 상수도·위생",  "Decision 1978/QD-TTg (2021.11.24) — MARD / UNICEF"),
        ("총 기사",           f"{total_arts}건 (3개 트래킹 합산)"),
        ("핵심 목표",         "도시 100% + 농촌 100%(2045) 안전 청정수 | 국가 수안보 달성"),
    ]):
        it.rows[i].cells[0].text=k; it.rows[i].cells[1].text=v
        _hex_bg(it.rows[i].cells[0], HEX["main"])
        it.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        it.rows[i].cells[0].paragraphs[0].runs[0].bold=True
        for cell in it.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(10)
    doc.add_paragraph()

    # 통합 구조 노트
    note_t = doc.add_table(rows=1,cols=1); note_t.style="Table Grid"
    nc = note_t.rows[0].cells[0]; _hex_bg(nc, "E0F2FE")
    np_ = nc.paragraphs[0]
    np_.add_run("ℹ️  보고서 구조 안내: ").bold=True
    np_.add_run(
        "이 보고서는 수자원 관리(유역·수안보) → 도시 상수도(정수·배관) → 농촌 급수(오지·위생)의 "
        "3개 하위 정책을 하나의 국가 물 공급 체계로 통합 추적합니다. "
        "VN-WW-2030(하수처리)은 별도 보고서에서 추적됩니다."
    ).font.size=Pt(9.5)
    doc.add_paragraph()
    doc.add_page_break()

    # ══ SECTIONS 1–3: 하부 트래킹별 ══
    for sec_no, track in enumerate(SUB_TRACKS, 1):
        tid = track["id"]; color = track["color"]; hexc = track["hex"]
        arts = buckets.get(tid, [])
        rel  = [a for a in arts if _relevance(a, tid)]
        rep, yr = _select_articles(rel) if rel else _select_articles(arts)

        _h(doc, f"{sec_no}. {track['emoji']}  {track['label']}", 1, color, 13)
        _hr(doc, hexc)

        # 통계 배지
        stat_t = doc.add_table(rows=1,cols=4); stat_t.style="Table Grid"
        for j,(k,v) in enumerate([
            ("결정문", track["legal"][:40]),
            ("역사 기사", f"{len(arts)}건"),
            ("관련성 통과", f"{len(rel)}건"),
            (f"보고({yr})", f"{len(rep)}건"),
        ]):
            stat_t.rows[0].cells[j].text=f"{k}\n{v}"
            _hex_bg(stat_t.rows[0].cells[j], hexc)
            for para in stat_t.rows[0].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size=Pt(9); run.bold=(j==0)
                    run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        doc.add_paragraph()

        # 성격 설명
        nature_p = doc.add_paragraph(); nature_p.paragraph_format.left_indent=Inches(0.2)
        nature_p.add_run("【성격】 ").bold=True; nature_p.runs[0].font.color.rgb=_rgb(color)
        nature_p.runs[0].font.size=Pt(10)
        nature_p.add_run(track["nature"]).font.size=Pt(10)
        doc.add_paragraph()

        # KPI 테이블
        kpi_tbl = doc.add_table(rows=1, cols=2); kpi_tbl.style="Table Grid"
        # 각 트래킹별 KPI 행 구성
        kpi_rows = []
        if tid == "VN-WAT-RESOURCES":
            kpi_rows = [
                ("2030 목표", track["kpi_2030"]),
                ("2050 비전", track["kpi_2050"]),
                ("13대 하천유역", "Bang Giang-Ky Cung | Red-Thai Binh (홍강) | Ma | Ca | Huong | Vu Gia-Thu Bon | Tra Khuc | Kon-Ha Thanh | Ba | Se San | Srepok | Dong Nai | Mekong"),
                ("우선 과제", "메콩델타 염수침입 | 홍강 지하수 고갈 | 동나이강 오염 | 중남부 만성가뭄 | 국제 수계(메콩강 상류댐) 협력"),
            ]
        elif tid == "VN-WAT-URBAN":
            kpi_rows = [
                ("2025년 현황", track["kpi_2025_achieved"]),
                ("2030 목표", track["kpi_2030"]),
                ("2035 목표", track["kpi_2035"]),
                ("2050 비전", track["kpi_2050"]),
            ]
        elif tid == "VN-WAT-RURAL":
            kpi_rows = [
                ("2021 기준선", track["baseline_2021"]),
                ("2030 목표", track["kpi_2030"]),
                ("2045 비전", track["kpi_2045"]),
                ("우선 지역", " | ".join(track["priority_regions"][:3])),
            ]

        for k, v in kpi_rows:
            row = kpi_tbl.add_row()
            row.cells[0].text = k; row.cells[1].text = v
            _hex_bg(row.cells[0], hexc)
            row.cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            row.cells[0].paragraphs[0].runs[0].bold=True
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.font.size=Pt(9.5)
        doc.add_paragraph()

        # 주요 프로젝트
        _h(doc, "주요 프로젝트 및 정책 과제", 2, color)
        for proj in track["key_projects"]:
            p = doc.add_paragraph(f"▪ {proj}", style="List Bullet")
            p.runs[0].font.size=Pt(10)
        doc.add_paragraph()

        # 최신 뉴스
        if rep:
            _h(doc, f"최신 뉴스 ({yr}, 상위 5건)", 2, color)
            _add_articles(doc, rep, color, max_n=5)
            if len(rep) > 5:
                mp = doc.add_paragraph(); mp.paragraph_format.left_indent=Inches(0.2)
                mp.add_run(f"  ※ 추가 {len(rep)-5}건 → Appendix 참조").font.size=Pt(9)
        else:
            np2 = doc.add_paragraph(); np2.paragraph_format.left_indent=Inches(0.2)
            np2.add_run(f"ℹ️ {yr}년 관련 기사 수집 중 (키워드 확장 예정)").font.size=Pt(9.5)

        _hr(doc,"D1D5DB"); doc.add_paragraph()
        doc.add_page_break()

    # ══ SECTION 4: Province 교차 분석 ══
    _h(doc, "4. Province 활동도 — 수자원·상수도 교차 분석", 1, COLOR_PROV, 13); _hr(doc, HEX["prov"])

    prov_track = defaultdict(lambda: defaultdict(int))
    prov_norm = {'vietnam':'전국','ho chi minh city':'Ho Chi Minh City','hcmc':'Ho Chi Minh City',
                 'ha noi':'Hanoi','hanoi':'Hanoi'}
    for track in SUB_TRACKS:
        for art in buckets.get(track["id"],[]):
            pv = (art.get('province','') or '전국').strip()
            pv = prov_norm.get(pv.lower(), pv) or '전국'
            prov_track[pv][track["id"]] += 1

    top_provs = sorted(prov_track.keys(), key=lambda p: sum(prov_track[p].values()), reverse=True)[:15]

    if top_provs:
        pt = doc.add_table(rows=1, cols=5); pt.style="Table Grid"
        for j,(h,hexc) in enumerate([("Province","prov"),("합계","prov"),
                                      ("🌊 수자원","resrc"),("🏙️ 도시급수","urban"),("🌾 농촌급수","rural")]):
            pt.rows[0].cells[j].text=h
            _hex_bg(pt.rows[0].cells[j], HEX[hexc])
            pt.rows[0].cells[j].paragraphs[0].runs[0].bold=True
            pt.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            pt.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)
        for prov in top_provs:
            row = pt.add_row(); row.cells[0].text=prov
            total = sum(prov_track[prov].values()); row.cells[1].text=str(total)
            for j,tid in enumerate(["VN-WAT-RESOURCES","VN-WAT-URBAN","VN-WAT-RURAL"]):
                cnt = prov_track[prov].get(tid,0)
                row.cells[j+2].text = str(cnt) if cnt else "—"
                if cnt >= 5: _hex_bg(row.cells[j+2], ["resrc","urban","rural"][j] and HEX[["resrc","urban","rural"][j]])
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.font.size=Pt(9)
    doc.add_paragraph()
    doc.add_page_break()

    # ══ SECTION 5: 한국 기업 기회 ══
    _h(doc, "5. 한국 기업 기회 분석", 1, COLOR_KOR, 13); _hr(doc, HEX["kor"])

    kor_t = doc.add_table(rows=1,cols=2); kor_t.style="Table Grid"
    for j,h in enumerate(["분야","기회"]):
        kor_t.rows[0].cells[j].text=h
        kor_t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _hex_bg(kor_t.rows[0].cells[j], HEX["kor"])
        kor_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

    for sec, opp in [
        ("🌊 수자원 관리",
         "HIGH: 유역 관리 디지털 시스템 (AI 저수지 운영) — 한국 K-water·WAMIS 기술 수출\n"
         "HIGH: 홍강·동나이강 오염 복원 사업 — 한국 4대강 기술 활용\n"
         "MEDIUM: 메콩델타 염수차단 수문·배수 인프라 설계·시공\n"
         "MEDIUM: 국제 수계 모니터링 시스템 — 위성·IoT 센서 네트워크"),
        ("🏙️ 도시 상수도",
         "HIGH: 노후 배관 교체 PE파이프·스마트 계측기 공급 — LS전선·동국S&C 등\n"
         "HIGH: 스마트 워터 시스템 (IoT 누수감지·압력관리) — SK텔레콤·KT 솔루션 수출\n"
         "HIGH: 정수장 고도처리 설비 — 막여과(MBR) 기술 공급\n"
         "MEDIUM: 수도요금 IT 관리 시스템 — 한국 수도 billing 시스템 수출\n"
         "투자 규모: 도시 상수도 $10B+ (2026-2035)"),
        ("🌾 농촌 상수도",
         "MEDIUM: 농촌 분산급수 시스템 패키지 (소형 정수+태양광 펌프) — K-water ODA 연계\n"
         "MEDIUM: 기후탄력적 농촌 급수 기술 (내염성 RO, 빗물집수) — KOICA 사업 연계\n"
         "MEDIUM: 소수민족 오지 위성 급수 시스템\n"
         "채널: KOICA·ADB·World Bank ODA 사업 수주"),
        ("📊 전략적 포인트",
         "① K-water: 수자원 마스터플랜 기술지원 + 디지털 유역 관리 → 컨설팅 수출\n"
         "② MOC/MOIT와 공동 스마트 워터 시범사업 → 규모 확대 (Hanoi·HCM 스마트시티 연계)\n"
         "③ ODA 연계 농촌급수: KOICA + 현지 건설사 패키지 수주 모델\n"
         "④ ADB $2.4B 베트남 수자원 투자 계획 — 한국 ADB 의결권 활용 수주 전략\n"
         "⑤ 주요 창구: MONRE(수자원) + MOC(도시상수도) + MARD(농촌급수) 별도 접근 필요"),
    ]:
        row = kor_t.add_row()
        row.cells[0].text=sec; row.cells[1].text=opp
        _hex_bg(row.cells[0], "FEF2F2"); row.cells[0].paragraphs[0].runs[0].bold=True
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(9.5)
    doc.add_paragraph()
    doc.add_page_break()

    # ══ APPENDIX ══
    _h(doc, "Appendix — 전체 관련 기사 목록", 1, COLOR_MAIN)
    _hr(doc, HEX["main"])

    seen = set(); all_flat = []
    for track in SUB_TRACKS:
        for art in buckets.get(track["id"],[]):
            url = art.get('url','') or str(id(art))
            if url not in seen:
                seen.add(url); art['_track'] = track["emoji"]+track["label"][:8]
                all_flat.append(art)
    all_flat.sort(key=_date_key, reverse=True)

    ap_note = doc.add_paragraph(f"총 {len(all_flat)}건 (3개 트래킹 중복 제거) — 상위 60건 표시")
    ap_note.runs[0].font.size=Pt(9); ap_note.runs[0].font.color.rgb=RGBColor(0x64,0x74,0x8B)
    doc.add_paragraph()

    ap_t = doc.add_table(rows=1,cols=5); ap_t.style="Table Grid"
    for j,h in enumerate(["No","날짜","제목","출처","분야"]):
        ap_t.rows[0].cells[j].text=h
        ap_t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _hex_bg(ap_t.rows[0].cells[j], "1E40AF")
        ap_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        ap_t.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)
    for ri, art in enumerate(all_flat[:60], 1):
        row = ap_t.add_row()
        row.cells[0].text=str(ri); row.cells[1].text=_date_key(art)
        row.cells[2].text=_title(art)[:60]; row.cells[3].text=(art.get('source','') or '')[:18]
        row.cells[4].text=art.get('_track','')[:15]
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(8.5)

    # 저장
    fname = f"MI_REPORT_WATER_INTEGRATED_W{week:02d}_{now.strftime('%Y%m%d')}.docx"
    fpath = os.path.join(output_dir, fname)
    doc.save(fpath)
    return fpath


if __name__ == "__main__":
    out = "/home/work/claw/outputs/reports/MI_Reports"
    Path(out).mkdir(parents=True, exist_ok=True)
    print("Water 통합 MI 보고서 생성 중...")
    fpath = generate_water_report(out)
    if fpath:
        sz = os.path.getsize(fpath) // 1024
        print(f"✅ {fpath} ({sz}KB)")

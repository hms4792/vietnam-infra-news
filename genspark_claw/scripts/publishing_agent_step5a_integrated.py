"""
Publishing Agent Step 5-A — Vietnam Infrastructure Intelligence Hub
Excel(18시트) + Word 보고서(마스터플랜별) + Executive Summary
원본 Genspark Agent 형식 완전 호환 (v3.0)
"""
import os, sys, json
from datetime import datetime, timedelta
from collections import Counter
from pathlib import Path

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 마스터플랜 정의 (18개, 한글명 포함) ──────────────────────────────
MASTERPLANS = [
    {"id": "VN-PWR-PDP8",       "short_name": "PDP8",        "name_ko": "베트남 제8차 전력개발계획 (PDP8)",          "color": "F59E0B"},
    {"id": "VN-ENV-IND-1894",   "short_name": "Decision1894","name_ko": "베트남 환경산업 발전 프로그램 (Decision 1894)", "color": "10B981"},
    {"id": "VN-TRAN-2055",      "short_name": "Transport2055","name_ko": "베트남 교통인프라 마스터플랜 2055",           "color": "3B82F6"},
    {"id": "VN-URB-METRO-2030", "short_name": "Metro2030",   "name_ko": "베트남 도시철도·메트로 개발계획 2030",         "color": "8B5CF6"},
    {"id": "VN-GAS-PDP8",       "short_name": "GasPDP8",     "name_ko": "베트남 가스 인프라 개발계획 (PDP8 연계)",       "color": "EF4444"},
    {"id": "VN-WAT-2050",       "short_name": "Water2050",   "name_ko": "베트남 수자원 개발 마스터플랜 2050",           "color": "06B6D4"},
    {"id": "VN-REN-NPP-2050",   "short_name": "Nuclear2050", "name_ko": "베트남 원자력·신에너지 개발계획 2050",          "color": "6366F1"},
    {"id": "VN-COAL-RETIRE",    "short_name": "CoalRetire",  "name_ko": "베트남 석탄화력 단계적 폐지계획 (JETP)",        "color": "78716C"},
    {"id": "VN-GRID-SMART",     "short_name": "SmartGrid",   "name_ko": "베트남 스마트그리드 고도화 계획",               "color": "0EA5E9"},
    {"id": "VN-EV-2030",        "short_name": "EV2030",      "name_ko": "베트남 전기차·친환경 모빌리티 2030",            "color": "22C55E"},
    {"id": "VN-CARBON-2050",    "short_name": "Carbon2050",  "name_ko": "베트남 탄소중립 2050 로드맵",                 "color": "84CC16"},
    {"id": "VN-LNG-HUB",        "short_name": "LNGHub",      "name_ko": "베트남 LNG 허브 개발 전략",                  "color": "F97316"},
    {"id": "VN-WW-2030",        "short_name": "Wastewater",  "name_ko": "베트남 국가 수처리 마스터플랜 2021-2030",        "color": "14B8A6"},
    {"id": "VN-WS-NORTH-2030",  "short_name": "WaterSupply", "name_ko": "북부 상수도 인프라 계획 2030",                "color": "38BDF8"},
    {"id": "VN-SW-MEKONG-2030", "short_name": "Mekong",      "name_ko": "메콩 델타 폐수처리 네트워크 2025-2030",         "color": "2DD4BF"},
    {"id": "VN-SC-2030",        "short_name": "SmartCity",   "name_ko": "베트남 스마트시티 개발 국가전략 2030",           "color": "A78BFA"},
    {"id": "VN-IP-NORTH-2030",  "short_name": "IndPark",     "name_ko": "북부 산업단지 개발 계획 2030",                "color": "FB923C"},
    {"id": "VN-OG-2030",        "short_name": "OilGas",      "name_ko": "베트남 석유가스 개발 계획 2030",               "color": "94A3B8"},
]

def _hex(h):
    """hex color → openpyxl PatternFill color"""
    return h.replace("#","")

def _matched_articles(articles, plan_id):
    """agent_pipeline의 matched_plans 리스트 기반 매칭 (policy_context 폴백 포함)"""
    result = []
    for a in articles:
        mp = a.get("matched_plans", [])
        pc = a.get("policy_context", {}) or {}
        if plan_id in mp or pc.get("plan_id") == plan_id:
            result.append(a)
    return result

def _style_header_row(ws, row, headers, bg_hex, text_hex="FFFFFF"):
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row, col, header)
        cell.font = Font(bold=True, color=text_hex, size=10)
        cell.fill = PatternFill(start_color=bg_hex, end_color=bg_hex, fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        thin = Side(style="thin", color="FFFFFF")
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)

def _autofit(ws, min_w=8, max_w=60):
    for col in ws.columns:
        best = min_w
        for cell in col:
            try:
                v = str(cell.value or "")
                best = max(best, min(len(v.split("\n")[0]) + 2, max_w))
            except:
                pass
        ws.column_dimensions[get_column_letter(col[0].column)].width = best

# ════════════════════════════════════════════════════════════════
# Excel 데이터베이스 (18개 시트 + All_Articles)
# ════════════════════════════════════════════════════════════════
def create_excel_database(articles, output_path):
    print(f"📊 Excel 데이터베이스 생성 중... ({len(articles)}개 기사)")
    wb = openpyxl.Workbook()

    # 공통 헤더
    HEADERS = [
        "No", "제목 (Title)", "URL", "발행일", "출처 (Source)",
        "섹터 (Sector)", "지역 (Province)", "매칭 마스터플랜",
        "요약 🇰🇷 한국어", "요약 🇺🇸 영어", "요약 🇻🇳 베트남어",
        "QC", "Genspark 섹터"
    ]
    COL_W = {2:52, 3:40, 9:45, 10:45, 11:40}

    def write_row(ws, row_idx, art, row_num):
        mp = art.get("matched_plans", [])
        pc = art.get("policy_context", {}) or {}
        plan_ids = mp if mp else ([pc.get("plan_id","")] if pc.get("plan_id") else [])
        plan_names = []
        for pid in plan_ids:
            meta = next((p for p in MASTERPLANS if p["id"]==pid), None)
            plan_names.append(meta["name_ko"] if meta else pid)

        vals = [
            row_num,
            art.get("title",""),
            art.get("url",""),
            (art.get("published_date","") or "")[:10],
            art.get("source",""),
            art.get("sector",""),
            art.get("province",""),
            "\n".join(plan_names) if plan_names else "No Match",
            art.get("summary_ko",""),
            art.get("summary_en",""),
            art.get("summary_vi",""),
            art.get("qc_status","PASS"),
            art.get("sector_genspark", art.get("sector","")),
        ]
        for col, val in enumerate(vals, 1):
            cell = ws.cell(row_idx, col, val)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if col == 3:  # URL
                cell.font = Font(color="1D4ED8", underline="single")
            if col == 12:  # QC
                cell.font = Font(
                    bold=True,
                    color="059669" if val=="PASS" else "DC2626"
                )
        # 행 높이
        ws.row_dimensions[row_idx].height = 60

    # ── Sheet 1: All_Articles ──
    ws_all = wb.active
    ws_all.title = "All_Articles"
    ws_all.freeze_panes = "A2"
    _style_header_row(ws_all, 1, HEADERS, "1E40AF")
    ws_all.row_dimensions[1].height = 20
    for i, art in enumerate(articles, 2):
        write_row(ws_all, i, art, i-1)
    for col, w in COL_W.items():
        ws_all.column_dimensions[get_column_letter(col)].width = w
    print(f"   ✓ All_Articles ({len(articles)}행)")

    # ── Sheet 2~19: 18개 마스터플랜별 ──
    for plan in MASTERPLANS:
        ws = wb.create_sheet(title=plan["short_name"][:31])
        ws.freeze_panes = "A2"
        bg = plan["color"].replace("#","")
        _style_header_row(ws, 1, HEADERS, bg)
        ws.row_dimensions[1].height = 20

        # 플랜 설명 헤더 (2행)
        ws.insert_rows(1)
        merged_cell = ws.cell(1, 1, f"📋 {plan['name_ko']}  |  Plan ID: {plan['id']}")
        merged_cell.font = Font(bold=True, size=12, color="FFFFFF")
        merged_cell.fill = PatternFill(start_color=bg, end_color=bg, fill_type="solid")
        merged_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(HEADERS))
        ws.row_dimensions[1].height = 22

        matched = _matched_articles(articles, plan["id"])
        for i, art in enumerate(matched, 3):
            write_row(ws, i, art, i-2)
        for col, w in COL_W.items():
            ws.column_dimensions[get_column_letter(col)].width = w
        print(f"   ✓ {plan['short_name']:12s} ({len(matched):3d}건)")

    wb.save(output_path)
    print(f"✅ Excel 저장: {output_path}")
    return output_path


# ════════════════════════════════════════════════════════════════
# Word 보고서 (마스터플랜별 개별)
# ════════════════════════════════════════════════════════════════
def _add_colored_heading(doc, text, level, hex_color):
    h = doc.add_heading(text, level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(hex_color)
    return h

def _add_article_block(doc, i, art, plan_id):
    """기사 1건 블록 — 원본 Genspark Agent 형식"""
    pc = art.get("policy_context",{}) or {}
    mp = art.get("matched_plans",[])
    score = pc.get("score", "—")

    # 기사 번호 + 제목
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] {art.get('title','(제목없음)')}")
    run.bold = True
    run.font.size = Pt(11)

    # 메타 테이블
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.autofit = True
    cells = tbl.rows[0].cells
    cells[0].text = f"📅 {(art.get('published_date','') or '')[:10]}"
    cells[1].text = f"📰 {art.get('source','')}"
    cells[2].text = f"🏙️ {art.get('province','')}"
    cells[3].text = f"🎯 매칭점수: {score}"
    for cell in cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    # 요약 3개국어
    def add_summary(label, key, flag):
        txt = art.get(key,"")
        if txt:
            p2 = doc.add_paragraph()
            r = p2.add_run(f"{flag} {label}: ")
            r.bold = True
            r.font.size = Pt(10)
            p2.add_run(txt).font.size = Pt(10)
            p2.paragraph_format.left_indent = Inches(0.3)

    add_summary("한국어 요약", "summary_ko", "🇰🇷")
    add_summary("English Summary", "summary_en", "🇺🇸")
    add_summary("Tóm tắt tiếng Việt", "summary_vi", "🇻🇳")

    # URL
    p3 = doc.add_paragraph()
    r = p3.add_run("🔗 URL: ")
    r.bold = True
    r.font.size = Pt(9)
    r2 = p3.add_run(art.get("url",""))
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    p3.paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph()
    # 구분선
    p4 = doc.add_paragraph("─" * 80)
    p4.runs[0].font.size = Pt(7)
    p4.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    doc.add_paragraph()

def create_project_report(plan, articles, output_dir):
    print(f"📄 {plan['short_name']} 보고서 생성 중...")

    matched = _matched_articles(articles, plan["id"])
    doc = Document()

    # ── 페이지 여백 ──
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── 표지 ──
    week_num  = datetime.now().isocalendar()[1]
    year      = datetime.now().year
    date_end  = datetime.now().strftime("%Y-%m-%d")
    date_start= (datetime.now()-timedelta(days=6)).strftime("%Y-%m-%d")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(f"🇻🇳 {plan['name_ko']}")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor.from_string(plan["color"].replace("#",""))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run(f"주간 추적 보고서  |  {year}-W{week_num:02d}  |  {date_start} ~ {date_end}").font.size = Pt(12)

    doc.add_paragraph()
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.add_run(f"Plan ID: {plan['id']}  |  총 매칭 기사: {len(matched)}건").font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph("─" * 80).runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── 1. 마스터플랜 개요 ──
    _add_colored_heading(doc, "1. 마스터플랜 개요", 1, plan["color"].replace("#",""))
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    rows_data = [
        ("마스터플랜명", plan["name_ko"]),
        ("Plan ID",     plan["id"]),
        ("보고 기간",   f"{date_start} ~ {date_end}"),
        ("매칭 기사 수", f"{len(matched)}건"),
    ]
    for r_idx, (k,v) in enumerate(rows_data):
        tbl.rows[r_idx].cells[0].text = k
        tbl.rows[r_idx].cells[1].text = v
        tbl.rows[r_idx].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # ── 2. 이번 주 관련 뉴스 ──
    _add_colored_heading(doc, f"2. 이번 주 관련 뉴스 ({len(matched)}건)", 1, plan["color"].replace("#",""))

    if matched:
        for i, art in enumerate(matched, 1):
            _add_article_block(doc, i, art, plan["id"])
    else:
        p = doc.add_paragraph("이번 주 해당 마스터플랜 관련 기사가 수집되지 않았습니다.")
        p.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # ── 3. 섹터 & 지역 분포 ──
    _add_colored_heading(doc, "3. 섹터 및 지역 분포 분석", 1, plan["color"].replace("#",""))
    if matched:
        sector_cnt = Counter(a.get("sector","기타") for a in matched)
        province_cnt = Counter(a.get("province","미분류") for a in matched)

        doc.add_heading("3-1. 섹터 분포", 2)
        for sec, cnt in sector_cnt.most_common():
            doc.add_paragraph(f"• {sec}: {cnt}건", style="List Bullet")

        doc.add_heading("3-2. 주요 지역", 2)
        for prov, cnt in province_cnt.most_common(8):
            doc.add_paragraph(f"• {prov}: {cnt}건", style="List Bullet")
    else:
        doc.add_paragraph("데이터 없음")

    # ── 4. 월간 트렌드 ──
    _add_colored_heading(doc, "4. 월간 트렌드 분석", 1, plan["color"].replace("#",""))
    doc.add_paragraph(f"이번 주({date_start} ~ {date_end}) 매칭 기사: {len(matched)}건")
    doc.add_paragraph("※ 누적 트렌드 데이터는 매주 자동 갱신됩니다.")

    # ── 5. 주요 키워드 ──
    _add_colored_heading(doc, "5. 주요 키워드", 1, plan["color"].replace("#",""))
    if matched:
        all_words = " ".join(a.get("title","") for a in matched).split()
        top_kw = [(w,c) for w,c in Counter(all_words).most_common(20) if len(w)>3][:10]
        for kw, cnt in top_kw:
            doc.add_paragraph(f"• {kw}: {cnt}회", style="List Bullet")
    else:
        doc.add_paragraph("키워드 데이터 없음")

    # ── 저장 ──
    fname = f"FINAL_{plan['short_name']}_with_history.docx"
    fpath = os.path.join(output_dir, fname)
    doc.save(fpath)
    print(f"   ✓ {fname} 저장")
    return fpath


# ════════════════════════════════════════════════════════════════
# Executive Summary
# ════════════════════════════════════════════════════════════════
def create_executive_summary(articles, qc_report, output_dir):
    print("📑 Executive Summary 생성 중...")
    doc = Document()

    for section in doc.sections:
        section.top_margin  = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    week_num = datetime.now().isocalendar()[1]
    year     = datetime.now().year
    now_str  = datetime.now().strftime("%Y-%m-%d %H:%M KST")

    # 표지
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("🇻🇳 Vietnam Infrastructure Intelligence Hub")
    tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = RGBColor(0x1E,0x3A,0x8A)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(f"Executive Summary  |  {year}-W{week_num:02d}  |  생성: {now_str}").font.size = Pt(11)
    doc.add_paragraph()

    # KPI
    _add_colored_heading(doc, "1. 주간 KPI 요약", 1, "1E3A8A")
    total   = len(articles)
    passed  = qc_report.get("passed", 0)
    matched_n = sum(1 for a in articles if a.get("matched_plans") or a.get("policy_context"))
    rate    = qc_report.get("qc_rate", 0)

    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    hdr = ["수집 기사", "QC 통과", "마스터플랜 매칭", "QC 통과율"]
    val = [str(total), str(passed), str(matched_n), f"{rate:.1f}%"]
    for i, (h,v) in enumerate(zip(hdr,val)):
        tbl.rows[0].cells[i].text = h
        tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        tbl.rows[1].cells[i].text = v
        tbl.rows[1].cells[i].paragraphs[0].runs[0].font.size = Pt(14)
        tbl.rows[1].cells[i].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # 마스터플랜별 매칭 현황
    _add_colored_heading(doc, "2. 마스터플랜별 매칭 현황 (18개)", 1, "1E3A8A")
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    for i, h in enumerate(["Plan ID","마스터플랜명","매칭 기사","섹터"]):
        tbl2.rows[0].cells[i].text = h
        tbl2.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for plan in MASTERPLANS:
        m = _matched_articles(articles, plan["id"])
        sectors = Counter(a.get("sector","") for a in m)
        top_sec = sectors.most_common(1)[0][0] if sectors else "—"
        row = tbl2.add_row()
        row.cells[0].text = plan["id"]
        row.cells[1].text = plan["name_ko"]
        row.cells[2].text = f"{len(m)}건"
        row.cells[3].text = top_sec
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # 섹터 분포
    _add_colored_heading(doc, "3. 섹터별 기사 분포", 1, "1E3A8A")
    sector_dist = qc_report.get("sector_distribution",
                  Counter(a.get("sector","기타") for a in articles))
    for sec, cnt in sorted(sector_dist.items(), key=lambda x:-x[1]):
        pct = round(cnt/total*100) if total else 0
        doc.add_paragraph(f"• {sec}: {cnt}건 ({pct}%)", style="List Bullet")
    doc.add_paragraph()

    # 상위 10 기사
    _add_colored_heading(doc, "4. 이번 주 주요 기사 Top 10", 1, "1E3A8A")
    top10 = sorted(
        [a for a in articles if a.get("qc_status")=="PASS"] or articles,
        key=lambda x: x.get("published_date",""), reverse=True
    )[:10]
    for i, art in enumerate(top10, 1):
        p = doc.add_paragraph()
        p.add_run(f"[{i}] ").bold = True
        p.add_run(art.get("title","")).font.size = Pt(10)
        ko = (art.get("summary_ko","") or "")[:150]
        if ko:
            p2 = doc.add_paragraph(f"    🇰🇷 {ko}")
            p2.paragraph_format.left_indent = Inches(0.3)
            p2.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # 저장
    fpath = os.path.join(output_dir, "Executive_Summary.docx")
    doc.save(fpath)
    print(f"✅ Executive_Summary.docx 저장")
    return fpath


# ════════════════════════════════════════════════════════════════
# main()
# ════════════════════════════════════════════════════════════════
def main(input_file=None, output_dir=None):
    print("="*60)
    print("Publishing Agent — Step 5-A Integrated (v3.0)")
    print("="*60)

    base = Path(os.getenv("PIPELINE_OUTPUT_DIR", "."))
    if input_file is None:
        for candidate in [base/"genspark_output.json", Path("genspark_output.json"),
                          base/"processed_articles.json"]:
            if candidate.exists():
                input_file = str(candidate)
                break
    if input_file is None:
        print("❌ genspark_output.json 없음")
        sys.exit(1)

    if output_dir is None:
        output_dir = str(base/"reports")
    os.makedirs(output_dir, exist_ok=True)

    with open(input_file, encoding="utf-8") as f:
        raw = json.load(f)
    articles = raw if isinstance(raw, list) else raw.get("articles", [])
    print(f"✓ 기사 로드: {len(articles)}건")

    # QC report
    for qc_cand in [base/"genspark_qc_report.json", Path("genspark_qc_report.json")]:
        if qc_cand.exists():
            with open(qc_cand, encoding="utf-8") as f:
                qc_report = json.load(f)
            break
    else:
        qc_report = {"passed": len(articles), "qc_rate": 100.0}

    # Excel
    xl_path = os.path.join(output_dir, "Vietnam_Infra_News_Database.xlsx")
    create_excel_database(articles, xl_path)

    # Word — 전체 마스터플랜 개별 보고서
    print("\n📚 마스터플랜별 Word 보고서 생성 중...")
    for plan in MASTERPLANS:
        create_project_report(plan, articles, output_dir)

    # Executive Summary
    print()
    create_executive_summary(articles, qc_report, output_dir)

    print("\n" + "="*60)
    print("✅ 모든 보고서 생성 완료!")
    print(f"📁 출력 위치: {output_dir}")
    print("="*60)

if __name__ == "__main__":
    main()

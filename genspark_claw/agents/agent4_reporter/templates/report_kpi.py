"""
report_kpi.py — 단위상품형 KPI 보고서 생성기 v4.0
==================================================
대상 플랜:
  VN-ENV-IND-1894, VN-SC-2030, VN-OG-2030,
  VN-EV-2030, VN-CARBON-2050
  + 통합형: PDP8 (6 sub), Water (3 sub)

보고서 구조:
  표지
  Section 1. 프로그램/시장 개요
  Section 2. KPI 현황표 ★ — 목표치/현재값/달성률/추세
  Section 3. 하부 트래킹 항목별 진행현황 — KPI박스 + 최신/역사기사
  Section 4. 시장 동향 분석 — 연도별 기사 트렌드 + 섹터 분포
  Section 5. 한국 기업 기회
  Appendix
"""
import os, json
from datetime import datetime
from pathlib import Path
from collections import defaultdict, Counter

from report_lib import (
    new_doc, add_heading, add_hr, add_note, add_bullet, add_label_value,
    styled_table, kpi_box, info_box, project_card, build_cover,
    render_articles, render_history_timeline, render_province_section,
    render_korean_opportunity, build_appendix, insert_map,
    load_history, load_kpi_db, save_report,
    select_articles, get_year_trend, _date_key, _is_vi,
    PALETTE, PALETTE_HEX, _rgb, cell_bg,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path("/home/work/claw")

# ─────────────────────────────────────────────────────────
# 플랜별 KPI 메타데이터
# ─────────────────────────────────────────────────────────
KPI_CONFIG = {

    "VN-ENV-IND-1894": {
        "name_ko": "환경산업 개발프로그램 (Decision 1894)",
        "name_en": "Environmental Industry Development Program (Decision 1894/QD-TTg)",
        "legal_basis": "Decision 1894/QD-TTg (September 4, 2025) — MOIT 주관",
        "lead_agency": "Ministry of Industry & Trade (MOIT)",
        "investment": "정부 R&D 예산 + 민간 매칭",
        "period": "2025 – 2030",
        "theme": ("green", "15803D", "F0FDF4"),
        "report_type": "📊 단위상품형 KPI 보고서",
        "overview": (
            "베트남 환경산업 기술 자립화 프로그램. 폐수처리·배기가스·고형폐기물 기술 국산화. "
            "2030년 폐수처리 기술 국산화율 70~80%, 배기가스 처리 60~70% 목표. "
            "⚠️ 주의: 인프라 건설 사업이 아닌 MOIT 주관 기술산업 육성 프로그램."
        ),
        "national_kpis": [
            ("폐수처리 기술 국산화", "70~80%", "2025년 기준 ~30%", "2030"),
            ("배기가스 처리 국산화",  "60~70%", "2025년 기준 ~25%", "2030"),
            ("고형폐기물 처리 국산화","50~60%", "2025년 기준 ~20%", "2030"),
            ("환경 모니터링 장비",    "20%",    "2025년 기준 ~5%",  "2030"),
            ("재활용 산업단지",       "파일럿 구축", "법적 근거 미비", "2027"),
        ],
        "sub_tracks": [
            {"id": "VN-ENV-IND-1894", "name": "환경산업 전체", "focus": "기술산업 동향 전반"},
        ],
        "market_segments": [
            ("폐수처리 장비", "연간 $500M+", "국산화율 목표 70-80%"),
            ("대기오염 방지", "연간 $200M+", "국산화율 목표 60-70%"),
            ("고형폐기물 기계", "연간 $300M+", "국산화율 목표 50-60%"),
            ("환경 모니터링", "연간 $150M+", "국산화율 목표 20%"),
            ("WtE 설비", "연간 $400M+", "기술 구축 단계"),
        ],
        "korean_opportunities": [
            ("기술이전 파트너십", "HIGH",
             "국산화 목표 달성 위한 한국 기술이전 — 세금 인센티브 + 우선 구매 보장"),
            ("JV 현지 생산", "HIGH",
             "한국 설계·베트남 생산 — MBR막·모니터링센서·집진기 현지 합작법인"),
            ("R&D 센터 설립", "MEDIUM",
             "호아락·하이테크파크 환경기술 R&D — 기술이전 + 인력양성"),
            ("재활용 산업단지", "MEDIUM",
             "파일럿 단지 개발 참여 — LH공사 + 한국 재활용 기업 컨소시엄"),
            ("MOIT 직접 접촉", "HIGH",
             "Decision 980 기술목록 등재 신청 → 정부 조달 우선 공급자 지위 확보"),
        ],
    },

    "VN-SC-2030": {
        "name_ko": "스마트시티 국가전략 2030 (기술·플랫폼)",
        "name_en": "Vietnam Smart City Development Strategy 2030",
        "legal_basis": "Decision 950/QD-TTg (2018) + Resolution 52/NQ-TW (2019)",
        "lead_agency": "Ministry of Construction (MOC) + Ministry of Information & Communications (MIC)",
        "investment": "$15B (2020-2030) — 공공+민간",
        "period": "2018 – 2030",
        "theme": ("violet", "7C3AED", "F5F3FF"),
        "report_type": "📊 단위상품형 KPI 보고서",
        "overview": (
            "베트남 스마트시티 기술·플랫폼 전략. 2030년 전국 63개 성시에 스마트시티 플랫폼 구축. "
            "핵심 기술: AI 교통관제, 디지털 트윈, IoT 인프라, 전자정부, 스마트 에너지. "
            "⚠️ 도시개발(공간)과 별도 — 기술 레이어 추적 전용 보고서."
        ),
        "national_kpis": [
            ("스마트시티 플랫폼 도입", "63개 성시", "현재 15개 시범 운영", "2030"),
            ("AI 교통관제", "주요 10개 도시", "하노이·호치민 시범", "2030"),
            ("전자정부 지수", "ASEAN Top 3", "현재 ASEAN 5위", "2030"),
            ("디지털 인프라", "5G 전국 커버", "현재 주요 도시만", "2030"),
            ("스마트 에너지", "30개 도시 스마트미터", "현재 7개 도시", "2030"),
        ],
        "sub_tracks": [
            {"id": "VN-SC-2030", "name": "스마트시티 전체", "focus": "플랫폼·기술·파일럿 동향"},
        ],
        "market_segments": [
            ("스마트 교통 ITS", "연간 $300M", "AI 신호제어·단속"),
            ("전자정부 플랫폼", "연간 $500M", "정부 디지털화"),
            ("스마트 에너지", "연간 $400M", "스마트미터·그리드"),
            ("도시 IoT 인프라", "연간 $200M", "센서·통신망"),
            ("디지털 트윈", "연간 $100M", "도시 모델링"),
        ],
        "korean_opportunities": [
            ("스마트 교통 ITS", "HIGH",
             "SK텔레콤·KT·현대차 — AI 교통관제·C-ITS·자율주행 인프라"),
            ("전자정부 플랫폼", "HIGH",
             "NIA(한국정보화진흥원)·LG CNS — 스마트시티 통합플랫폼 수출"),
            ("스마트 에너지", "MEDIUM",
             "한전KDN·LS Electric — 스마트미터·AMI 시스템"),
            ("CCTV·안전 인프라", "MEDIUM",
             "KT·한화비전 — 도시 안전망 솔루션"),
            ("ODA 스마트시티", "HIGH",
             "KOICA 스마트시티 ODA — 빈증·다낭 시범사업 한국 참여"),
        ],
    },

    "VN-OG-2030": {
        "name_ko": "석유가스 개발 계획 2030",
        "name_en": "Vietnam Oil & Gas Development Plan 2030, Vision 2045",
        "legal_basis": "Decision 893/QD-TTg (2023) + PetroVietnam Corporate Strategy",
        "lead_agency": "Ministry of Industry & Trade (MOIT) + PetroVietnam (PVN)",
        "investment": "$12B (2021-2030) — PVN + JV",
        "period": "2021 – 2030 / Vision 2045",
        "theme": ("amber", "D97706", "FFFBEB"),
        "report_type": "📊 단위상품형 KPI 보고서",
        "overview": (
            "PetroVietnam(PVN) 주도 석유가스 탐사·생산·정제 전략. "
            "2030년 원유 생산 800만 톤/년, 가스 생산 130억 ㎥/년 목표. "
            "Bach Ho(백호) 고갈 대응 — 심해 탐사·신규 가스전 개발 핵심."
        ),
        "national_kpis": [
            ("원유 생산량", "800만톤/년", "2024년 ~1,050만톤(감소추세)", "2030"),
            ("가스 생산량", "130억㎥/년", "2024년 ~100억㎥", "2030"),
            ("정제 설비용량", "12만배럴/일", "현재 Dung Quat+Nghi Son", "2030"),
            ("LNG 수입 터미널", "5개소 이상", "현재 1개(Thi Vai)", "2030"),
            ("신규 탐사 블록", "15+ 블록", "JV 입찰 진행중", "2030"),
        ],
        "sub_tracks": [
            {"id": "VN-OG-2030", "name": "석유가스 전체", "focus": "탐사·생산·정제·LNG"},
        ],
        "market_segments": [
            ("탐사·생산 JV", "블록당 $1-5B", "심해 탐사 6개 블록"),
            ("LNG 수입 터미널", "터미널당 $0.5-1B", "5개소 건설 예정"),
            ("정제설비 고도화", "$2B+", "수소화처리·고도화"),
            ("가스 파이프라인", "$1B+", "육상·해저 배관"),
            ("O&M 서비스", "연간 $500M", "해상플랜트 유지관리"),
        ],
        "korean_opportunities": [
            ("심해 탐사 참여", "HIGH",
             "GS에너지·SK이노베이션 — 심해 블록 JV 참여 (PVN 입찰)"),
            ("LNG 터미널 건설", "HIGH",
             "삼성물산·현대건설 — Thi Vai 2단계·Ca Mau LNG 터미널 EPC"),
            ("해상 플랜트 제조", "MEDIUM",
             "현대중공업·삼성중공업 — FPSO·FSO 발주 수주"),
            ("O&M 서비스", "MEDIUM",
             "한국가스공사(KOGAS)·SK — LNG 터미널 운영 노하우 수출"),
            ("플랜트 엔지니어링", "MEDIUM",
             "삼성엔지니어링·GS건설 — 정제설비 고도화 EPC"),
        ],
    },

    "VN-EV-2030": {
        "name_ko": "전기차·친환경 모빌리티 2030",
        "name_en": "Electric Vehicle & Green Mobility Development Strategy 2030",
        "legal_basis": "Decision 876/QD-TTg (2022) + Decision 1119/QD-TTg (2022)",
        "lead_agency": "Ministry of Transport (MOT) + Ministry of Industry & Trade (MOIT)",
        "investment": "$5B (인프라) + 민간 투자",
        "period": "2022 – 2040",
        "theme": ("teal", "0D9488", "F0FDFA"),
        "report_type": "📊 단위상품형 KPI 보고서",
        "overview": (
            "2040년 내연기관 판매 금지, 2050년 탄소중립 모빌리티 전환. "
            "VinFast 주도 국내 EV 생산 + 충전 인프라 구축. "
            "2030년 EV 판매 비중 30%, 충전소 10만 기 목표."
        ),
        "national_kpis": [
            ("EV 판매 비중", "30%", "2024년 ~8% (VinFast 급성장)", "2030"),
            ("EV 충전소", "10만 기", "2024년 ~15,000기", "2030"),
            ("전기버스 전환", "50%", "2024년 ~5% (Hanoi Bus 선도)", "2030"),
            ("전기오토바이", "70%", "2024년 ~15%", "2040"),
            ("EV 배터리 생산", "국내 공장 3개+", "VinFast Vu Hung 공장", "2030"),
        ],
        "sub_tracks": [
            {"id": "VN-EV-2030", "name": "EV·모빌리티 전체", "focus": "차량·충전·정책 동향"},
        ],
        "market_segments": [
            ("EV 차량", "연간 $3B+", "VinFast + 수입 EV"),
            ("충전 인프라", "연간 $500M", "공공·민간 충전소"),
            ("EV 배터리", "연간 $1B+", "생산+교체 시장"),
            ("전기버스", "연간 $300M", "도시 대중교통"),
            ("V2G·스마트충전", "연간 $100M", "신기술 시장"),
        ],
        "korean_opportunities": [
            ("배터리 공급", "HIGH",
             "삼성SDI·LG에너지솔루션·SK온 — VinFast 배터리 공급 확대"),
            ("충전 인프라", "HIGH",
             "현대차·SK·LS Electric — 급속충전기·V2G 시스템 공급"),
            ("전기버스 제조", "MEDIUM",
             "현대차 Elec City 버스 — 하노이·호치민 전기버스 입찰"),
            ("배터리 소재", "MEDIUM",
             "포스코퓨처엠·에코프로 — 양극재·음극재 JV"),
            ("차량용 반도체", "MEDIUM",
             "삼성전자·SK하이닉스 — 차량용 칩 공급"),
        ],
    },

    "VN-CARBON-2050": {
        "name_ko": "탄소중립 2050 로드맵",
        "name_en": "Vietnam Net Zero Emissions Roadmap 2050",
        "legal_basis": "NDC (2022 Update) + Decision 896/QD-TTg (2022) + COP26 공약",
        "lead_agency": "Ministry of Natural Resources & Environment (MONRE)",
        "investment": "$368B (2022-2050) — 국제 기후금융 필요",
        "period": "2022 – 2050",
        "theme": ("green", "166534", "F0FDF4"),
        "report_type": "📊 단위상품형 KPI 보고서",
        "overview": (
            "COP26 탄소중립 2050 공약 이행 로드맵. "
            "2030년 NDC: GHG 43.5% 감축(국제지원 시). "
            "2050년 탄소중립 달성을 위한 에너지전환·산업전환·산림·탄소거래 체계 구축."
        ),
        "national_kpis": [
            ("GHG 감축 (2030)", "43.5% (국제지원)", "현재 NDC 목표 설정", "2030"),
            ("탄소중립 달성", "Net Zero", "2050년 목표", "2050"),
            ("ETS 탄소거래소", "운영 개시", "시범운영 2025, 공식 2028", "2028"),
            ("산림 피복률", "42%", "현재 42.02%", "2030"),
            ("REDD+ 프로젝트", "30+ 건", "현재 12건", "2030"),
        ],
        "sub_tracks": [
            {"id": "VN-CARBON-2050", "name": "탄소중립 전체", "focus": "NDC·ETS·탄소시장 동향"},
        ],
        "market_segments": [
            ("탄소 크레딧", "연간 $500M+ (2030~)", "VCM + 의무시장"),
            ("녹색채권", "연간 $1B+", "기후금융"),
            ("CCUS 기술", "$2B+ 시장", "포집·저장·활용"),
            ("탄소경영 컨설팅", "연간 $200M", "기업 탄소중립 지원"),
            ("자연기반 솔루션", "연간 $100M", "맹그로브·산림 REDD+"),
        ],
        "korean_opportunities": [
            ("탄소 크레딧 사업", "HIGH",
             "한국 기업 → GHG 크레딧 확보 — REDD+·재생에너지 크레딧"),
            ("ETS 시스템 구축", "HIGH",
             "한국 K-ETS 운영 경험 — 베트남 ETS 설계·운영 기술자문 (KEITI)"),
            ("CCUS 기술", "MEDIUM",
             "포스코·POSCO인터내셔널 — CCS 파일럿 프로젝트"),
            ("녹색채권 투자", "MEDIUM",
             "한국 금융기관 — 베트남 녹색채권 인수·투자"),
            ("탄소경영 솔루션", "MEDIUM",
             "한국 ESG 컨설팅사 — 베트남 기업 탄소중립 전환 지원"),
        ],
    },
}

# ─────────────────────────────────────────────────────────
# KPI형 보고서 생성 함수
# ─────────────────────────────────────────────────────────
def generate_kpi_report(plan_id: str, output_dir: str,
                        extra_arts=None) -> str:
    """
    extra_arts: 외부에서 넘긴 추가 기사 (통합 보고서 하부 항목용)
    """
    if plan_id not in KPI_CONFIG:
        raise ValueError(f"KPI_CONFIG에 {plan_id}가 없습니다.")

    cfg = KPI_CONFIG[plan_id]
    theme_color, theme_hex, theme_light = cfg["theme"]
    now = datetime.now()
    week = now.isocalendar()[1]
    week_str = f"W{week:02d}/{now.year}"

    # 기사 로드
    buckets = load_history([plan_id])
    arts = sorted(buckets.get(plan_id, []), key=_date_key, reverse=True)
    if extra_arts:
        # 중복 제거 후 병합
        seen = {a.get("url","") or str(id(a)) for a in arts}
        for a in extra_arts:
            k = a.get("url","") or str(id(a))
            if k not in seen:
                seen.add(k); arts.append(a)
        arts = sorted(arts, key=_date_key, reverse=True)

    rec_arts, year_label = select_articles(arts)
    kpi_db = load_kpi_db()

    doc = new_doc()

    # ── 표지 ─────────────────────────────────────────────
    build_cover(
        doc, plan_id,
        cfg["name_ko"], cfg["name_en"],
        cfg["legal_basis"], cfg["lead_agency"],
        cfg["investment"], cfg["period"],
        len(arts), week_str, cfg["report_type"],
        theme_hex
    )

    # ── Section 1: 프로그램/시장 개요 ────────────────────
    add_heading(doc, "1. 프로그램/시장 개요", 1, theme_color, 13)
    add_hr(doc, theme_hex)

    ot = doc.add_table(rows=1, cols=1); ot.style = "Table Grid"
    oc = ot.rows[0].cells[0]; cell_bg(oc, theme_light)
    op = oc.paragraphs[0]
    op.add_run(cfg["overview"]).font.size = Pt(10)
    doc.add_paragraph()

    # 시장 세그먼트
    add_heading(doc, "1-1. 시장 세그먼트", 2, theme_color)
    seg_rows = []
    for seg_name, size, note in cfg.get("market_segments", []):
        seg_rows.append([seg_name, size, note])
    if seg_rows:
        styled_table(doc,
            ["세그먼트", "시장 규모", "비고"],
            seg_rows,
            col_widths=[2.0, 1.8, 3.5],
            header_bg=theme_color, font_size=9.5)

    doc.add_page_break()

    # ── Section 2: KPI 현황표 ★ ──────────────────────────
    add_heading(doc, "2. KPI 현황표 ★", 1, theme_color, 13)
    add_hr(doc, theme_hex)
    add_note(doc, "마스터플랜 공식 KPI 목표치 대비 현재 달성 현황", color_key="gray")
    doc.add_paragraph()

    # KPI 대형 박스
    kpi_items_top = []
    for kpi_name, target, current, deadline in cfg.get("national_kpis", [])[:4]:
        kpi_items_top.append((kpi_name[:20], target, f"목표: {deadline}"))
    if kpi_items_top:
        kpi_box(doc, kpi_items_top,
                bg_hex=theme_light, border_hex=theme_hex,
                title=f"📊 {cfg['name_ko']} — 핵심 KPI 목표", font_size=10)

    # KPI 상세 테이블
    kpi_rows = []
    for kpi_name, target, current, deadline in cfg.get("national_kpis", []):
        # 간단 달성률 계산 (텍스트 기반)
        kpi_rows.append([kpi_name, target, current, deadline, "▶ 진행중"])
    styled_table(doc,
        ["KPI 지표", "목표치", "현재 수준", "목표연도", "상태"],
        kpi_rows,
        col_widths=[2.2, 1.5, 2.0, 0.9, 0.9],
        header_bg=theme_color, font_size=9.5)

    # KPI DB 상세 (있는 경우)
    if plan_id in kpi_db:
        nkpi = kpi_db[plan_id].get("national_kpi", {})
        if nkpi:
            add_heading(doc, "2-1. KPI DB 상세 데이터", 2, theme_color)
            kpi_detail_rows = []
            for yr_key, yr_data in nkpi.items():
                if isinstance(yr_data, dict):
                    for metric, val in yr_data.items():
                        kpi_detail_rows.append([yr_key.replace("_"," "), metric.replace("_"," "), str(val)])
                else:
                    kpi_detail_rows.append([yr_key.replace("_"," "), "투자 총액", str(yr_data)])
            if kpi_detail_rows:
                styled_table(doc,
                    ["기준연도/구분", "지표", "수치"],
                    kpi_detail_rows,
                    col_widths=[1.6, 2.8, 3.0],
                    header_bg=theme_color, font_size=9)

    doc.add_page_break()

    # ── Section 3: 하부 트래킹 항목별 진행현황 ───────────
    add_heading(doc, "3. 트래킹 항목별 진행현황", 1, theme_color, 13)
    add_hr(doc, theme_hex)

    for sub in cfg.get("sub_tracks", []):
        sub_id = sub["id"]
        sub_name = sub["name"]
        sub_focus = sub.get("focus", "")

        add_heading(doc, f"▶ {sub_name}  [{sub_id}]", 2, theme_color)
        add_note(doc, f"📌 {sub_focus}", color_key="gray")
        doc.add_paragraph()

        # 해당 하부 ID 기사
        sub_buckets = load_history([sub_id])
        sub_arts = sorted(sub_buckets.get(sub_id, []), key=_date_key, reverse=True)

        # 기사 통계 미니 KPI
        yr_trend = get_year_trend(sub_arts, sub_id)
        trend_items = [(y, yr_trend.get(y, 0), "건") for y in ["2023","2024","2025","2026"]]
        kpi_box(doc, trend_items,
                bg_hex=theme_light, border_hex=theme_hex,
                title=f"연도별 기사 건수 — {sub_name}", font_size=9)

        # 최신 뉴스
        rec, yr_lbl = select_articles(sub_arts)
        if rec:
            add_heading(doc, f"  📰 최신 뉴스 ({yr_lbl}, 상위 5건)", 3, theme_color)
            render_articles(doc, rec, theme_color, max_n=5)

        # 역사 타임라인
        hist = [a for a in sub_arts if _date_key(a) < "2026-01-01"]
        if hist:
            add_heading(doc, "  📚 역사 기사 타임라인 (2019–2025)", 3, theme_color)
            render_history_timeline(doc, hist, theme_color)

        add_hr(doc, "E5E7EB"); doc.add_paragraph()

    doc.add_page_break()

    # ── Section 4: 시장 동향 분석 ────────────────────────
    add_heading(doc, "4. 시장 동향 분석", 1, theme_color, 13)
    add_hr(doc, theme_hex)

    # 연도별 기사 건수 트렌드
    add_heading(doc, "4-1. 연도별 기사 건수 트렌드", 2, theme_color)
    yr_all = get_year_trend(arts, plan_id)
    trend_rows = [["연도"] + [str(y) for y in range(2019, 2027)],
                  ["기사수"] + [str(yr_all.get(str(y), 0)) for y in range(2019, 2027)]]
    styled_table(doc,
        ["구분", "2019","2020","2021","2022","2023","2024","2025","2026"],
        [["기사수"] + [str(yr_all.get(str(y),0)) for y in range(2019,2027)]],
        header_bg=theme_color, font_size=9)

    # 성별 활동
    add_heading(doc, "4-2. 지역별 활동 현황", 2, theme_color)
    render_province_section(doc, arts, plan_id, kpi_db, theme_hex)

    doc.add_page_break()

    # ── Section 5: 한국 기업 기회 ────────────────────────
    add_heading(doc, "5. 한국 기업 기회 분석", 1, theme_color, 13)
    add_hr(doc, theme_hex)
    render_korean_opportunity(doc, cfg["korean_opportunities"], theme_hex)
    doc.add_page_break()

    # ── Appendix ─────────────────────────────────────────
    build_appendix(doc, arts, theme_hex)

    # ── 저장 ─────────────────────────────────────────────
    safe_id = plan_id.replace("-","_")
    fname = f"MI_REPORT_{safe_id}_{week_str.replace('/','_')}_{now.strftime('%Y%m%d')}.docx"
    fpath, sz = save_report(doc, output_dir, fname)
    return fpath


def generate_all_kpi(output_dir: str) -> list:
    results = []
    for plan_id in KPI_CONFIG:
        print(f"  📊 {plan_id} ...", end=" ", flush=True)
        try:
            fpath = generate_kpi_report(plan_id, output_dir)
            sz = os.path.getsize(fpath) // 1024
            print(f"✅ {sz}KB  [{fpath.split('/')[-1]}]")
            results.append((plan_id, fpath, sz))
        except Exception as e:
            print(f"❌ {e}")
            results.append((plan_id, None, 0))
    return results


if __name__ == "__main__":
    import sys
    sys.path.insert(0, "/home/work/claw/scripts")
    out = "/home/work/claw/outputs/reports/MI_Reports_v4"
    Path(out).mkdir(parents=True, exist_ok=True)

    plan = sys.argv[1] if len(sys.argv) > 1 else "VN-OG-2030"
    print(f"KPI 보고서 생성: {plan}")
    fpath = generate_kpi_report(plan, out)
    print(f"✅ {fpath}  ({os.path.getsize(fpath)//1024}KB)")

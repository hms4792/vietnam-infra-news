"""
mi_water_report_v4.py — Water 수자원 통합 보고서 v4.0
"""
import os, json
from datetime import datetime
from pathlib import Path

from report_lib import (
    new_doc, add_heading, add_hr, add_note,
    styled_table, kpi_box, build_cover,
    render_articles, render_history_timeline,
    render_province_section, render_korean_opportunity,
    build_appendix, load_history, load_kpi_db, save_report,
    select_articles, get_year_trend, _date_key,
    PALETTE, _rgb, cell_bg,
)
from docx.shared import Pt

BASE_DIR = Path("/home/work/claw")

WATER_SUBS = [
    {"id":"VN-WAT-RESOURCES","name":"수자원 관리·유역 계획",
     "color":"sky","hex":"0EA5E9",
     "legal":"Decision 1622/QD-TTg (2022), MONRE",
     "kpis":[("유역 통합관리","10개 주요 유역","현재 7개 운영"),
             ("지하수 보호","Red River·Mekong","고갈 경보 단계"),
             ("물 안보 지수","ASEAN Top 5","현재 6위")]},
    {"id":"VN-WAT-URBAN","name":"도시 상수도 인프라",
     "color":"blue","hex":"1E40AF",
     "legal":"Decision 1929/QD-TTg (2021) + Decision 1566/QD-TTg (2021), MOC",
     "kpis":[("도시 수도보급률","95~100%","현재 ~92%"),
             ("수돗물 품질","국제기준 100%","현재 ~75%"),
             ("NRW(비수익 손실)","15% 이하","현재 평균 18%")]},
    {"id":"VN-WAT-RURAL","name":"농촌 상수도·위생",
     "color":"green","hex":"15803D",
     "legal":"Decision 1978/QD-TTg (2021), MARD",
     "kpis":[("농촌 수도보급률","62M명 청정수 접근","현재 ~85% 공급"),
             ("농촌 위생시설","70% 2025, 90% 2030","현재 ~65%"),
             ("농촌 NRW","20% 이하","현재 ~25%")]},
]

THEME_HEX   = "0EA5E9"
THEME_LIGHT = "F0F9FF"
THEME_COLOR = "sky"

def generate_water_report(output_dir: str) -> str:
    now = datetime.now(); week = now.isocalendar()[1]
    week_str = f"W{week:02d}/{now.year}"

    all_ids = [s["id"] for s in WATER_SUBS]
    buckets = load_history(all_ids)
    all_arts = []; seen = set()
    for pid, arts in buckets.items():
        for a in arts:
            k = a.get("url","") or str(id(a))
            if k not in seen: seen.add(k); all_arts.append(a)
    all_arts.sort(key=_date_key, reverse=True)

    doc = new_doc()

    build_cover(doc,
        "WATER-INTEGRATED",
        "베트남 수자원·상수도 통합 보고서",
        "Vietnam Water Resources & Water Supply Integrated Report",
        "Decision 1622/QD-TTg + Decision 1929/QD-TTg + Decision 1978/QD-TTg",
        "MONRE (수자원) + MOC (도시상수도) + MARD (농촌급수)",
        "$15B+ (2021-2030, 수자원·상수도·농촌급수 합산)",
        "2021 – 2030 / Vision 2050",
        len(all_arts), week_str,
        "💧 수자원 통합 KPI 보고서", THEME_HEX)

    # Section 1: 개요
    add_heading(doc, "1. 수자원·상수도 정책 체계 개요", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)
    ot = doc.add_table(rows=1, cols=1); ot.style = "Table Grid"
    oc = ot.rows[0].cells[0]; cell_bg(oc, THEME_LIGHT)
    oc.paragraphs[0].add_run(
        "베트남 3개 수자원 마스터플랜 통합 추적. "
        "VN-WAT-RESOURCES: MONRE 주관 유역관리·지하수 보호. "
        "VN-WAT-URBAN: MOC 주관 전국 도시 상수도 현대화. "
        "VN-WAT-RURAL: MARD 주관 농촌 6,200만 인구 청정수 공급. "
        "총 투자 $15B+, 2030년 전 국민 안전한 식수 접근 목표."
    ).font.size = Pt(10)
    doc.add_paragraph()

    kpi_box(doc,
        [("도시 보급률 2030","95~100%","현재 ~92%"),
         ("농촌 보급률 2030","90%","현재 ~85%"),
         ("NRW 목표","15~20%","현재 18~25%"),
         ("수자원 유역관리","10개 유역","현재 7개")],
        bg_hex=THEME_LIGHT, border_hex=THEME_HEX,
        title="수자원·상수도 핵심 KPI (2030 목표)", font_size=10)

    # Sub-track 정보
    styled_table(doc,
        ["하부 트래킹", "법적 근거", "담당부처", "핵심 목표"],
        [[s["name"], s["legal"][:40], s["legal"].split(",")[-1].strip()[:10],
          s["kpis"][0][1][:35]] for s in WATER_SUBS],
        col_widths=[2.0, 2.5, 1.2, 2.6],
        header_bg=THEME_COLOR, font_size=9.5)
    doc.add_page_break()

    # Section 2: KPI 현황표
    add_heading(doc, "2. KPI 현황표", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)
    kpi_rows = []
    for sub in WATER_SUBS:
        for kpi_name, target, current in sub["kpis"]:
            kpi_rows.append([sub["name"][:22], kpi_name[:28], target[:28], current[:30]])
    styled_table(doc,
        ["하부 트래킹","KPI 지표","목표치","현재 수준"],
        kpi_rows,
        col_widths=[2.0,2.2,1.8,2.3],
        header_bg=THEME_COLOR, font_size=9)
    doc.add_page_break()

    # Section 3: 하부별 상세
    add_heading(doc, "3. 하부 트래킹별 상세 진행현황", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)

    for sub in WATER_SUBS:
        pid = sub["id"]; color = sub["color"]; hex_c = sub["hex"]
        sub_arts = sorted(buckets.get(pid,[]), key=_date_key, reverse=True)

        add_heading(doc, f"▶ {sub['name']}  [{pid}]", 2, color)
        add_note(doc, f"📋 {sub['legal']}", "gray")
        doc.add_paragraph()

        kpi_items = [(k[:20], v[:25], "") for k,v,_ in sub["kpis"][:4]]
        kpi_box(doc, kpi_items, bg_hex=THEME_LIGHT, border_hex=hex_c,
                title=f"{sub['name']} — 목표 KPI", font_size=9)

        yr_trend = get_year_trend(sub_arts, pid)
        trend_items = [(y, yr_trend.get(y,0), "건") for y in ["2023","2024","2025","2026"]]
        kpi_box(doc, trend_items, bg_hex="F8FAFC", border_hex=hex_c,
                title="연도별 기사 건수", font_size=9)

        rec, yr_lbl = select_articles(sub_arts)
        if rec:
            add_heading(doc, f"  📰 최신 뉴스 ({yr_lbl}, 상위 4건)", 3, color)
            render_articles(doc, rec, color, max_n=4)

        hist = [a for a in sub_arts if _date_key(a) < "2026-01-01"]
        if hist:
            add_heading(doc, "  📚 역사 기사 타임라인", 3, color)
            render_history_timeline(doc, hist, color)

        if not sub_arts:
            add_note(doc, "  ℹ️ 관련 기사 수집 중.", "gray")

        add_hr(doc, "E5E7EB"); doc.add_paragraph()

    doc.add_page_break()

    # Section 4: 시장 동향
    add_heading(doc, "4. 시장 동향 분석", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)
    yr_all = get_year_trend(all_arts, "WATER")
    styled_table(doc,
        ["2019","2020","2021","2022","2023","2024","2025","2026"],
        [[str(yr_all.get(str(y),0)) for y in range(2019,2027)]],
        header_bg=THEME_COLOR, font_size=9)
    render_province_section(doc, all_arts, "VN-WAT-RESOURCES", None, THEME_HEX)
    doc.add_page_break()

    # Section 5: 한국 기업 기회
    add_heading(doc, "5. 한국 기업 기회", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)
    render_korean_opportunity(doc, [
        ("상수도 EPC", "HIGH",
         "현대건설·GS건설 — 하노이·호치민·다낭 WTP 건설 EPC 패키지"),
        ("막여과(UF/MBR) 기술", "HIGH",
         "코오롱인더스트리·도레이·롯데케미칼 — UF 막 모듈 현지 생산"),
        ("스마트 수도관리", "HIGH",
         "K-water·한국수자원공사 — NRW 저감 스마트 파이프망 기술"),
        ("농촌 급수 시스템", "MEDIUM",
         "K-water·KOICA ODA — 농촌 소규모 급수시스템 패키지"),
        ("수자원 디지털 트윈", "MEDIUM",
         "한국건설기술연구원 — 홍강·메콩 유역 디지털트윈 구축"),
        ("ODA 패키지", "HIGH",
         "EDCF + KOICA — 도시 상수도 현대화 ODA ($500M+)"),
    ], THEME_HEX)
    doc.add_page_break()

    build_appendix(doc, all_arts, THEME_HEX)

    fname = f"MI_REPORT_WATER_INTEGRATED_W{week:02d}_{now.strftime('%Y%m%d')}.docx"
    fpath, sz = save_report(doc, output_dir, fname)
    return fpath


if __name__ == "__main__":
    import sys; sys.path.insert(0,"/home/work/claw/scripts")
    out = "/home/work/claw/outputs/reports/MI_Reports_v4"
    Path(out).mkdir(parents=True, exist_ok=True)
    fpath = generate_water_report(out)
    print(f"✅ {fpath}  ({os.path.getsize(fpath)//1024}KB)")

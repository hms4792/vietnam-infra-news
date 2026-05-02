"""
하노이 도시개발 통합 MI 보고서 생성기
=====================================
Decision 1668/QD-TTg (2024.12.27) 기반 — 하노이 마스터플랜 2045/2065

구조:
  표지  : 하노이 마스터플랜 2045/2065 개요 + 투자 규모
  Sec 1 : 하노이 도시개발 프레임워크 — 5대 도시권 + 용어 정리
  Sec 2 : 북부 신도시 (동아인·BRG스마트시티·노이바이)  [HN-URBAN-NORTH]
  Sec 3 : 서부 과학기술도시 (호아락·하이테크·선따이)    [HN-URBAN-WEST]
  Sec 4 : 동부 도시권 (롱비엔·자람·링로드4)             [HN-URBAN-EAST]
  Sec 5 : 남부 물류도시 (푸쑤엔·자빈공항)               [HN-URBAN-SOUTH]
  Sec 6 : 도시 인프라 공통 (메트로·링로드·홍강·도심)    [HN-URBAN-INFRA]
  Sec 7 : 스마트시티 vs 도시개발 — 용어 정리 + 기사 분류
  Sec 8 : 한국 기업 기회 분석
  Appendix: 전체 관련 기사
"""
import os, json, re
from datetime import datetime
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path("/home/work/claw")

COLOR_MAIN  = (0xDC, 0x26, 0x26)   # red     — 하노이
COLOR_NORTH = (0xEF, 0x44, 0x44)   # red-500 — 북부
COLOR_WEST  = (0x7C, 0x3A, 0xED)   # violet  — 서부
COLOR_EAST  = (0x0E, 0xA5, 0xE9)   # sky     — 동부
COLOR_SOUTH = (0x16, 0xA3, 0x4A)   # green   — 남부
COLOR_INFRA = (0xF5, 0x9E, 0x0B)   # amber   — 인프라
COLOR_SC    = (0xA7, 0x8B, 0xFA)   # violet  — 스마트시티
COLOR_KOR   = (0x0F, 0x17, 0x2A)   # navy    — 한국

HEX = {
    "main":"DC2626","north":"EF4444","west":"7C3AED","east":"0EA5E9",
    "south":"16A34A","infra":"F59E0B","sc":"A78BFA","kor":"1E3A5F","light":"FEF2F2"
}

def _rgb(t): return RGBColor(*t)
def _hex_bg(cell, h):
    tc=cell._tc; p=tc.get_or_add_tcPr()
    s=OxmlElement('w:shd'); s.set(qn('w:val'),'clear')
    s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),h); p.append(s)
def _hr(doc, c="CCCCCC"):
    p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr()
    pb=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),c); pb.append(bot); pPr.append(pb)
def _h(doc, text, lvl, color=COLOR_MAIN, size=None):
    h=doc.add_heading(text,lvl)
    for r in h.runs:
        r.font.color.rgb=_rgb(color)
        if size: r.font.size=Pt(size)
    return h
def _date_key(a):
    d=str(a.get('published_date','') or '')
    try:
        if re.match(r'\d{4}-\d{2}-\d{2}',d): return d[:10]
        return datetime.strptime(d.strip(),'%b %d, %Y').strftime('%Y-%m-%d')
    except: return '1900-01-01'
def _is_vi(t): return any(c in (t or '') for c in ['ă','ơ','ư','đ','ấ','ề','ộ','ừ','ạ','ọ','ổ','ị','ế'])
def _title(a):
    t=a.get('title','')
    if _is_vi(t):
        en=a.get('summary_en','')
        if en and not _is_vi(en): return f"[VI] {en[:80]}"
        ko=a.get('summary_ko','')
        if ko and not _is_vi(ko): return f"[VI] {ko[:70]}"
    return t[:85]
def _summary(a):
    ko=a.get('summary_ko',''); en=a.get('summary_en','')
    if ko and not _is_vi(ko): return '🇰🇷', ko[:200]
    if en and not _is_vi(en): return '🇺🇸', en[:200]
    return None, None
def _select_arts(arts):
    s=sorted(arts,key=_date_key,reverse=True)
    for yr in ['2026','2025','2024']:
        y=[a for a in s if yr in str(a.get('published_date',''))]
        if y: return y,yr
    return s[:20],'전체'
def _add_articles(doc, arts, color, max_n=5):
    for i,art in enumerate(arts[:max_n],1):
        t=_title(art); flag,sm=_summary(art); date=_date_key(art)
        src=(art.get('source','') or '')[:30]
        ap=doc.add_paragraph(); ap.paragraph_format.left_indent=Inches(0.2)
        nr=ap.add_run(f"[{i}] "); nr.bold=True; nr.font.size=Pt(10); nr.font.color.rgb=_rgb(color)
        ap.add_run(t).bold=True; ap.runs[-1].font.size=Pt(10)
        mp=doc.add_paragraph(); mp.paragraph_format.left_indent=Inches(0.4)
        mp.add_run(f"📅 {date}  |  📰 {src}").font.size=Pt(8.5)
        if flag and sm:
            sp=doc.add_paragraph(); sp.paragraph_format.left_indent=Inches(0.4)
            sp.add_run(f"{flag} ").font.size=Pt(9.5); sp.add_run(sm).font.size=Pt(9.5)
            en=art.get('summary_en','')
            if flag=='🇰🇷' and en and not _is_vi(en):
                ep=doc.add_paragraph(); ep.paragraph_format.left_indent=Inches(0.4)
                ep.add_run("🇺🇸 ").font.size=Pt(9.5); ep.add_run(en[:180]).font.size=Pt(9.5)
        doc.add_paragraph()

def _proj_table(doc, projects, color):
    if not projects: return
    t=doc.add_table(rows=1,cols=4); t.style="Table Grid"
    for j,h in enumerate(["프로젝트명","규모/투자","개발사","현황"]):
        t.rows[0].cells[j].text=h
        t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _hex_bg(t.rows[0].cells[j], HEX["main"])
        t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        t.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)
    for p in projects:
        row=t.add_row()
        area=f"{p.get('area_ha','')}" if p.get('area_ha') else ""
        inv=p.get('investment_usd','') or p.get('investment','') or 'TBD'
        if area: inv=f"{area}ha / ${inv}" if '$' not in str(inv) else f"{area}ha / {inv}"
        row.cells[0].text=p.get('name','')[:45]
        row.cells[1].text=str(inv)[:30]
        row.cells[2].text=p.get('developer','')[:25] or p.get('note','')[:25] or '-'
        row.cells[3].text=p.get('status','')[:35]
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(8.5)
    doc.add_paragraph()


def generate_hanoi_report(output_dir: str) -> str:
    with open(BASE_DIR/"config/history_db.json", encoding='utf-8') as f:
        hdb=json.load(f)
    with open(BASE_DIR/"config/hanoi_urban_structure.json", encoding='utf-8') as f:
        hn=json.load(f)

    now=datetime.now(); week=now.isocalendar()[1]

    # 기사 버킷 — HN-URBAN-* + VN-SC-2030 + VN-URB-METRO-2030
    HANOI_IDS=["HN-URBAN-NORTH","HN-URBAN-WEST","HN-URBAN-INFRA","VN-SC-2030","VN-URB-METRO-2030"]
    buckets=defaultdict(list)
    for art in hdb['articles'].values():
        for pid in (art.get('matched_plans') or []):
            if pid in HANOI_IDS: buckets[pid].append(art)

    # 추가: 하노이 지역 필터 (어떤 플랜에도 없지만 하노이 관련 기사)
    hanoi_kws=["dong anh","me linh","soc son","noi bai","hoa lac","xuan mai","son tay",
               "gia lam","long bien","phu xuyen","ring road 4","hanoi master plan",
               "đông anh","mê linh","hòa lạc","gia lâm","vành đai 4","quy hoạch hà nội",
               "red river corridor hanoi","to lich"]
    extra_hanoi=[]
    for art in hdb['articles'].values():
        text=((art.get('title','') or '')+(art.get('summary_en','') or '')+(art.get('content','') or '')[:200]).lower()
        if any(kw in text for kw in hanoi_kws):
            # 기 버킷된 기사 제외
            already=any(art in buckets[pid] for pid in HANOI_IDS)
            if not already: extra_hanoi.append(art)
    extra_hanoi=sorted(extra_hanoi,key=_date_key,reverse=True)

    total_arts=sum(len(v) for v in buckets.values())+len(extra_hanoi)

    doc=Document()
    for sec in doc.sections:
        sec.top_margin=Inches(0.9); sec.bottom_margin=Inches(0.9)
        sec.left_margin=Inches(1.2); sec.right_margin=Inches(1.2)

    # ══ COVER ══
    cov=doc.add_table(rows=1,cols=1); cov.style="Table Grid"
    c=cov.rows[0].cells[0]; _hex_bg(c,"0F172A")
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("VIETNAM INFRASTRUCTURE INTELLIGENCE HUB")
    r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    doc.add_paragraph()

    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=tp.add_run("하노이 도시개발 마스터플랜 2045/2065")
    r2.bold=True; r2.font.size=Pt(19); r2.font.color.rgb=_rgb(COLOR_MAIN)

    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run("Hanoi General Planning to 2045, Vision to 2065 — Decision 1668/QD-TTg | December 27, 2024").font.size=Pt(10)
    doc.add_paragraph()

    it=doc.add_table(rows=6,cols=2); it.style="Table Grid"
    for i,(k,v) in enumerate([
        ("결정문",      "Decision 1668/QD-TTg (2024.12.27) 마스터플랜 + Decision 1569/QD-TTg (2024.12.12) 수도계획"),
        ("통합 결의",   "Resolution 258/2025/QH15 (2025.12.11) — 단일 통합 계획 수립 허가 (Capital Plan + Master Plan 통합)"),
        ("도시 모델",   "9개 성장축 × 9개 거점 × 9개 중심 다핵다중심 도시 클러스터 / 행정면적 3,360km²"),
        ("투자 규모",   "2026–2035: $556.4B | 2036–2045: $1.9T | 총 2026–2045: $2.5T"),
        ("인구 목표",   "2030: 12M → 2045: 14.6–16M → 2065: 17–19M → 2100: ≤20M"),
        ("보고 기사",   f"하노이 관련 기사: {total_arts}건 | 5개 도시권 추적"),
    ]):
        it.rows[i].cells[0].text=k; it.rows[i].cells[1].text=v
        _hex_bg(it.rows[i].cells[0],HEX["main"])
        it.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        it.rows[i].cells[0].paragraphs[0].runs[0].bold=True
        for cell in it.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(9.5)
    doc.add_paragraph()
    doc.add_paragraph().alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].add_run(
        f"Generated: {now.strftime('%Y-%m-%d %H:%M')}  |  W{week}/{now.year}  |  CONFIDENTIAL"
    ).font.size=Pt(9)
    doc.add_page_break()

    # ══ SECTION 1: 도시개발 프레임워크 + 용어 정리 ══
    _h(doc,"1. 하노이 도시개발 프레임워크 — 5대 도시권 + 용어 정리",1,COLOR_MAIN,13); _hr(doc,HEX["main"])

    # 용어 정리 박스
    term_t=doc.add_table(rows=1,cols=1); term_t.style="Table Grid"
    tc=term_t.rows[0].cells[0]; _hex_bg(tc,"FEF3C7")
    tp2=tc.paragraphs[0]; tp2.add_run("📌  스마트시티 vs 도시개발 마스터플랜 — 용어 정리").bold=True
    tp2.runs[0].font.size=Pt(10); tp2.runs[0].font.color.rgb=_rgb(COLOR_INFRA)
    tc.add_paragraph()
    desc=tc.add_paragraph()
    desc.add_run(
        "스마트시티 (VN-SC-2030, 국가전략): 기술·플랫폼 레이어 — IoT, AI교통, 전자정부, 디지털트윈, 스마트에너지.\n"
        "                                        기술 시스템을 '어디에' 구축하느냐의 문제.\n\n"
        "도시개발 마스터플랜 (이 보고서): 공간·인프라 레이어 — '어디에' 무엇을 짓느냐의 문제.\n"
        "                                        Decision 1668: 북부·서부·동부·남부·중심 5개 도시권 공간 구조.\n\n"
        "⚡ 핵심: BRG-Sumitomo 동아인 스마트시티는 '스마트시티 기술'을 채택한 '북부 신도시 개발 프로젝트'.\n"
        "         → 도시개발(공간) + 스마트시티(기술) 두 레이어가 동시 적용되는 사례."
    ).font.size=Pt(9.5)
    doc.add_paragraph()

    # 5대 도시권 요약 테이블
    _h(doc,"1-1. 5대 도시권 개요",2,COLOR_MAIN)
    zones=hn["five_urban_zones"]
    zt=doc.add_table(rows=1,cols=5); zt.style="Table Grid"
    for j,h in enumerate(["도시권","면적/인구","행정구역","기능","핵심 프로젝트"]):
        zt.rows[0].cells[j].text=h
        zt.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _hex_bg(zt.rows[0].cells[j],HEX["main"])
        zt.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        zt.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)
    zone_rows=[
        ("🏛️ 중심 (구도심)", "133km² / 2.81M(2045)", "Hoan Kiem, Ba Dinh, Dong Da 등",
         "정치·행정·문화·역사·서비스", "홍강 코리도, 또리히 복원, 구도심 재개발"),
        ("🔴 북부 (동아인·BRG)", "632km² / 2.7-2.9M(2045)", "Dong Anh, Me Linh, Soc Son",
         "스마트시티·금융·국제 엔터·노이바이 관문", "BRG 스마트시티 $4.2B, 코로아 1,400ha"),
        ("🟣 서부 (호아락·선따이)", "17,000km²", "Hoa Lac, Xuan Mai, Son Tay, Ba Vi",
         "과학기술·고등교육·문화관광·생태", "호아락 하이테크파크, Tien Xuan 600ha"),
        ("🔵 동부 (롱비엔·자람)", "176km² / 850K(2030)", "Long Bien, Gia Lam",
         "동부 도시 확장·주거·산업·물류", "링로드4 자람구간, 홍강교량 5+개"),
        ("🟢 남부 (푸쑤엔·남부물류)", "303km² / 1.6M", "Phu Xuyen, Thuong Tin, Thanh Oai",
         "물류·산업지원·제2공항·공공서비스", "자빈 제2공항, 남부 물류허브"),
    ]
    for r_data in zone_rows:
        row=zt.add_row()
        for j,v in enumerate(r_data):
            row.cells[j].text=v
            for para in row.cells[j].paragraphs:
                for run in para.runs: run.font.size=Pt(8.5)
    doc.add_paragraph()
    doc.add_page_break()

    # ══ SECTIONS 2–6: 각 도시권 상세 ══
    zone_configs=[
        ("2. 🔴 북부 신도시",  "HN-NORTH", COLOR_NORTH, HEX["north"], zones["zone_2_northern"]),
        ("3. 🟣 서부 과학기술도시", "HN-WEST",  COLOR_WEST,  HEX["west"],  zones["zone_4_western"]),
        ("4. 🔵 동부 도시권",   "HN-EAST",  COLOR_EAST,  HEX["east"],  zones["zone_3_eastern"]),
        ("5. 🟢 남부 물류도시", "HN-SOUTH", COLOR_SOUTH, HEX["south"], zones["zone_5_southern"]),
        ("6. ⚡ 도시 인프라 공통", "HN-URBAN-INFRA", COLOR_INFRA, HEX["infra"], None),
    ]
    bucket_map={"HN-NORTH":"HN-URBAN-NORTH","HN-WEST":"HN-URBAN-WEST",
                "HN-EAST":"HN-URBAN-INFRA","HN-SOUTH":"HN-URBAN-INFRA",
                "HN-URBAN-INFRA":"HN-URBAN-INFRA"}

    for sec_title, zone_id, color, hexc, zone_data in zone_configs:
        _h(doc, sec_title, 1, color, 13); _hr(doc, hexc)

        bucket_id=bucket_map.get(zone_id, zone_id)
        arts=buckets.get(bucket_id, [])
        # 동부·남부는 extra_hanoi에서 보충
        if zone_id in ("HN-EAST","HN-SOUTH"):
            kws={"HN-EAST":["gia lam","long bien","ring road 4","gia lâm","long biên","vành đai 4"],
                 "HN-SOUTH":["phu xuyen","gia binh","southern hanoi","phú xuyên","sân bay gia bình"]}
            add=[a for a in extra_hanoi
                 if any(kw in ((a.get('title','') or '')+(a.get('summary_en','') or '')+(a.get('content','') or '')[:150]).lower()
                        for kw in kws.get(zone_id,[]))]
            arts=arts+add

        rep,yr=_select_arts(arts) if arts else ([],"-")

        if zone_data:
            # KPI 테이블
            kpi_t=doc.add_table(rows=3,cols=2); kpi_t.style="Table Grid"
            for ri,(k,v) in enumerate([
                ("기능", zone_data.get("function","-")),
                ("면적·인구",
                 f"{zone_data.get('area_sqkm','?')}km² | 2030: {zone_data.get('population_2030','?')} | 2045: {zone_data.get('population_2045','?')}"),
                ("행정구역", ", ".join(zone_data.get("districts",[])[:8])),
            ]):
                kpi_t.rows[ri].cells[0].text=k; kpi_t.rows[ri].cells[1].text=str(v)
                _hex_bg(kpi_t.rows[ri].cells[0], hexc)
                kpi_t.rows[ri].cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
                kpi_t.rows[ri].cells[0].paragraphs[0].runs[0].bold=True
                for cell in kpi_t.rows[ri].cells:
                    for para in cell.paragraphs:
                        for run in para.runs: run.font.size=Pt(9.5)
            doc.add_paragraph()

            # 공간 축
            if zone_data.get("key_axes"):
                _h(doc,"주요 공간 개발 축",2,color)
                for ax in zone_data["key_axes"]:
                    p=doc.add_paragraph(f"▪ {ax}",style="List Bullet"); p.runs[0].font.size=Pt(10)
                doc.add_paragraph()

            # 주요 프로젝트 테이블
            _h(doc,"주요 개발 프로젝트",2,color)
            _proj_table(doc, zone_data.get("key_projects",[]), color)

            # 한국 존재감
            if zone_data.get("korean_presence"):
                kp=doc.add_paragraph(); kp.paragraph_format.left_indent=Inches(0.2)
                kp.add_run("🇰🇷 한국 참여 현황: ").bold=True; kp.runs[0].font.color.rgb=_rgb(COLOR_KOR); kp.runs[0].font.size=Pt(10)
                kp.add_run(zone_data["korean_presence"]).font.size=Pt(10)
                doc.add_paragraph()
        else:
            # 인프라 공통 섹션
            _h(doc,"도시 인프라 공통 — 메트로·링로드·홍강·도심재개발",2,color)
            metro=hn["urban_rail_network"]; ring=hn["ring_roads"]
            it2=doc.add_table(rows=3,cols=2); it2.style="Table Grid"
            for ri,(k,v) in enumerate([
                ("메트로 목표","2035: 410km+ | 2065: 616km+ 15개 노선"),
                ("현재 운영","Line 2A 껫린-하동 (13km) + Line 3 년-하노이역 (12.5km 부분)"),
                ("링로드","RR3 완공 | RR3.5 건설중 | RR4 공사중(2023-2027, 112km, $1.5B)"),
            ]):
                it2.rows[ri].cells[0].text=k; it2.rows[ri].cells[1].text=v
                _hex_bg(it2.rows[ri].cells[0],hexc)
                it2.rows[ri].cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
                it2.rows[ri].cells[0].paragraphs[0].runs[0].bold=True
                for cell in it2.rows[ri].cells:
                    for para in cell.paragraphs:
                        for run in para.runs: run.font.size=Pt(9.5)
            doc.add_paragraph()

        # 최신 뉴스
        if rep:
            _h(doc,f"최신 뉴스 ({yr}, 상위 5건)",2,color)
            _add_articles(doc,rep,color,max_n=5)
            if len(rep)>5:
                mp=doc.add_paragraph(); mp.paragraph_format.left_indent=Inches(0.2)
                mp.add_run(f"  ※ 추가 {len(rep)-5}건 → Appendix").font.size=Pt(9)
        else:
            np2=doc.add_paragraph(); np2.paragraph_format.left_indent=Inches(0.2)
            np2.add_run(
                "ℹ️ 역사 DB에 이 도시권 전용 기사가 아직 충분하지 않습니다.\n"
                "   다음 주간 수집부터 키워드 매칭이 강화되어 기사가 누적됩니다."
            ).font.size=Pt(9.5)
        _hr(doc,"D1D5DB"); doc.add_paragraph()
        doc.add_page_break()

    # ══ SECTION 7: 스마트시티 기사 ══
    _h(doc,"7. 스마트시티 기술·플랫폼 기사 (VN-SC-2030)",1,COLOR_SC,13); _hr(doc,HEX["sc"])
    sc_note=doc.add_paragraph(); sc_note.paragraph_format.left_indent=Inches(0.2)
    sc_note.add_run(
        "이 섹션은 도시개발 '공간' 기사가 아닌, 스마트시티 '기술·플랫폼' 관련 기사를 분리 추적합니다.\n"
        "BRG 동아인 프로젝트처럼 양쪽에 해당하는 기사는 해당 도시권 섹션과 여기 중복 표시될 수 있습니다."
    ).font.size=Pt(9.5)
    doc.add_paragraph()
    sc_arts=buckets.get("VN-SC-2030",[])
    sc_rep,sc_yr=_select_arts(sc_arts) if sc_arts else ([],"-")
    if sc_rep:
        _add_articles(doc,sc_rep,COLOR_SC,max_n=8)
    else:
        doc.add_paragraph().add_run("ℹ️ 스마트시티 기술 기사가 수집되지 않았습니다.").font.size=Pt(9.5)
    doc.add_page_break()

    # ══ SECTION 8: 한국 기업 기회 ══
    _h(doc,"8. 한국 기업 기회 분석 — 하노이 5대 도시권",1,COLOR_KOR,13); _hr(doc,HEX["kor"])
    kor_t=doc.add_table(rows=1,cols=2); kor_t.style="Table Grid"
    for j,h in enumerate(["도시권/분야","한국 기업 기회"]):
        kor_t.rows[0].cells[j].text=h
        kor_t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _hex_bg(kor_t.rows[0].cells[j],HEX["kor"])
        kor_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for sec,opp in [
        ("🔴 북부 신도시 (동아인)",
         "HIGH: BRG-Sumitomo 스마트시티 — 삼성물산·현대건설 Phase 2-5 수주 기회\n"
         "HIGH: LH공사 스마트시티 기술 수출 (도시관리 플랫폼·스마트 인프라)\n"
         "HIGH: 코로아-동아인 1,400ha — 한국 건설사 EPC 컨소시엄\n"
         "HIGH: 노이바이 T3 — 인천공항공사 설계 자문 + 한국 건설사"),
        ("🟣 서부 과학기술도시 (호아락)",
         "HIGH: 호아락 R&D 센터 — 삼성·LG·현대차 연구소 설립\n"
         "HIGH: 하이테크파크 스마트팩토리 — 한국 제조 기업 입주\n"
         "MEDIUM: Tien Xuan 600ha 스마트 주거단지 — 한국 건설사 설계·시공\n"
         "MEDIUM: Urban Rail Line 8 (Metro, 호아락행) — 한국 메트로 컨소시엄"),
        ("🔵 동부·🟢 남부",
         "MEDIUM: 자람 도시개발 — 한국 부동산 개발사 (롯데건설·현대E&C)\n"
         "MEDIUM: 링로드4 — 한국 건설사 교량 패키지\n"
         "MEDIUM: 자빈 제2공항 — 한국항공우주산업 + 인천공항 컨소시엄\n"
         "MEDIUM: 남부 물류허브 — CJ대한통운·롯데로지스틱스 거점"),
        ("🏛️ 중심 (도심재개발)",
         "HIGH: 홍강 양안 개발 $1.5B+ — 한국 건설사 수상도시 설계·시공\n"
         "MEDIUM: 또리히 복원 — 한국 4대강 노하우 (환경부·K-water)\n"
         "MEDIUM: 내부도로 혼잡 해소 — 한국 스마트교통 ITS 솔루션"),
        ("⚡ 메트로·인프라",
         "HIGH: 하노이 메트로 3·5·6 — 한국 컨소시엄 ODA + 상업차관 패키지\n"
         "         Korea Eximbank ODA + 현대로템 차량 + LS Electric 변전\n"
         "MEDIUM: 스마트그리드 하노이 도심 — 한전KDN·LS Electric\n"
         "TIP: $2.5T 투자 중 민간 비중 높음 → PPP 구조 활용"),
        ("📊 전략 요약",
         "① 하노이 2026-2035 $556B 중 한국 기업 접근 가능 시장: 약 $30-50B 추정\n"
         "② 진입 시점: 지금이 가장 좋음 — 2026년 대규모 착공 시작\n"
         "③ 핵심 채널: 하노이 인민위원회 + 투자기획국 (DPI Hanoi)\n"
         "④ KOTRA 하노이 + 대한상의 베트남 활용 필수\n"
         "⑤ 한국 정부 ODA: EDCF (수출입은행) + KOICA — 메트로·환경 프로젝트"),
    ]:
        row=kor_t.add_row()
        row.cells[0].text=sec; row.cells[1].text=opp
        _hex_bg(row.cells[0],HEX["light"]); row.cells[0].paragraphs[0].runs[0].bold=True
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(9.5)
    doc.add_paragraph()
    doc.add_page_break()

    # ══ APPENDIX ══
    _h(doc,"Appendix — 하노이 관련 전체 기사",1,COLOR_MAIN); _hr(doc,HEX["main"])
    seen=set(); all_flat=[]
    for pid in HANOI_IDS:
        for art in buckets.get(pid,[]):
            url=art.get('url','') or str(id(art))
            if url not in seen:
                seen.add(url); art['_track']=pid[:15]; all_flat.append(art)
    for art in extra_hanoi:
        url=art.get('url','') or str(id(art))
        if url not in seen:
            seen.add(url); art['_track']='HN-Extra'; all_flat.append(art)
    all_flat.sort(key=_date_key,reverse=True)

    doc.add_paragraph(f"총 {len(all_flat)}건 — 상위 60건 표시").runs[0].font.size=Pt(9)
    doc.add_paragraph()
    ap_t=doc.add_table(rows=1,cols=5); ap_t.style="Table Grid"
    for j,h in enumerate(["No","날짜","제목","출처","분류"]):
        ap_t.rows[0].cells[j].text=h
        ap_t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _hex_bg(ap_t.rows[0].cells[j],"1E40AF")
        ap_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        ap_t.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)
    for ri,art in enumerate(all_flat[:60],1):
        row=ap_t.add_row()
        row.cells[0].text=str(ri); row.cells[1].text=_date_key(art)
        row.cells[2].text=_title(art)[:60]; row.cells[3].text=(art.get('source','') or '')[:18]
        row.cells[4].text=art.get('_track','')[:15]
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(8.5)

    fname=f"MI_REPORT_HANOI_URBAN_W{week:02d}_{now.strftime('%Y%m%d')}.docx"
    fpath=os.path.join(output_dir, fname)
    doc.save(fpath)
    return fpath


if __name__ == "__main__":
    out="/home/work/claw/outputs/reports/MI_Reports"
    Path(out).mkdir(parents=True,exist_ok=True)
    print("하노이 도시개발 MI 보고서 생성 중...")
    from pathlib import Path
    fpath=generate_hanoi_report(out)
    if fpath:
        sz=os.path.getsize(fpath)//1024
        print(f"✅ {fpath} ({sz}KB)")

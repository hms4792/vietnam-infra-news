"""
Decision 1894 / VN-ENV-IND-1894 Dedicated MI Report Generator
==============================================================
보고서 구조:
  Cover
  1. What is Decision 1894? (프로그램 성격 정확한 설명)
  2. National KPI 2030 (기술산업 자립도 목표)
  3. 6 Major Task Areas & Implementation Status
  4. Key Technology Sectors (세부 산업별)
  5. Province Activity — Recent News by Province
  6. Ministry Responsibilities & Progress
  7. Risk & Business Opportunity for Korean Companies
  8. Appendix
"""
import os, sys, json, re
from datetime import datetime, timedelta
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = Path("/home/work/claw")
COLOR  = (0x10, 0xB9, 0x81)   # emerald green
HEX_C  = "10B981"

def _rgb(t): return RGBColor(*t)
def _set_bg(cell, h):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),h)
    p.append(s)

def _hr(doc, color="CCCCCC"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'), color)
    pb.append(bot); pPr.append(pb)

def _h(doc, text, lvl, color=COLOR, size=None):
    h = doc.add_heading(text, lvl)
    for r in h.runs:
        r.font.color.rgb = _rgb(color)
        if size: r.font.size = Pt(size)
    return h

def _date_key(a):
    d = str(a.get('published_date','') or '')
    m = re.search(r'(20\d{2})', d)
    yr = m.group(1) if m else '1900'
    if not d[:4].isdigit():
        try: return datetime.strptime(d.strip(),'%b %d, %Y').strftime('%Y-%m-%d')
        except: return f'{yr}-01-01'
    return d[:10]

def _is_vi(t):
    return any(c in (t or '') for c in ['ă','ơ','ư','đ','ấ','ề','ộ','ừ','ạ','ọ','ổ','ị','ế'])

def _display_title(art):
    t = art.get('title','')
    if _is_vi(t):
        en = art.get('summary_en','')
        if en and not _is_vi(en): return f"[VI] {en[:80]}"
        ko = art.get('summary_ko','')
        if ko and not _is_vi(ko): return f"[VI] {ko[:70]}"
    return t

def _summary(art):
    ko = art.get('summary_ko','')
    en = art.get('summary_en','')
    if ko and not _is_vi(ko): return '🇰🇷', ko
    if en and not _is_vi(en): return '🇺🇸', en
    c = art.get('content','')
    if c: return '📄', c[:250]
    return None, None

def _relevance_check_d1894(art):
    """D1894 전용 관련성 필터 — 기술산업 기사만"""
    # 반드시 포함해야 할 키워드
    must_kws = [
        "environmental industry","công nghiệp môi trường","decision 1894","1894",
        "recycling industrial","eco-industrial","environmental technology",
        "waste treatment technology","wastewater technology","wte technology",
        "emission control technology","environmental equipment","monitoring equipment",
        "circular economy","kinh tế tuần hoàn","epr","extended producer",
        "green industry","environmental enterprise","CNMT","moit environment",
        "environmental monitoring","pollution control equipment",
        "biodegradable","green material","landfill gas","biogas technology",
    ]
    # 제외 키워드 (건설/시공 계약 기사 — 인프라 플랜에 해당)
    exclude_kws = [
        "wwtp construction contract","wastewater plant contract","landfill operation contract",
        "solid waste collection bid","garbage truck procurement"
    ]
    text = (
        (art.get('title','') or '') + ' ' +
        (art.get('summary_en','') or '') + ' ' +
        (art.get('summary_ko','') or '') + ' ' +
        (art.get('content','') or '')[:300]
    ).lower()
    
    if any(ex in text for ex in exclude_kws):
        return False
    return any(kw.lower() in text for kw in must_kws)

def generate_d1894_report(output_dir: str) -> str:
    # ── 데이터 로드 ──
    with open(BASE/"config/history_db.json", encoding='utf-8') as f:
        hdb = json.load(f)
    with open(BASE/"config/d1894_program_structure.json", encoding='utf-8') as f:
        prog = json.load(f)

    # D1894 매칭 기사 수집
    all_arts = [a for a in hdb['articles'].values()
                if 'VN-ENV-IND-1894' in (a.get('matched_plans') or [])]
    # 관련성 재필터
    relevant = [a for a in all_arts if _relevance_check_d1894(a)]
    
    # 2026년 기사 우선
    arts_2026 = sorted([a for a in relevant if re.search(r'2026', str(a.get('published_date','')))],
                       key=_date_key, reverse=True)
    report_arts = arts_2026 if arts_2026 else sorted(
        [a for a in relevant if re.search(r'2025', str(a.get('published_date','')))],
        key=_date_key, reverse=True
    )
    yr_label = "2026" if arts_2026 else "2025 (most recent)"
    
    now   = datetime.now()
    week  = now.isocalendar()[1]
    
    doc = Document()
    for sec in doc.sections:
        sec.top_margin=Inches(1.0); sec.bottom_margin=Inches(1.0)
        sec.left_margin=Inches(1.2); sec.right_margin=Inches(1.2)
    
    # ══ COVER ══
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1); tbl.style="Table Grid"
    c = tbl.rows[0].cells[0]; _set_bg(c, "0F172A")
    p = c.paragraphs[0]
    r = p.add_run("VIETNAM INFRASTRUCTURE INTELLIGENCE HUB")
    r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2 = tp.add_run("Vietnam Environmental Industry Development Program")
    r2.bold=True; r2.font.size=Pt(18); r2.font.color.rgb=_rgb(COLOR)
    
    sp = doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run("Decision No. 1894/QD-TTg | September 4, 2025").font.size=Pt(11)
    
    doc.add_paragraph()
    it = doc.add_table(rows=5, cols=2); it.style="Table Grid"
    for i,(k,v) in enumerate([
        ("Plan ID", "VN-ENV-IND-1894"),
        ("Official Title", "Program for the Development of Vietnam's Environmental Industry Sector 2025–2030"),
        ("Lead Ministry", "Ministry of Industry and Trade (MOIT)"),
        ("Signed By", "Deputy PM Tran Hong Ha — September 4, 2025"),
        ("Report Articles", f"Matched: {len(all_arts)} | Relevant: {len(relevant)} | Report ({yr_label}): {len(report_arts)}"),
    ]):
        it.rows[i].cells[0].text=k; it.rows[i].cells[1].text=v
        _set_bg(it.rows[i].cells[0], HEX_C)
        it.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        it.rows[i].cells[0].paragraphs[0].runs[0].bold=True
        for cell in it.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(10)
    
    doc.add_paragraph()
    conf_p = doc.add_paragraph(); conf_p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    conf_p.add_run(f"Generated: {now.strftime('%Y-%m-%d %H:%M')}  |  W{week}/{now.year}  |  CONFIDENTIAL").font.size=Pt(9)
    doc.add_page_break()
    
    # ══ SECTION 1: What is Decision 1894? ══
    _h(doc, "1. What is Decision 1894?", 1, COLOR, 14); _hr(doc, HEX_C)
    
    distinction_box = doc.add_table(rows=1, cols=1); distinction_box.style="Table Grid"
    dc = distinction_box.rows[0].cells[0]; _set_bg(dc, "ECFDF5")
    dp = dc.paragraphs[0]
    dr = dp.add_run("⚠️  IMPORTANT DISTINCTION")
    dr.bold=True; dr.font.size=Pt(10); dr.font.color.rgb=_rgb(COLOR)
    dc.add_paragraph()
    d2 = dc.add_paragraph()
    d2.add_run(
        "Decision 1894 is NOT a waste infrastructure construction plan.\n"
        "It is a NATIONAL INDUSTRIAL DEVELOPMENT PROGRAM to build up Vietnam's domestic "
        "environmental technology industry — i.e., developing Vietnamese-made technology, "
        "equipment, and services for environmental protection.\n\n"
        "This is distinct from:\n"
        "  • VN-WW-2030 (wastewater treatment plant construction)\n"
        "  • Solid Waste Infrastructure Plans (landfill/incinerator construction)\n"
        "  • VN-ENV infrastructure programs\n\n"
        "D1894 focuses on WHO MAKES the technology, not WHO BUILDS the plants."
    ).font.size=Pt(10)
    doc.add_paragraph()
    
    overview_p = doc.add_paragraph()
    overview_p.paragraph_format.left_indent=Inches(0.2)
    overview_p.add_run(prog['program_identity']['nature']).font.size=Pt(10)
    doc.add_paragraph()
    
    _h(doc, "Core Philosophy", 2, COLOR)
    for item in prog['core_philosophy']:
        p = doc.add_paragraph(f"▪ {item}", style="List Bullet")
        p.runs[0].font.size=Pt(10)
    doc.add_paragraph()
    
    # ══ SECTION 2: National KPI 2030 ══
    _h(doc, "2. National KPI Targets — 2030", 1, COLOR, 13); _hr(doc, HEX_C)
    
    kpi = prog['national_kpi_2030']
    kpi_tbl = doc.add_table(rows=1, cols=4); kpi_tbl.style="Table Grid"
    for j,h in enumerate(["Technology Sector","2030 Target (Domestic Market Share)","Scope","Current Gap"]):
        kpi_tbl.rows[0].cells[j].text=h
        kpi_tbl.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _set_bg(kpi_tbl.rows[0].cells[j], HEX_C)
        kpi_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        kpi_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)
    
    kpi_rows = [
        ("Wastewater Treatment Technology", "70–80% of domestic demand + export", "WWTP equipment, modules, chemicals", "High import dependency (Korea, Japan, EU)"),
        ("Exhaust Gas Treatment Technology", "60–70% of domestic demand", "Air pollution control equipment", "Advanced systems almost fully imported"),
        ("Solid & Hazardous Waste Technology", "50–60% (collection/transport)\n60–70% (sorting/recycling)", "Collection vehicles, sorting lines, recycling machinery", "Incineration/hazardous treatment relies on imports"),
        ("Environmental Monitoring Equipment", "20% of domestic demand", "IoT sensors, analyzers, AMS stations", "Almost entirely import-dependent"),
        ("Waste-to-Energy Technology", "Domestic capability established", "WTE boilers, biogas, RDF systems", "Large-scale WTE uses Chinese/EU technology"),
    ]
    for sector, target, scope, gap in kpi_rows:
        row = kpi_tbl.add_row()
        row.cells[0].text=sector; row.cells[1].text=target
        row.cells[2].text=scope; row.cells[3].text=gap
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(9)
    doc.add_paragraph()
    
    # Policy & Institutional KPIs
    _h(doc, "2-1. Policy & Institutional KPIs", 2, COLOR)
    pol_tbl = doc.add_table(rows=1, cols=3); pol_tbl.style="Table Grid"
    for j,h in enumerate(["Milestone","Target","Timeline"]):
        pol_tbl.rows[0].cells[j].text=h
        pol_tbl.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _set_bg(pol_tbl.rows[0].cells[j], "F0FDF4")
        pol_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)
    for milestone, target, tl in [
        ("Recycling Industrial Parks", "Legal framework + pilot parks established", "By 2027"),
        ("HS Codes for Env. Goods", "Environmental goods HS codes issued", "2025 (immediate)"),
        ("Tax Policy Reform", "Tax incentives for environmental goods + tech transfer", "2025–2026"),
        ("Government Decree", "Decree specifically regulating environmental industry", "By 2030"),
        ("National Env. Industry Center", "Established and operational", "By 2028"),
        ("National Statistics", "Environmental industry added to national stat system", "2025–2026"),
        ("Foreign Investment", "FDI attracted into environmental technology sector", "Ongoing"),
    ]:
        row = pol_tbl.add_row()
        row.cells[0].text=milestone; row.cells[1].text=target; row.cells[2].text=tl
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(9)
    doc.add_paragraph()
    
    # ══ SECTION 3: 6 Major Task Areas ══
    _h(doc, "3. Six Major Task Areas & Implementation Status", 1, COLOR, 13); _hr(doc, HEX_C)
    
    tasks = prog['six_major_task_areas']
    status_map = {
        "Task_1_Policy_Institutional": ("🟡 In Progress", "HS code issuance initiated; tax policy review underway; Decree drafting not yet started"),
        "Task_2_Market_Development":   ("🟡 Early Stage", "Government procurement mechanisms being designed; incubator network concept phase"),
        "Task_3_RnD_Innovation_Digital":("🟡 In Progress", "Decision 980 catalogue review underway; RE waste recycling technology R&D initiated"),
        "Task_4_Human_Resources":       ("⚪ Planning", "Curriculum development not yet reported"),
        "Task_5_Digital_Platform":      ("🟡 In Progress", "Environmental industry portal development announced by MOIT (Oct 2025)"),
        "Task_6_Communications":        ("🟢 Active", "MOIT held environmental industry development conference (Oct 2025); media campaigns launched"),
    }
    
    for task_key, task_data in tasks.items():
        status, note = status_map.get(task_key, ("⚪ Not Started", ""))
        _h(doc, f"{task_data['title']}", 2, COLOR)
        
        info_t = doc.add_table(rows=2, cols=2); info_t.style="Table Grid"
        info_t.rows[0].cells[0].text="Lead Ministry"
        info_t.rows[0].cells[1].text=task_data.get('lead','-')
        info_t.rows[1].cells[0].text="Implementation Status"
        info_t.rows[1].cells[1].text=f"{status}  {note}"
        for ri in range(2):
            _set_bg(info_t.rows[ri].cells[0], "F0FDF4")
            info_t.rows[ri].cells[0].paragraphs[0].runs[0].bold=True
            for cell in info_t.rows[ri].cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.font.size=Pt(9)
        doc.add_paragraph()
        
        _h(doc, "Sub-Tasks", 3, COLOR)
        for sub in task_data.get('sub_tasks', []):
            p = doc.add_paragraph(f"▪ {sub}", style="List Bullet")
            p.runs[0].font.size=Pt(9.5)
        doc.add_paragraph()
        _hr(doc, "E2E8F0"); doc.add_paragraph()
    
    doc.add_page_break()
    
    # ══ SECTION 4: Key Technology Sectors ══
    _h(doc, "4. Key Technology Sectors — Market Potential", 1, COLOR, 13); _hr(doc, HEX_C)
    
    for sector_key, sector_data in prog['key_technology_sectors'].items():
        _h(doc, sector_data['description'], 2, COLOR)
        
        st = doc.add_table(rows=4, cols=2); st.style="Table Grid"
        for i,(k,v) in enumerate([
            ("2030 Target", sector_data.get('target_2030', sector_data.get('target','-'))),
            ("Key Products / Equipment", ", ".join(sector_data.get('key_products',[]))),
            ("Current Gap", sector_data.get('current_gaps', sector_data.get('current_status','-'))),
            ("Concept", sector_data.get('concept','')),
        ]):
            if not v: continue
            st.rows[i].cells[0].text=k; st.rows[i].cells[1].text=v
            _set_bg(st.rows[i].cells[0], "F0FDF4")
            st.rows[i].cells[0].paragraphs[0].runs[0].bold=True
            for cell in st.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.font.size=Pt(9.5)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ══ SECTION 5: Province Activity — Recent News ══
    _h(doc, f"5. Province Activity — Recent News ({yr_label})", 1, COLOR, 13); _hr(doc, HEX_C)
    
    note_p = doc.add_paragraph()
    note_p.paragraph_format.left_indent=Inches(0.2)
    nr = note_p.add_run("ℹ️ Note: ")
    nr.bold=True; nr.font.size=Pt(10)
    note_p.add_run(
        "The articles below are collected from news sources and represent recent activities "
        "related to the D1894 program areas. They serve as progress indicators, NOT as official "
        "program reports. Some articles may cover related environmental industry activities "
        "without explicitly referencing Decision 1894."
    ).font.size=Pt(9.5)
    doc.add_paragraph()
    
    # Province별 기사 버킷팅
    art_by_prov = defaultdict(list)
    for a in report_arts:
        pv = a.get('province','') or 'National'
        pv_map = {'vietnam':'National','national / unspecified':'National',
                  'ho chi minh city':'Ho Chi Minh City','hcmc':'Ho Chi Minh City',
                  'hanoi':'Hanoi','ha noi':'Hanoi','da nang':'Da Nang','danang':'Da Nang'}
        pv = pv_map.get(pv.lower(), pv)
        art_by_prov[pv].append(a)
    
    # Province KPI 데이터
    prov_kpis = prog.get('province_role',{}).get('key_provinces',{})
    
    # 순서: KPI 있는 지역 먼저, 그 다음 기사 있는 지역
    ordered_provs = list(prov_kpis.keys())
    extra = [p for p in art_by_prov.keys() if p not in ordered_provs]
    
    for prov in ordered_provs + sorted(extra, key=lambda x:-len(art_by_prov.get(x,[]))):
        arts = art_by_prov.get(prov, [])
        prov_kpi = prov_kpis.get(prov, {})
        if not arts and not prov_kpi:
            continue
        
        _h(doc, f"📍 {prov}", 2, COLOR)
        
        if prov_kpi:
            kp = doc.add_paragraph()
            kp.paragraph_format.left_indent=Inches(0.2)
            r1 = kp.add_run("Role: "); r1.bold=True; r1.font.size=Pt(10)
            kp.add_run(prov_kpi.get('role','')).font.size=Pt(10)
            op = doc.add_paragraph()
            op.paragraph_format.left_indent=Inches(0.2)
            r2 = op.add_run("Opportunities: "); r2.bold=True; r2.font.size=Pt(10); r2.font.color.rgb=_rgb(COLOR)
            op.add_run(prov_kpi.get('opportunities','')).font.size=Pt(10)
            doc.add_paragraph()
        
        if arts:
            _h(doc, f"Recent News ({len(arts)} articles)", 3, COLOR)
            for i, art in enumerate(arts[:5], 1):
                title = _display_title(art)
                flag, summary = _summary(art)
                date  = _date_key(art)
                src   = art.get('source','') or ''
                
                ap = doc.add_paragraph()
                ap.paragraph_format.left_indent=Inches(0.2)
                nr2 = ap.add_run(f"[{i}] "); nr2.bold=True; nr2.font.size=Pt(10); nr2.font.color.rgb=_rgb(COLOR)
                ap.add_run(title[:85]).bold=True; ap.runs[-1].font.size=Pt(10)
                
                mp = doc.add_paragraph()
                mp.paragraph_format.left_indent=Inches(0.4)
                mp.add_run(f"📅 {date}  |  📰 {src[:30]}").font.size=Pt(8.5)
                
                if summary and flag:
                    sp2 = doc.add_paragraph()
                    sp2.paragraph_format.left_indent=Inches(0.4)
                    sp2.add_run(f"{flag} ").font.size=Pt(9.5)
                    sp2.add_run(summary[:220]).font.size=Pt(9.5)
                
                # English summary
                en = art.get('summary_en','')
                if flag == '🇰🇷' and en and not _is_vi(en):
                    ep = doc.add_paragraph()
                    ep.paragraph_format.left_indent=Inches(0.4)
                    ep.add_run("🇺🇸 ").font.size=Pt(9.5)
                    ep.add_run(en[:200]).font.size=Pt(9.5)
                
                up = doc.add_paragraph()
                up.paragraph_format.left_indent=Inches(0.4)
                ur = up.add_run(art.get('url','')[:80])
                ur.font.size=Pt(8); ur.font.color.rgb=RGBColor(0x1D,0x4E,0xD8)
                doc.add_paragraph()
            
            if len(arts) > 5:
                mp2 = doc.add_paragraph()
                mp2.paragraph_format.left_indent=Inches(0.2)
                mp2.add_run(f"  ※ {len(arts)-5} more articles → Appendix").font.size=Pt(9)
        else:
            np2 = doc.add_paragraph()
            np2.paragraph_format.left_indent=Inches(0.2)
            np2.add_run("ℹ️ No recent news collected for this province.").font.size=Pt(9)
        
        _hr(doc, "E2E8F0"); doc.add_paragraph()
    
    doc.add_page_break()
    
    # ══ SECTION 6: Ministry Responsibilities ══
    _h(doc, "6. Ministry Responsibilities & Coordination", 1, COLOR); _hr(doc, HEX_C)
    
    min_tbl = doc.add_table(rows=1, cols=2); min_tbl.style="Table Grid"
    for j,h in enumerate(["Ministry / Agency","Responsibility"]):
        min_tbl.rows[0].cells[j].text=h
        min_tbl.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _set_bg(min_tbl.rows[0].cells[j], HEX_C)
        min_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for ministry, resp in prog['ministry_responsibilities'].items():
        row = min_tbl.add_row()
        row.cells[0].text=ministry.replace('_',' '); row.cells[1].text=resp
        if ministry == 'MOIT': _set_bg(row.cells[0], "D1FAE5")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(9.5)
    doc.add_paragraph()
    
    # Implementation Timeline
    _h(doc, "6-1. Implementation Timeline", 2, COLOR)
    for phase, tasks_list in prog['implementation_timeline'].items():
        tp2 = doc.add_paragraph()
        tr3 = tp2.add_run(f"📅 {phase.replace('_',' ')}: ")
        tr3.bold=True; tr3.font.size=Pt(10); tr3.font.color.rgb=_rgb(COLOR)
        tp2.add_run("  ·  ".join(tasks_list)).font.size=Pt(9.5)
    doc.add_paragraph()
    
    # ══ SECTION 7: Risk & Korean Business Opportunity ══
    _h(doc, "7. Risk Assessment & Korean Business Opportunity", 1, COLOR); _hr(doc, HEX_C)
    
    rb_tbl = doc.add_table(rows=1, cols=2); rb_tbl.style="Table Grid"
    for j,h in enumerate(["Category","Analysis"]):
        rb_tbl.rows[0].cells[j].text=h
        rb_tbl.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _set_bg(rb_tbl.rows[0].cells[j], HEX_C)
        rb_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    
    for cat, analysis in [
        ("🚨 Key Risks",
         "1. Implementation speed uncertain — policy framework still being drafted\n"
         "2. Private sector capacity to deliver domestic technology remains low\n"
         "3. Competing priorities: infrastructure construction demand vs. technology industry development\n"
         "4. Limited access to green credit for SME environmental tech companies\n"
         "5. Foreign technology still dominant — localization targets ambitious"),
        
        ("💡 Korean Company Opportunities",
         "HIGH PRIORITY:\n"
         "▪ Technology Transfer Partnership: Korean env. tech companies can transfer technology to Vietnamese partners → preferred vendor status + tax incentives\n"
         "▪ Joint Venture in Environmental Equipment: Local manufacturing of Korean-designed WWTP modules, monitoring systems\n"
         "▪ Environmental Industry Incubator: Korean companies can co-develop SME incubators (Korea has proven models via KOICA/KOTRA)\n\n"
         "MEDIUM PRIORITY:\n"
         "▪ Recycling Industrial Park Development: Korea LH Corporation already active (Hung Yen Smart City) — D1894 recycling park framework creates new entry points\n"
         "▪ Monitoring Equipment Manufacturing: Korean sensors/analyzers → partner with local assembly to meet 20% domestic target\n"
         "▪ WTE Technology: Korean WTE companies (KEPCO, Samsung C&T) can position for technology transfer\n\n"
         "SPECIFIC ENTRY POINTS:\n"
         "▪ Decision 980 catalogue review → submit Korean technologies for inclusion\n"
         "▪ National Environmental Industry Center tender when announced\n"
         "▪ Government procurement programs → preferred domestic partner status through JV"),
        
        ("📈 Market Outlook",
         "2025–2026: Policy framework phase — limited immediate market but critical for positioning\n"
         "2026–2028: Market formation — procurement programs begin; incubators launch; early technology transfer deals\n"
         "2028–2030: Scale-up phase — domestic production ramps; export ambitions materialize\n\n"
         "Total environmental technology market size: est. $2–3B by 2030\n"
         "Technology transfer deals: expected $500M+ over program period\n"
         "MOIT actively seeking international partners — Korean companies have strong positioning via existing government relationships"),
        
        ("⚠️ Strategic Note",
         "D1894 is a TECHNOLOGY INDUSTRY program. Business strategies should focus on:\n"
         "1. PARTNERSHIP MODEL (not just export) — Vietnamese companies need technology, not finished goods\n"
         "2. EARLY POSITIONING — policy framework phase is the best time to engage\n"
         "3. MOIT is the primary engagement target (not MONRE) — this is an industry program\n"
         "4. KOTRA/KOICA/EXIM Bank programs can support Korean company entry"),
    ]:
        row = rb_tbl.add_row()
        row.cells[0].text=cat; row.cells[1].text=analysis
        _set_bg(row.cells[0], "F0FDF4")
        row.cells[0].paragraphs[0].runs[0].bold=True
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(9.5)
    doc.add_paragraph()
    
    # ══ SECTION 8: Appendix ══
    _h(doc, f"8. Appendix — All Relevant Articles ({len(relevant)} articles, filter-passed)", 1, COLOR)
    _hr(doc, HEX_C)
    
    ap_note = doc.add_paragraph(
        f"Total matched: {len(all_arts)} | Relevance-filtered: {len(relevant)} | "
        f"Excluded as infrastructure construction news (not D1894-specific): {len(all_arts)-len(relevant)}\n"
        "Full article history available in Database Excel — 'D1894' sheet."
    )
    ap_note.runs[0].font.size=Pt(9); ap_note.runs[0].font.color.rgb=RGBColor(0x64,0x74,0x8B)
    doc.add_paragraph()
    
    if relevant:
        ap_t = doc.add_table(rows=1, cols=5); ap_t.style="Table Grid"
        for j,h in enumerate(["No","Date","Title (EN/KO)","Source","Province"]):
            ap_t.rows[0].cells[j].text=h
            ap_t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
            _set_bg(ap_t.rows[0].cells[j], "1E40AF")
            ap_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            ap_t.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)
        for ri, art in enumerate(sorted(relevant, key=_date_key, reverse=True)[:60], 1):
            row = ap_t.add_row()
            row.cells[0].text=str(ri)
            row.cells[1].text=_date_key(art)
            row.cells[2].text=_display_title(art)[:60]
            row.cells[3].text=(art.get('source','') or '')[:20]
            row.cells[4].text=(art.get('province','') or '')[:20]
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.font.size=Pt(8.5)
    
    # ── 저장 ──
    fname = f"MI_REPORT_D1894_DEEP_W{week:02d}_{now.strftime('%Y%m%d')}.docx"
    fpath = os.path.join(output_dir, fname)
    doc.save(fpath)
    return fpath

if __name__ == "__main__":
    out = "/home/work/claw/outputs/reports/MI_Reports"
    Path(out).mkdir(parents=True, exist_ok=True)
    print("D1894 심층 보고서 생성 중...")
    fpath = generate_d1894_report(out)
    if fpath:
        sz = os.path.getsize(fpath) // 1024
        print(f"✅ {fpath} ({sz}KB)")

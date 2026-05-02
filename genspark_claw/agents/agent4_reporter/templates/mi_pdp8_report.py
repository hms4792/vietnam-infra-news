"""
mi_pdp8_report_v4.py — PDP8 에너지 통합 보고서 v4.0
=====================================================
Decision 768/QD-TTg (April 15, 2025) 기반
6개 하부 트래킹 → 1개 통합 KPI 보고서
"""
import os, json
from datetime import datetime
from pathlib import Path
from collections import defaultdict

from report_lib import (
    new_doc, add_heading, add_hr, add_note, add_bullet,
    styled_table, kpi_box, info_box, build_cover,
    render_articles, render_history_timeline,
    render_province_section, render_korean_opportunity,
    build_appendix, load_history, load_kpi_db, save_report,
    select_articles, get_year_trend, _date_key,
    PALETTE, PALETTE_HEX, _rgb, cell_bg,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path("/home/work/claw")

PDP8_SUBS = [
    {"id":"VN-PDP8-RENEWABLE","name":"재생에너지 (태양광·풍력·수력)",
     "color":"amber","hex":"D97706",
     "kpis":[("태양광 2030","46,459–73,416 MW","현재 ~21,000MW"),
             ("육상풍력 2030","26,066–38,029 MW","현재 ~5,000MW"),
             ("해상풍력 2030","6,000–17,032 MW","현재 ~1,000MW")]},
    {"id":"VN-PDP8-LNG","name":"LNG 발전 & 인프라",
     "color":"orange","hex":"EA580C",
     "kpis":[("LNG 발전용량","22,524 MW by 2030","현재 ~0MW (건설중)"),
             ("LNG 수입 터미널","5개소 이상","현재 Thi Vai 1개"),
             ("LNG 연간 수요","14~18 BCM by 2030","현재 ~6 BCM")]},
    {"id":"VN-PDP8-NUCLEAR","name":"원자력 (닌투언 1·2호기)",
     "color":"violet","hex":"7C3AED",
     "kpis":[("원자력 설비용량","4,000–6,400 MW","2030-2035 상업운전"),
             ("닌투언 1호기","2,000 MW (Russia ROSATOM)","설계 재개 2025"),
             ("닌투언 2호기","2,000 MW (Japan JV)","착공 계획 2028")]},
    {"id":"VN-PDP8-COAL","name":"석탄 전환·폐지 로드맵",
     "color":"gray","hex":"374151",
     "kpis":[("석탄 설비용량 2030","31,055 MW (신규 없음)","현재 ~26,000MW"),
             ("석탄 폐지 목표","2050년 0MW","Just Energy Transition"),
             ("조기 전환 검토","10개 노후 발전소","ADB·WB 지원")]},
    {"id":"VN-PDP8-GRID","name":"송전망·스마트그리드",
     "color":"teal","hex":"0D9488",
     "kpis":[("500kV 초고압 송전","회로 증설 2,400km+","현재 6,500km"),
             ("스마트미터 보급","100% 도시 2030","현재 ~30%"),
             ("배터리 저장(BESS)","3,000 MW by 2030","현재 시범단계")]},
    {"id":"VN-PDP8-HYDROGEN","name":"수소·그린에너지 (RE 수출)",
     "color":"green","hex":"15803D",
     "kpis":[("그린수소 생산","시범 → 상업화 2030","전략 수립 단계"),
             ("RE 수출 용량","5,000–10,000 MW by 2035","싱가포르·말레이시아"),
             ("수소 경제 로드맵","2030 마스터플랜","MPI 수립 중")]},
]

THEME_HEX = "D97706"
THEME_LIGHT = "FFFBEB"
THEME_COLOR = "amber"

def generate_pdp8_report(output_dir: str) -> str:
    now = datetime.now()
    week = now.isocalendar()[1]
    week_str = f"W{week:02d}/{now.year}"

    all_ids = [s["id"] for s in PDP8_SUBS]
    buckets = load_history(all_ids)
    all_arts = []
    seen = set()
    for pid, arts in buckets.items():
        for a in arts:
            k = a.get("url","") or str(id(a))
            if k not in seen: seen.add(k); all_arts.append(a)
    all_arts.sort(key=_date_key, reverse=True)

    doc = new_doc()

    # 표지
    build_cover(doc,
        "PDP8-INTEGRATED",
        "PDP8 에너지 통합 보고서",
        "Power Development Plan VIII — Integrated Report (Decision 768/QD-TTg)",
        "Decision 768/QD-TTg (April 15, 2025) — Revised PDP8",
        "Ministry of Industry & Trade (MOIT)",
        "$136.3B (2026–2030) | 총 설비 183,291–236,363 MW (2030)",
        "2021 – 2030 / Vision 2050",
        len(all_arts), week_str,
        "🔋 에너지 통합 KPI 보고서", THEME_HEX)

    # Section 1: PDP8 개요
    add_heading(doc, "1. PDP8 개요 — Decision 768/QD-TTg", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)

    ot = doc.add_table(rows=1, cols=1); ot.style = "Table Grid"
    oc = ot.rows[0].cells[0]; cell_bg(oc, THEME_LIGHT)
    op = oc.paragraphs[0]
    op.add_run(
        "2025년 4월 15일 결정된 Decision 768/QD-TTg는 2023년 원본(Decision 500)을 대폭 개정한 "
        "베트남 2030 전력개발계획 최신 버전. "
        "3대 기둥: ① 에너지 안보 ② 공정전환(Just Transition) ③ RE 산업생태계 육성. "
        "$136.3B(2026-2030) 투자로 재생에너지 주도 전력 시스템 구축."
    ).font.size = Pt(10)
    doc.add_paragraph()

    # 핵심 수치 KPI 박스
    kpi_box(doc,
        [("총 설비용량 2030","183,291~236,363","MW"),
         ("태양광 비중","25~31%","최대 73,416MW"),
         ("투자 규모","$136.3B","2026-2030"),
         ("원자력 재개","4,000~6,400","MW (2030-35)")],
        bg_hex=THEME_LIGHT, border_hex=THEME_HEX,
        title="PDP8 핵심 수치 (Decision 768, 2025.04.15)", font_size=10)

    # 전원믹스 테이블
    add_heading(doc, "1-1. 전원별 설비용량 목표 (2030)", 2, THEME_COLOR)
    styled_table(doc,
        ["전원","설비용량 (MW)","비중","비고"],
        [["태양광 (육상+부유식)",  "46,459–73,416", "25–31%", "최대 비중"],
         ["육상 풍력",             "26,066–38,029", "14–16%", "4개 풍력벨트"],
         ["해상 풍력",             "6,000–17,032",  "3–7%",   "남중국해·통킹만"],
         ["LNG 발전",              "22,524",         "10%",    "수입 LNG"],
         ["석탄 (신규 없음)",       "31,055",         "13%",    "2050년 0"],
         ["원자력",                "4,000–6,400",    "2–3%",   "닌투언 1·2"],
         ["수력",                  "29,346",         "12%",    "증설 최소"],
         ["BESS 저장",             "3,000+",         "-",      "신설"]],
        col_widths=[2.0, 1.8, 0.9, 2.6],
        header_bg=THEME_COLOR, font_size=9.5)
    doc.add_page_break()

    # Section 2: KPI 현황표
    add_heading(doc, "2. 하부 트래킹 KPI 현황표", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)
    kpi_rows = []
    for sub in PDP8_SUBS:
        for kpi_name, target, current in sub["kpis"]:
            kpi_rows.append([sub["name"][:25], kpi_name[:30], target[:30], current[:35]])
    styled_table(doc,
        ["하부 트래킹", "KPI 지표", "목표치", "현재 수준"],
        kpi_rows,
        col_widths=[2.0, 2.2, 1.8, 2.3],
        header_bg=THEME_COLOR, font_size=9)
    doc.add_page_break()

    # Section 3: 하부 트래킹별 상세
    add_heading(doc, "3. 하부 트래킹별 상세 진행현황", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)

    for sub in PDP8_SUBS:
        pid = sub["id"]; color = sub["color"]; hex_c = sub["hex"]
        sub_arts = sorted(buckets.get(pid, []), key=_date_key, reverse=True)

        add_heading(doc, f"▶ {sub['name']}  [{pid}]", 2, color)

        # KPI 미니박스
        kpi_items = [(k[:20], v[:25], "") for k,v,_ in sub["kpis"][:4]]
        kpi_box(doc, kpi_items, bg_hex="FFFBEB", border_hex=hex_c,
                title=f"{sub['name']} — 목표 KPI", font_size=9)

        # 연도별 트렌드
        yr_trend = get_year_trend(sub_arts, pid)
        trend_items = [(y, yr_trend.get(y, 0), "건") for y in ["2023","2024","2025","2026"]]
        kpi_box(doc, trend_items, bg_hex="F8FAFC", border_hex=hex_c,
                title="연도별 기사 건수", font_size=9)

        # 최신 뉴스
        rec, yr_lbl = select_articles(sub_arts)
        if rec:
            add_heading(doc, f"  📰 최신 뉴스 ({yr_lbl}, 상위 4건)", 3, color)
            render_articles(doc, rec, color, max_n=4)

        # 역사 타임라인
        hist = [a for a in sub_arts if _date_key(a) < "2026-01-01"]
        if hist:
            add_heading(doc, "  📚 역사 기사 타임라인", 3, color)
            render_history_timeline(doc, hist, color)

        if not sub_arts:
            add_note(doc, "  ℹ️ 관련 기사 수집 중 — 다음 주간 수집 후 업데이트됩니다.", "gray")

        add_hr(doc, "E5E7EB"); doc.add_paragraph()

    doc.add_page_break()

    # Section 4: 시장 동향
    add_heading(doc, "4. 시장 동향 분석", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)
    add_heading(doc, "4-1. 전체 연도별 기사 트렌드", 2, THEME_COLOR)
    yr_all = get_year_trend(all_arts, "PDP8")
    styled_table(doc,
        ["2019","2020","2021","2022","2023","2024","2025","2026"],
        [[str(yr_all.get(str(y),0)) for y in range(2019,2027)]],
        header_bg=THEME_COLOR, font_size=9)
    add_heading(doc, "4-2. 지역별 활동", 2, THEME_COLOR)
    render_province_section(doc, all_arts, "VN-PDP8-RENEWABLE", None, THEME_HEX)
    doc.add_page_break()

    # Section 5: 한국 기업 기회
    add_heading(doc, "5. 한국 기업 기회", 1, THEME_COLOR, 13)
    add_hr(doc, THEME_HEX)
    render_korean_opportunity(doc, [
        ("태양광 EPC", "HIGH",
         "한화큐셀·OCI·현대에너지 — 대형 태양광 EPC + 모듈 공급 ($5B+ 시장)"),
        ("해상풍력", "HIGH",
         "SK에코플랜트·삼성물산 — 해상풍력 EPC + 터빈 공급 컨소시엄"),
        ("LNG 터미널", "HIGH",
         "현대건설·삼성물산 — LNG 수입터미널 EPC (터미널 5개, $3B+)"),
        ("원자력 참여", "HIGH",
         "한국수력원자력(KHNP) — APR1400 수출 협상 (닌투언 대안 제안)"),
        ("BESS·스마트그리드", "MEDIUM",
         "삼성SDI·LG에너지솔루션 — Grid-Scale BESS + LS Electric 변전"),
        ("그린수소", "MEDIUM",
         "POSCO·현대차 — 그린수소 생산·수출 파일럿 (싱가포르 루트)"),
        ("ODA 에너지", "HIGH",
         "EDCF(수출입은행) — 재생에너지·송전망 ODA 패키지 ($1B+)"),
    ], THEME_HEX)
    doc.add_page_break()

    build_appendix(doc, all_arts, THEME_HEX)

    now2 = datetime.now(); week2 = now2.isocalendar()[1]
    fname = f"MI_REPORT_PDP8_INTEGRATED_W{week2:02d}_{now2.strftime('%Y%m%d')}.docx"
    fpath, sz = save_report(doc, output_dir, fname)
    return fpath


if __name__ == "__main__":
    import sys; sys.path.insert(0,"/home/work/claw/scripts")
    out = "/home/work/claw/outputs/reports/MI_Reports_v4"
    Path(out).mkdir(parents=True, exist_ok=True)
    fpath = generate_pdp8_report(out)
    print(f"✅ {fpath}  ({os.path.getsize(fpath)//1024}KB)")

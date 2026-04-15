"""
PDP8 통합 MI 보고서 생성기 — Vietnam Infrastructure Hub
=========================================================
Option A: PDP8 모(母)플랜 단일 보고서
  - 표지: Decision 768/QD-TTg 기반 전체 개요
  - Section 1: PDP8 전체 국가 KPI (2030/2050)
  - Section 2: 재생에너지 (VN-PDP8-RENEWABLE)
  - Section 3: LNG 발전 & 인프라 (VN-PDP8-LNG)
  - Section 4: 원자력 (VN-PDP8-NUCLEAR)
  - Section 5: 석탄 전환·폐지 (VN-PDP8-COAL)
  - Section 6: 송전망·스마트그리드 (VN-PDP8-GRID)
  - Section 7: 수소·그린에너지 수출 (VN-PDP8-HYDROGEN)
  - Section 8: Province 활동도 (분야별 교차 분석)
  - Section 9: 한국 기업 기회 분석
  - Appendix:  전체 기사 목록 (분야 태그 포함)
"""
import os, sys, json, re
from datetime import datetime
from collections import defaultdict, Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR   = Path("/home/work/claw")
CONFIG_DIR = BASE_DIR / "config"
OUTPUT_DIR = Path(os.getenv("PIPELINE_OUTPUT_DIR", "/home/work/claw/outputs"))

# PDP8 색상 팔레트
COLOR_MAIN   = (0xF5, 0x9E, 0x0B)   # amber  — 표지/섹션 헤더
COLOR_RE     = (0x16, 0xA3, 0x4A)   # green  — 재생에너지
COLOR_LNG    = (0xEF, 0x44, 0x44)   # red    — LNG
COLOR_NUC    = (0x63, 0x66, 0xF1)   # indigo — 원자력
COLOR_COAL   = (0x78, 0x71, 0x6C)   # stone  — 석탄
COLOR_GRID   = (0x0E, 0xA5, 0xE9)   # sky    — 송전망
COLOR_H2     = (0x10, 0xB9, 0x81)   # emerald— 수소
COLOR_PROV   = (0x7C, 0x3A, 0xED)   # violet — Province
COLOR_KOR    = (0xDC, 0x26, 0x26)   # red    — 한국 기회

HEX = {
    "main":  "F59E0B", "re":    "16A34A", "lng":   "EF4444",
    "nuc":   "6366F1", "coal":  "78716C", "grid":  "0EA5E9",
    "h2":    "10B981", "prov":  "7C3AED", "kor":   "DC2626",
    "light": "FEF3C7",  # amber-50
}

# 6개 하부 트래킹 정의
SUB_TRACKS = [
    {
        "id":      "VN-PDP8-RENEWABLE",
        "label":   "재생에너지 (태양광·풍력·수력)",
        "emoji":   "☀️",
        "color":   COLOR_RE,
        "hex":     HEX["re"],
        "kpi_2030": "태양광 46,459–73,416MW | 육상풍력 26,066–38,029MW | 해상풍력 6,000–17,032MW | RE비중 28–36%",
        "kpi_2050": "태양광 293,088–295,646MW | 해상풍력 113,503–139,097MW | RE비중 74–75%",
        "key_projects": [
            "Binh Thuan / Ninh Thuan 해상풍력 클러스터 (남부)",
            "Quang Binh / Ha Tinh / Nghe An 해상풍력 (북중부)",
            "Ba Ria Vung Tau 해상풍력",
            "메콩델타 수상태양광",
            "중부 고원 양수발전 (Bac Ai, Ninh Son)",
            "Trung Nam 태양광 + 배터리 (Ninh Thuan)",
            "RE 수출 허브 → 싱가포르/말레이시아 (2035)",
        ],
    },
    {
        "id":      "VN-PDP8-LNG",
        "label":   "LNG 발전 & 인프라",
        "emoji":   "🔥",
        "color":   COLOR_LNG,
        "hex":     HEX["lng"],
        "kpi_2030": "LNG 발전 22,524MW | 국내가스 10,861–14,930MW | LNG 비중 9.5–12.3%",
        "kpi_2050": "LNG+CCS 1,887–2,269MW | LNG+수소혼소 18,200–26,123MW | 전환완료 8,576–11,325MW",
        "key_projects": [
            "Son My LNG 터미널 + 발전소 (Binh Thuan) — 3,600MW",
            "Ca Mau LNG 터미널 + 발전소 — 1,500MW",
            "Nhon Trach 3·4호기 LNG (Dong Nai) — 1,500MW",
            "Quang Ninh LNG 발전 — 1,500MW",
            "Hai Lang LNG (Quang Tri) — 1,500MW",
            "Long Son LNG (Ba Ria Vung Tau) — 3,000MW",
            "FSRU 터미널: Son My, Ca Mau, Quang Ninh, Thi Vai",
        ],
    },
    {
        "id":      "VN-PDP8-NUCLEAR",
        "label":   "원자력 (닌투언 1·2호기)",
        "emoji":   "⚛️",
        "color":   COLOR_NUC,
        "hex":     HEX["nuc"],
        "kpi_2030": "원자력 4,000–6,400MW (2030–2035 상업운전)",
        "kpi_2050": "추가 8,000MW | 총 원자력 10,500–14,000MW",
        "key_projects": [
            "닌투언 1호기 NPP (공급자 협상 중 — Russia Rosatom 유력)",
            "닌투언 2호기 NPP (공급자 협상 중 — Japan JAEA/JAPC 유력)",
            "Resolution 174/2024/QH15 (2024.11.30) — 원자력 재개 결의",
            "Resolution 189/2025/QH15 (2025.02.19) — 닌투언 건설 특별 메커니즘",
        ],
    },
    {
        "id":      "VN-PDP8-COAL",
        "label":   "석탄 전환·폐지 로드맵",
        "emoji":   "🏭",
        "color":   COLOR_COAL,
        "hex":     HEX["coal"],
        "kpi_2030": "석탄 31,055MW (신규 없음) | 비중 13.1–16.9%",
        "kpi_2050": "석탄 0MW — 완전 퇴출 | 바이오매스/암모니아 혼소 25,632–32,432MW",
        "key_projects": [
            "Vung Ang 1 (1,200MW, Ha Tinh) — 20년 이상 → 바이오매스 전환 검토",
            "Vinh Tan 4 (1,200MW, Binh Thuan)",
            "Duyen Hai 1 (1,245MW, Tra Vinh)",
            "Quang Ninh 클러스터 (2,400MW+)",
            "Nghi Son 2 (1,200MW, Thanh Hoa)",
            "JETP 이행 로드맵 — 석탄 피크 2030, 40년 초과 설비 순차 폐쇄",
        ],
    },
    {
        "id":      "VN-PDP8-GRID",
        "label":   "송전망·스마트그리드",
        "emoji":   "⚡",
        "color":   COLOR_GRID,
        "hex":     HEX["grid"],
        "kpi_2030": "500kV 신설 12,944km | 220kV 신설 15,307km | 변전소 신설 102,900MVA | 투자 $18.1B",
        "kpi_2050": "HVDC 26,000–36,000MW / 3,500–6,600km | 초고압교류 24,000MVA",
        "key_projects": [
            "북남 500kV 3회선 확충",
            "HVDC 북남 백본 (2031–2035, 재생에너지 송전용)",
            "EVN 스마트그리드 고도화 (변전소 자동화)",
            "중부 베트남 RE 계통 연계",
            "라오스·캄보디아·중국 국경간 연계선",
            "직접전력구매계약 (DPPA) 시장 인프라",
        ],
    },
    {
        "id":      "VN-PDP8-HYDROGEN",
        "label":   "수소·그린에너지 신산업",
        "emoji":   "🌿",
        "color":   COLOR_H2,
        "hex":     HEX["h2"],
        "kpi_2030": "광역 RE 산업서비스 센터 2개 설립",
        "kpi_2050": "수소발전 7,030MW | 수소혼소 LNG 16,400–20,900MW | RE수출 10,000MW 유지",
        "key_projects": [
            "북부 RE 센터 (해상풍력 + 수소 생산)",
            "남부 RE 센터 (해상풍력 + 태양광 + 수출 케이블)",
            "싱가포르 수출 케이블 (2035, 5,000–10,000MW)",
            "말레이시아 RE 수출 협정",
            "수소혼소 LNG 전환 로드맵 (기존 LNG 플랜트 순차 전환)",
            "그린 암모니아 생산 기지 (해안 지역)",
        ],
    },
]


# ── 유틸 ──────────────────────────────────────────────────────
def _rgb(t): return RGBColor(*t)
def _hex_bg(cell, h):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),h)
    p.append(s)
def _hr(doc, color="CCCCCC"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'), color)
    pb.append(bot); pPr.append(pb)
def _h(doc, text, lvl, color=COLOR_MAIN, size=None):
    h = doc.add_heading(text, lvl)
    for r in h.runs:
        r.font.color.rgb = _rgb(color)
        if size: r.font.size = Pt(size)
    return h
def _date_key(a):
    d = str(a.get('published_date','') or '')
    try:
        if re.match(r'\d{4}-\d{2}-\d{2}', d): return d[:10]
        return datetime.strptime(d.strip(), '%b %d, %Y').strftime('%Y-%m-%d')
    except: return d[:10] if len(d) >= 10 else '1900-01-01'
def _is_vi(t):
    return any(c in (t or '') for c in ['ă','ơ','ư','đ','ấ','ề','ộ','ừ','ạ','ọ','ổ','ị','ế'])
def _title(a):
    t = a.get('title','')
    if _is_vi(t):
        en = a.get('summary_en','')
        if en and not _is_vi(en): return f"[VI] {en[:80]}"
        ko = a.get('summary_ko','')
        if ko and not _is_vi(ko): return f"[VI] {ko[:70]}"
    return t[:80]
def _summary(a):
    ko = a.get('summary_ko','')
    en = a.get('summary_en','')
    if ko and not _is_vi(ko): return '🇰🇷', ko[:200]
    if en and not _is_vi(en): return '🇺🇸', en[:200]
    return None, None


def _select_articles(arts):
    """2026년 기사 우선, 없으면 2025년, 없으면 전체 최신"""
    s = sorted(arts, key=_date_key, reverse=True)
    y2026 = [a for a in s if '2026' in str(a.get('published_date',''))]
    if y2026: return y2026, "2026"
    y2025 = [a for a in s if '2025' in str(a.get('published_date',''))]
    if y2025: return y2025, "2025"
    return s[:20], "전체"


def _relevance(a, track_id):
    """트래킹 ID별 필수 키워드 최소 1개"""
    must = {
        "VN-PDP8-RENEWABLE": ["solar","wind","offshore wind","onshore wind","floating solar",
                               "rooftop solar","hydropower","pumped storage","battery storage",
                               "renewable energy","điện gió","điện mặt trời","năng lượng tái tạo"],
        "VN-PDP8-LNG":       ["lng","fsru","gas terminal","gas power","gas-fired","natural gas",
                               "khí lng","cảng lng","điện khí","nhập khẩu lng"],
        "VN-PDP8-NUCLEAR":   ["nuclear","ninh thuan","hạt nhân","điện hạt nhân","smr",
                               "resolution 174","resolution 189"],
        "VN-PDP8-COAL":      ["coal","jetp","coal phase","coal retire","biomass cofire",
                               "ammonia cofire","nhiệt điện than","chuyển đổi than"],
        "VN-PDP8-GRID":      ["500kv","220kv","transmission","substation","smart grid","hvdc",
                               "dppa","electricity market","lưới điện","trạm biến áp","500kv"],
        "VN-PDP8-HYDROGEN":  ["green hydrogen","hydrogen","ammonia fuel","re export",
                               "energy export","hydro xanh","xuất khẩu điện"],
    }
    text = ((a.get('title','') or '') + ' ' +
            (a.get('summary_en','') or '') + ' ' +
            (a.get('summary_ko','') or '') + ' ' +
            (a.get('content','') or '')[:300]).lower()
    return any(kw in text for kw in must.get(track_id, []))


def _add_article_block(doc, arts, color, max_n=5):
    """기사 블록 렌더링"""
    for i, art in enumerate(arts[:max_n], 1):
        t = _title(art)
        flag, sm = _summary(art)
        date = _date_key(art)
        src  = (art.get('source','') or '')[:30]

        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Inches(0.2)
        nr = ap.add_run(f"[{i}] "); nr.bold=True; nr.font.size=Pt(10); nr.font.color.rgb=_rgb(color)
        ap.add_run(t).bold = True; ap.runs[-1].font.size = Pt(10)

        mp = doc.add_paragraph()
        mp.paragraph_format.left_indent = Inches(0.4)
        mp.add_run(f"📅 {date}  |  📰 {src}").font.size = Pt(8.5)

        if flag and sm:
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.4)
            sp.add_run(f"{flag} ").font.size = Pt(9.5)
            sp.add_run(sm).font.size = Pt(9.5)
            # 영문도 함께
            en = art.get('summary_en','')
            if flag == '🇰🇷' and en and not _is_vi(en):
                ep = doc.add_paragraph()
                ep.paragraph_format.left_indent = Inches(0.4)
                ep.add_run("🇺🇸 ").font.size = Pt(9.5)
                ep.add_run(en[:180]).font.size = Pt(9.5)
        doc.add_paragraph()


def generate_pdp8_report(output_dir: str) -> str:
    # ── 데이터 로드 ──
    with open(BASE_DIR/"config/history_db.json", encoding='utf-8') as f:
        hdb = json.load(f)
    with open(BASE_DIR/"config/pdp8_structure.json", encoding='utf-8') as f:
        pdp8 = json.load(f)

    now  = datetime.now()
    week = now.isocalendar()[1]

    # 트래킹별 기사 버킷
    buckets = defaultdict(list)
    for art in hdb['articles'].values():
        for pid in (art.get('matched_plans') or []):
            if pid.startswith('VN-PDP8-'):
                buckets[pid].append(art)

    # 전체 PDP8 관련 기사 (중복 제거)
    all_pdp8_ids = set()
    for arts in buckets.values():
        for a in arts: all_pdp8_ids.add(id(a))
    total_arts = sum(len(v) for v in buckets.values())

    # ── Document 초기화 ──
    doc = Document()
    for sec in doc.sections:
        sec.top_margin=Inches(1.0); sec.bottom_margin=Inches(1.0)
        sec.left_margin=Inches(1.2); sec.right_margin=Inches(1.2)

    # ══ COVER ══
    cov = doc.add_table(rows=1, cols=1); cov.style="Table Grid"
    c = cov.rows[0].cells[0]; _hex_bg(c, "0F172A")
    p = c.paragraphs[0]
    r = p.add_run("VIETNAM INFRASTRUCTURE INTELLIGENCE HUB")
    r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    tp = doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2 = tp.add_run("베트남 제8차 전력개발계획 (PDP8)")
    r2.bold=True; r2.font.size=Pt(20); r2.font.color.rgb=_rgb(COLOR_MAIN)

    sp = doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run("Vietnam National Power Development Plan — Decision 768/QD-TTg | April 15, 2025 (Revised)").font.size=Pt(11)

    doc.add_paragraph()
    # 커버 인포 테이블
    it = doc.add_table(rows=6, cols=2); it.style="Table Grid"
    for i,(k,v) in enumerate([
        ("결정문",    "Decision 768/QD-TTg (2025.04.15) — Revised PDP8 | 원본: Decision 500/QD-TTg (2023.05.15)"),
        ("서명자",    "Deputy Prime Minister Bui Thanh Son"),
        ("이행계획",  "Decision 1509/QD-BCT (2025.05.30) — MOIT Implementation Plan"),
        ("3대 기둥",  "① 에너지 안보  ② 공정전환 (JETP)  ③ RE 산업생태계 육성"),
        ("투자 규모", "2026–2030: $136.3B  |  2031–2035: $130.0B  |  2036–2050: $569.1B"),
        ("보고서 기사",f"PDP8 관련 역사 기사: {total_arts}건 | 6개 하부 트래킹"),
    ]):
        it.rows[i].cells[0].text=k; it.rows[i].cells[1].text=v
        _hex_bg(it.rows[i].cells[0], "F59E0B")
        it.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        it.rows[i].cells[0].paragraphs[0].runs[0].bold=True
        for cell in it.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(10)
    doc.add_paragraph()
    doc.add_paragraph().alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].add_run(
        f"Generated: {now.strftime('%Y-%m-%d %H:%M')}  |  W{week}/{now.year}  |  CONFIDENTIAL"
    ).font.size=Pt(9)
    doc.add_page_break()

    # ══ SECTION 1: PDP8 전체 국가 KPI ══
    _h(doc, "1. PDP8 국가 KPI — 2030 / 2050", 1, COLOR_MAIN, 14); _hr(doc, HEX["main"])

    doc.add_paragraph()
    kpi_t = doc.add_table(rows=1, cols=3); kpi_t.style="Table Grid"
    for j,h in enumerate(["구분","2030 목표","2050 목표"]):
        kpi_t.rows[0].cells[j].text=h
        kpi_t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _hex_bg(kpi_t.rows[0].cells[j], "0F172A")
        kpi_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        kpi_t.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)

    kpi30 = pdp8["national_kpi_2030"]
    kpi50 = pdp8["national_kpi_2050"]
    cap30 = pdp8["capacity_by_source_2030_mw"]
    invest= pdp8["investment_requirements_usd_billion"]

    rows_data = [
        ("상업전력 (억 kWh)",       kpi30["commercial_electricity_billion_kwh"],   kpi50["commercial_electricity_billion_kwh"]),
        ("최대 수요 (MW)",          kpi30["maximum_capacity_mw"],                  kpi50["maximum_capacity_mw"]),
        ("총 설비 용량 (MW)",       kpi30["total_installed_capacity_mw"],          kpi50["total_installed_capacity_mw"]),
        ("☀️ 태양광 (MW)",          f"{cap30['solar_power']['low']:,}–{cap30['solar_power']['high']:,}  ({cap30['solar_power']['pct']})",  "293,088–295,646 MW (35–38%)"),
        ("🌬️ 육상풍력 (MW)",        f"{cap30['onshore_nearshore_wind']['low']:,}–{cap30['onshore_nearshore_wind']['high']:,}  ({cap30['onshore_nearshore_wind']['pct']})", "84,696–91,400 MW"),
        ("🌊 해상풍력 (MW)",        f"{cap30['offshore_wind']['low']:,}–{cap30['offshore_wind']['high']:,}  ({cap30['offshore_wind']['pct']})", "113,503–139,097 MW"),
        ("💧 수력 (MW)",            f"{cap30['hydropower']['low']:,}–{cap30['hydropower']['high']:,}  ({cap30['hydropower']['pct']})", "40,624 MW"),
        ("🔋 배터리 저장 (MW)",     f"{cap30['battery_storage']['low']:,}–{cap30['battery_storage']['high']:,}  ({cap30['battery_storage']['pct']})", "95,983–96,120 MW"),
        ("⚛️ 원자력 (MW)",          f"{cap30['nuclear_power']['low']:,}–{cap30['nuclear_power']['high']:,}  ({cap30['nuclear_power']['pct']})", "10,500–14,000 MW"),
        ("🔥 LNG 발전 (MW)",        f"{cap30['lng_thermal']['fixed']:,}  ({cap30['lng_thermal']['pct']})", "CCS/수소혼소 전환"),
        ("🏭 석탄 (MW)",            f"{cap30['coal_thermal']['fixed']:,}  ({cap30['coal_thermal']['pct']})  ※ 신규 없음", "0 MW — 완전 퇴출"),
        ("🌿 RE 비중 (수력 제외)", kpi30["renewable_share_excl_hydro_pct"],       kpi50["renewable_share_excl_hydro_pct"]),
        ("💨 온실가스 (백만 톤)",   kpi30["ghg_emissions_million_tons"],           kpi50["ghg_emissions_million_tons"]),
        ("📤 RE 수출 (MW)",         "—",                                            "5,000–10,000 MW (2035부터, 싱가포르·말레이시아)"),
        ("💰 투자 (발전)",          f"${invest['2026_2030']['generation']}B (2026–30)", f"${invest['2036_2050']['generation']}B (2036–50)"),
        ("💰 투자 (송전)",          f"${invest['2026_2030']['transmission']}B (2026–30)", f"${invest['2036_2050']['transmission']}B (2036–50)"),
    ]
    for r_data in rows_data:
        row = kpi_t.add_row()
        for j, val in enumerate(r_data):
            row.cells[j].text = str(val)
            if j == 0: _hex_bg(row.cells[j], HEX["light"])
            for para in row.cells[j].paragraphs:
                for run in para.runs: run.font.size = Pt(9)
    doc.add_paragraph()

    # 3대 기둥
    _h(doc, "1-1. 3대 전략 기둥", 2, COLOR_MAIN)
    for item in pdp8["three_pillars"]:
        p = doc.add_paragraph(f"▪ {item}", style="List Bullet")
        p.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    doc.add_page_break()

    # ══ SECTIONS 2–7: 하부 트래킹별 ══
    for sec_no, track in enumerate(SUB_TRACKS, 2):
        tid   = track["id"]
        color = track["color"]
        hexc  = track["hex"]
        arts  = buckets.get(tid, [])
        rel   = [a for a in arts if _relevance(a, tid)]
        rep, yr = _select_articles(rel) if rel else _select_articles(arts)

        # 섹션 헤더
        _h(doc, f"{sec_no}. {track['emoji']}  {track['label']}", 1, color, 13)
        _hr(doc, hexc)

        # 통계 배지
        stat_t = doc.add_table(rows=1, cols=3); stat_t.style="Table Grid"
        for j,(k,v) in enumerate([
            ("역사 기사", f"{len(arts)}건"),
            ("관련성 통과", f"{len(rel)}건"),
            (f"보고 대상({yr})", f"{len(rep)}건"),
        ]):
            stat_t.rows[0].cells[j].text=f"{k}\n{v}"
            _hex_bg(stat_t.rows[0].cells[j], hexc)
            for para in stat_t.rows[0].cells[j].paragraphs:
                for run in para.runs:
                    run.font.size=Pt(10); run.bold=True
                    run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        doc.add_paragraph()

        # KPI 테이블
        kpi_tbl = doc.add_table(rows=2, cols=2); kpi_tbl.style="Table Grid"
        kpi_tbl.rows[0].cells[0].text="2030 목표"
        kpi_tbl.rows[0].cells[1].text=track["kpi_2030"]
        kpi_tbl.rows[1].cells[0].text="2050 목표"
        kpi_tbl.rows[1].cells[1].text=track["kpi_2050"]
        for ri in range(2):
            _hex_bg(kpi_tbl.rows[ri].cells[0], hexc)
            kpi_tbl.rows[ri].cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            kpi_tbl.rows[ri].cells[0].paragraphs[0].runs[0].bold=True
            for cell in kpi_tbl.rows[ri].cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.font.size=Pt(9.5)
        doc.add_paragraph()

        # 주요 프로젝트
        _h(doc, "주요 프로젝트·정책", 2, color)
        for proj in track["key_projects"]:
            p = doc.add_paragraph(f"▪ {proj}", style="List Bullet")
            p.runs[0].font.size = Pt(10)
        doc.add_paragraph()

        # 최신 뉴스
        if rep:
            _h(doc, f"최신 뉴스 ({yr}, 상위 5건)", 2, color)
            _add_article_block(doc, rep, color, max_n=5)
            if len(rep) > 5:
                more_p = doc.add_paragraph()
                more_p.paragraph_format.left_indent = Inches(0.2)
                more_p.add_run(f"  ※ 추가 {len(rep)-5}건 → Appendix 참조").font.size=Pt(9)
        else:
            np2 = doc.add_paragraph()
            np2.paragraph_format.left_indent = Inches(0.2)
            np2.add_run(f"ℹ️ {yr}년 관련 기사가 아직 수집되지 않았습니다.").font.size=Pt(9.5)

        _hr(doc, "D1D5DB"); doc.add_paragraph()
        doc.add_page_break()

    # ══ SECTION 8: Province 활동도 (분야별 교차 분석) ══
    _h(doc, "8. Province 활동도 — PDP8 분야별 교차 분석", 1, COLOR_PROV, 13); _hr(doc, HEX["prov"])

    prov_note = doc.add_paragraph()
    prov_note.paragraph_format.left_indent = Inches(0.2)
    prov_note.add_run(
        "각 Province의 PDP8 관련 기사를 분야별로 집계합니다. "
        "기사 수는 해당 Province의 에너지 인프라 개발 활동도를 반영합니다."
    ).font.size = Pt(10)
    doc.add_paragraph()

    # Province × Track 매트릭스 집계
    prov_track = defaultdict(lambda: defaultdict(int))
    prov_norm = {
        'vietnam':'전국', 'national / unspecified':'전국',
        'ho chi minh city':'Ho Chi Minh City', 'hcmc':'Ho Chi Minh City',
        'ha noi':'Hanoi', 'hanoi':'Hanoi',
    }
    track_ids = [t["id"] for t in SUB_TRACKS]
    track_short = {t["id"]: t["label"][:12] for t in SUB_TRACKS}

    for track in SUB_TRACKS:
        for art in buckets.get(track["id"], []):
            pv = (art.get('province','') or '전국').strip()
            pv = prov_norm.get(pv.lower(), pv)
            if not pv: pv = '전국'
            prov_track[pv][track["id"]] += 1

    # 상위 15개 Province
    top_provs = sorted(prov_track.keys(),
                       key=lambda p: sum(prov_track[p].values()), reverse=True)[:15]

    if top_provs:
        short_labels = [t["emoji"] + t["label"][:8] for t in SUB_TRACKS]
        pt = doc.add_table(rows=1, cols=len(SUB_TRACKS)+2); pt.style="Table Grid"
        # 헤더
        pt.rows[0].cells[0].text="Province"
        pt.rows[0].cells[1].text="합계"
        for j, track in enumerate(SUB_TRACKS):
            pt.rows[0].cells[j+2].text = track["emoji"]
            _hex_bg(pt.rows[0].cells[j+2], track["hex"])
            pt.rows[0].cells[j+2].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        _hex_bg(pt.rows[0].cells[0], HEX["prov"])
        _hex_bg(pt.rows[0].cells[1], HEX["prov"])
        for cell in pt.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold=True; run.font.size=Pt(9)
                    if cell == pt.rows[0].cells[0] or cell == pt.rows[0].cells[1]:
                        run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

        for prov in top_provs:
            row = pt.add_row()
            row.cells[0].text = prov
            total = sum(prov_track[prov].values())
            row.cells[1].text = str(total)
            for j, track in enumerate(SUB_TRACKS):
                cnt = prov_track[prov].get(track["id"], 0)
                row.cells[j+2].text = str(cnt) if cnt else "—"
                if cnt >= 10: _hex_bg(row.cells[j+2], track["hex"])
                elif cnt >= 5: _hex_bg(row.cells[j+2], "FEF3C7")
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.font.size=Pt(9)

    doc.add_paragraph()
    # 이모지 범례
    leg_p = doc.add_paragraph()
    leg_p.paragraph_format.left_indent = Inches(0.2)
    legend_text = "  |  ".join([f"{t['emoji']} {t['label'][:15]}" for t in SUB_TRACKS])
    leg_p.add_run(f"범례: {legend_text}").font.size=Pt(8.5)
    doc.add_paragraph()
    doc.add_page_break()

    # ══ SECTION 9: 한국 기업 기회 분석 ══
    _h(doc, "9. 한국 기업 기회 분석", 1, COLOR_KOR, 13); _hr(doc, HEX["kor"])

    kor_opps = [
        ("☀️ 재생에너지",
         "HIGH: 해상풍력 EPC — 삼성물산·현대건설·SK에코 등 대형 플랜트 경험 보유\n"
         "HIGH: 태양광 모듈 공급 — 한화큐셀 현지 생산 또는 수출\n"
         "MEDIUM: 양수발전 설계·시공 — 중부고원 Bac Ai, Ninh Son 프로젝트\n"
         "MEDIUM: 배터리 저장 시스템 — 삼성SDI·LG에너지솔루션 BESS 공급\n"
         "TIP: Decision 768 해상풍력 인센티브 (세금면제, 최소 전력구매 보장) 적극 활용"),
        ("🔥 LNG 인프라",
         "HIGH: LNG FSRU/터미널 EPC — KOGAS·현대중공업·삼성중공업\n"
         "HIGH: 가스복합화력 발전소 — 두산에너빌리티(구 두산중공업) 가스터빈\n"
         "MEDIUM: LNG 장기공급 계약 — KOGAS 국제 트레이딩\n"
         "핵심 프로젝트: Son My 3,600MW, Ca Mau 1,500MW, Nhon Trach 3·4"),
        ("⚛️ 원자력",
         "HIGH: 닌투언 1·2호기 — 한국 공급자 참여 가능성 (Resolution 189 국제 입찰)\n"
         "HIGH: 두산에너빌리티 원자로 압력용기·증기발생기 공급\n"
         "MEDIUM: 원전 O&M, 핵연료 공급, 방사선 관리 서비스\n"
         "NOTE: Resolution 189/2025 특별 메커니즘 — 외국 기업 참여 명시"),
        ("⚡ 송전망·그리드",
         "HIGH: 500kV 송전선 EPC — LS Electric·현대일렉트릭 변압기/GIS\n"
         "HIGH: 스마트그리드 시스템 — 한전KDN·LS Electric AMI/EMS\n"
         "MEDIUM: HVDC 장비 공급 (2031–2035 북남 HVDC)\n"
         "MEDIUM: DPPA 전력시장 IT 인프라 — 전력거래소 시스템 수출"),
        ("🌿 수소·RE 수출",
         "HIGH: 싱가포르 RE 수출 케이블 컨소시엄 참여 — 포스코·한전\n"
         "HIGH: 그린 암모니아 생산 플랜트 — 롯데케미칼·현대엔지니어링\n"
         "MEDIUM: 수소혼소 LNG 기술 전환 — 두산에너빌리티\n"
         "LONG-TERM: 2035년 5,000–10,000MW 수출 시장 선점"),
        ("🏭 석탄 전환",
         "MEDIUM: 바이오매스 전환 컨설팅 및 공사 — 한국남동발전 등 운영 경험\n"
         "MEDIUM: 암모니아 혼소 기술 공급\n"
         "LOW: 폐쇄 부지 환경 복원 서비스\n"
         "TIP: JETP 파트너십 자금 활용 가능 ($15.5B JETP 약정)"),
    ]

    kor_t = doc.add_table(rows=1, cols=2); kor_t.style="Table Grid"
    for j,h in enumerate(["분야","한국 기업 기회"]):
        kor_t.rows[0].cells[j].text=h
        kor_t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _hex_bg(kor_t.rows[0].cells[j], HEX["kor"])
        kor_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for sector, opp in kor_opps:
        row = kor_t.add_row()
        row.cells[0].text = sector
        row.cells[1].text = opp
        _hex_bg(row.cells[0], "FEF2F2")
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(9.5)
    doc.add_paragraph()

    # 전략적 시사점
    _h(doc, "9-1. 전략적 시사점", 2, COLOR_KOR)
    strat = doc.add_paragraph()
    strat.paragraph_format.left_indent = Inches(0.2)
    strat.add_run(
        "① PDP8 $136.3B 투자 중 한국 기업 접근 가능 시장: 약 $20–30B 추정\n"
        "② DPPA 도입 → 민간 RE 개발자로 직접 진입 가능 (한국 IPP 경험 활용)\n"
        "③ 원자력 재개 — Resolution 189 특별 메커니즘 국제 입찰 조기 대응 필수\n"
        "④ 해상풍력 인허가 창 좁음 (2031년 이전 원칙결정 또는 승인 필요) → 즉시 로비\n"
        "⑤ JETP $15.5B 자금 → ADB·KDB 공동 파이낸싱 구조 활용\n"
        "⑥ 핵심 창구: MOIT (ERAV) + EVN — 한국전력공사 MOU 확대 활용"
    ).font.size = Pt(10)
    doc.add_paragraph()
    doc.add_page_break()

    # ══ APPENDIX: 전체 기사 목록 ══
    _h(doc, "Appendix — PDP8 전체 관련 기사 목록", 1, COLOR_MAIN)
    _hr(doc, HEX["main"])

    # 전체 기사 수집 (중복 제거, 최신순)
    seen_urls = set()
    all_arts_flat = []
    for track in SUB_TRACKS:
        for art in buckets.get(track["id"], []):
            url = art.get('url','') or str(id(art))
            if url not in seen_urls:
                seen_urls.add(url)
                art['_track'] = track["emoji"] + track["label"][:10]
                all_arts_flat.append(art)
    all_arts_flat.sort(key=_date_key, reverse=True)

    ap_note = doc.add_paragraph(
        f"총 {len(all_arts_flat)}건 (6개 트래킹 중복 제거) — 상위 80건 표시\n"
        "전체 이력: Database Excel 파일의 각 플랜 시트 참조"
    )
    ap_note.runs[0].font.size=Pt(9); ap_note.runs[0].font.color.rgb=RGBColor(0x64,0x74,0x8B)
    doc.add_paragraph()

    ap_t = doc.add_table(rows=1, cols=5); ap_t.style="Table Grid"
    for j,h in enumerate(["No","날짜","제목 (EN/KO)","출처","분야"]):
        ap_t.rows[0].cells[j].text=h
        ap_t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
        _hex_bg(ap_t.rows[0].cells[j], "1E40AF")
        ap_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        ap_t.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)

    for ri, art in enumerate(all_arts_flat[:80], 1):
        row = ap_t.add_row()
        row.cells[0].text = str(ri)
        row.cells[1].text = _date_key(art)
        row.cells[2].text = _title(art)[:60]
        row.cells[3].text = (art.get('source','') or '')[:18]
        row.cells[4].text = art.get('_track', '')[:15]
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(8.5)

    # ── 저장 ──
    fname = f"MI_REPORT_PDP8_INTEGRATED_W{week:02d}_{now.strftime('%Y%m%d')}.docx"
    fpath = os.path.join(output_dir, fname)
    doc.save(fpath)
    return fpath


if __name__ == "__main__":
    out = "/home/work/claw/outputs/reports/MI_Reports"
    Path(out).mkdir(parents=True, exist_ok=True)
    print("PDP8 통합 MI 보고서 생성 중...")
    fpath = generate_pdp8_report(out)
    if fpath:
        sz = os.path.getsize(fpath) // 1024
        print(f"✅ {fpath} ({sz}KB)")

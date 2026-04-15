"""
report_lib.py — 보고서 공유 라이브러리 v4.0
=============================================
모든 MI 보고서(지역성/KPI형 공통)에서 import하여 사용하는 빌딩 블록
"""
import os, re, json, ast
from pathlib import Path
from datetime import datetime
from collections import defaultdict, Counter

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path("/home/work/claw")

# ─────────────────────────────────────────────
# 색상 팔레트
# ─────────────────────────────────────────────
PALETTE = {
    "navy":    (0x0F, 0x17, 0x2A),
    "blue":    (0x1E, 0x40, 0xAF),
    "sky":     (0x0E, 0xA5, 0xE9),
    "red":     (0xDC, 0x26, 0x26),
    "orange":  (0xEA, 0x58, 0x0C),
    "amber":   (0xD9, 0x77, 0x06),
    "green":   (0x15, 0x80, 0x3D),
    "teal":    (0x0D, 0x94, 0x88),
    "violet":  (0x7C, 0x3A, 0xED),
    "gray":    (0x4B, 0x55, 0x63),
    "light":   (0xF8, 0xFA, 0xFC),
    "white":   (0xFF, 0xFF, 0xFF),
    "gold":    (0xB4, 0x5A, 0x09),
}
PALETTE_HEX = {k: "%02X%02X%02X" % v for k, v in PALETTE.items()}

# 보고서 유형별 테마 색상
THEME = {
    "REGIONAL":   ("red",    "EF4444", "FEF2F2"),   # 지역성 — 빨강
    "KPI":        ("blue",   "1E40AF", "EFF6FF"),   # KPI형  — 파랑
    "ENERGY":     ("amber",  "D97706", "FFFBEB"),   # 에너지 — 앰버
    "WATER":      ("sky",    "0EA5E9", "F0F9FF"),   # 수자원 — 스카이
    "HANOI":      ("red",    "DC2626", "FEF2F2"),   # 하노이 — 빨강
    "ENV":        ("green",  "15803D", "F0FDF4"),   # 환경   — 그린
    "TRANSPORT":  ("orange", "EA580C", "FFF7ED"),   # 교통   — 오렌지
}

# ─────────────────────────────────────────────
# 문서 초기화
# ─────────────────────────────────────────────
def new_doc(margins=(0.9, 0.9, 1.1, 1.1)) -> Document:
    """여백 설정된 새 문서 반환 (top, bottom, left, right in inches)"""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(margins[0])
        sec.bottom_margin = Inches(margins[1])
        sec.left_margin   = Inches(margins[2])
        sec.right_margin  = Inches(margins[3])
    return doc

# ─────────────────────────────────────────────
# XML 헬퍼
# ─────────────────────────────────────────────
def _rgb(t): return RGBColor(t[0], t[1], t[2])
def _rgb_hex(h): return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def cell_bg(cell, hex_color: str):
    """셀 배경색 설정"""
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), hex_color.upper())
    p.append(s)

def cell_width(cell, width_inches: float):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(width_inches * 1440)))
    w.set(qn('w:type'), 'dxa'); p.append(w)

def add_hr(doc, color_hex="CCCCCC", thickness=6):
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), str(thickness))
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), color_hex.upper())
    pb.append(bot); pPr.append(pb)

def no_space_para(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p

# ─────────────────────────────────────────────
# 타이포그래피 헬퍼
# ─────────────────────────────────────────────
def add_heading(doc, text, level, color_key="navy", size=None, bold=True):
    h = doc.add_heading(text, level)
    for r in h.runs:
        r.font.color.rgb = _rgb(PALETTE[color_key])
        r.bold = bold
        if size: r.font.size = Pt(size)
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after  = Pt(2)
    return h

def add_label_value(doc, label, value, label_color="navy", indent=0.15, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    lr = p.add_run(f"{label}: "); lr.bold = True; lr.font.size = Pt(size)
    lr.font.color.rgb = _rgb(PALETTE[label_color])
    vr = p.add_run(str(value)); vr.font.size = Pt(size)
    return p

def add_note(doc, text, color_key="gray", indent=0.15, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); r.font.size = Pt(size)
    r.font.color.rgb = _rgb(PALETTE[color_key])
    return p

def add_bullet(doc, text, level=1, size=10, color_key=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.2 * level)
    r = p.add_run(text); r.font.size = Pt(size)
    if color_key: r.font.color.rgb = _rgb(PALETTE[color_key])
    return p

# ─────────────────────────────────────────────
# 테이블 빌더
# ─────────────────────────────────────────────
def styled_table(doc, headers, rows, col_widths=None,
                 header_bg="navy", header_fg="white",
                 alt_bg="F8FAFC", border_color="D1D5DB",
                 font_size=9.5):
    """
    헤더 + 데이터 행 테이블.
    rows: list of list (strings)
    col_widths: list of float (inches), optional
    """
    n = len(headers)
    t = doc.add_table(rows=1, cols=n)
    t.style = "Table Grid"

    # 헤더 행
    hrow = t.rows[0]
    for j, h in enumerate(headers):
        cell = hrow.cells[j]
        cell_bg(cell, PALETTE_HEX[header_bg])
        if col_widths: cell_width(cell, col_widths[j])
        p = cell.paragraphs[0]
        r = p.add_run(h); r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = _rgb(PALETTE[header_fg])

    # 데이터 행
    for ri, row_data in enumerate(rows):
        drow = t.add_row()
        bg = alt_bg if ri % 2 == 1 else "FFFFFF"
        for j, val in enumerate(row_data[:n]):
            cell = drow.cells[j]
            if bg != "FFFFFF": cell_bg(cell, bg)
            if col_widths: cell_width(cell, col_widths[j])
            p = cell.paragraphs[0]
            r = p.add_run(str(val) if val is not None else "—")
            r.font.size = Pt(font_size)

    doc.add_paragraph()
    return t

def kpi_box(doc, items, bg_hex="EFF6FF", border_hex="1E40AF", title=None, font_size=10):
    """
    KPI 요약 박스 (1열 테이블로 구현)
    items: list of (label, value, unit)
    """
    cols = min(4, len(items))
    t = doc.add_table(rows=2 if title else 1, cols=cols)
    t.style = "Table Grid"

    row_offset = 0
    if title:
        # 타이틀 행 (병합)
        for j in range(cols):
            cell_bg(t.rows[0].cells[j], border_hex)
        title_cell = t.rows[0].cells[0]
        # 병합
        for j in range(1, cols):
            title_cell = title_cell.merge(t.rows[0].cells[j])
        p = title_cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title); r.bold = True; r.font.size = Pt(font_size)
        r.font.color.rgb = _rgb(PALETTE["white"])
        row_offset = 1
        # 새 데이터행 추가
        data_row = t.add_row()
    else:
        data_row = t.rows[0]

    for j, (label, value, unit) in enumerate(items[:cols]):
        cell = data_row.cells[j]
        cell_bg(cell, bg_hex)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp = cell.paragraphs[0]; lr = lp.add_run(label + "\n")
        lr.bold = True; lr.font.size = Pt(font_size - 1)
        lr.font.color.rgb = _rgb(PALETTE["gray"])
        vr = lp.add_run(str(value)); vr.bold = True; vr.font.size = Pt(font_size + 2)
        vr.font.color.rgb = _rgb(PALETTE["navy"])
        if unit:
            ur = lp.add_run(f"\n{unit}"); ur.font.size = Pt(font_size - 2)
            ur.font.color.rgb = _rgb(PALETTE["gray"])

    doc.add_paragraph()
    return t

def info_box(doc, kv_list, bg_hex="F8FAFC", label_bg="1E40AF", font_size=9.5):
    """
    키-값 정보 박스 (2열 테이블)
    kv_list: list of (key, value)
    """
    t = doc.add_table(rows=len(kv_list), cols=2)
    t.style = "Table Grid"
    for ri, (k, v) in enumerate(kv_list):
        cell_bg(t.rows[ri].cells[0], label_bg)
        cell_bg(t.rows[ri].cells[1], bg_hex)
        cell_width(t.rows[ri].cells[0], 1.4)
        cell_width(t.rows[ri].cells[1], 4.5)
        kr = t.rows[ri].cells[0].paragraphs[0].add_run(k)
        kr.bold = True; kr.font.size = Pt(font_size)
        kr.font.color.rgb = _rgb(PALETTE["white"])
        vr = t.rows[ri].cells[1].paragraphs[0].add_run(str(v))
        vr.font.size = Pt(font_size)
    doc.add_paragraph()
    return t

def project_card(doc, proj, theme_hex, theme_color_key="navy", idx=None):
    """
    개별 프로젝트 카드 (지역성 보고서용)
    proj: dict with name, investment, area_ha, status, developer, note, deadline 등
    """
    # 카드 헤더
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    hcell = t.rows[0].cells[0]; cell_bg(hcell, theme_hex)
    p = hcell.paragraphs[0]
    prefix = f"[{idx}] " if idx else ""
    r = p.add_run(prefix + proj.get("name", "(이름 없음)"))
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = _rgb(PALETTE["white"])

    # 상세 정보 행
    fields = []
    if proj.get("investment") or proj.get("investment_usd"):
        inv = proj.get("investment") or proj.get("investment_usd","")
        if inv and not str(inv).startswith("$"): inv = f"${inv}"
        fields.append(("💰 투자규모", inv))
    if proj.get("area_ha"):   fields.append(("📐 면적", f"{proj['area_ha']} ha"))
    if proj.get("length_km"): fields.append(("📏 연장", f"{proj['length_km']} km"))
    if proj.get("status"):    fields.append(("📊 현황", proj["status"]))
    if proj.get("developer"): fields.append(("🏢 시행사", proj["developer"]))
    if proj.get("deadline"):  fields.append(("📅 완료목표", proj["deadline"]))
    if proj.get("note"):      fields.append(("📌 비고", proj["note"][:80]))

    if fields:
        ft = doc.add_table(rows=len(fields), cols=2)
        ft.style = "Table Grid"
        for ri, (k, v) in enumerate(fields):
            cell_bg(ft.rows[ri].cells[0], "F1F5F9")
            cell_width(ft.rows[ri].cells[0], 1.3)
            cell_width(ft.rows[ri].cells[1], 4.8)
            kr = ft.rows[ri].cells[0].paragraphs[0].add_run(k)
            kr.bold = True; kr.font.size = Pt(9.5)
            kr.font.color.rgb = _rgb(PALETTE["gray"])
            vr = ft.rows[ri].cells[1].paragraphs[0].add_run(str(v))
            vr.font.size = Pt(9.5)

    doc.add_paragraph()

# ─────────────────────────────────────────────
# 표지 빌더
# ─────────────────────────────────────────────
def build_cover(doc, plan_id, plan_name_ko, plan_name_en,
                legal_basis, lead_agency, investment, period,
                article_count, week_str, report_type, theme_hex,
                subtitle=None):
    """통합 표지 생성"""
    now = datetime.now()

    # 상단 브랜드 배너
    bt = doc.add_table(rows=1, cols=1); bt.style = "Table Grid"
    bc = bt.rows[0].cells[0]; cell_bg(bc, PALETTE_HEX["navy"])
    bp = bc.paragraphs[0]; bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run("VIETNAM INFRASTRUCTURE INTELLIGENCE HUB")
    br.bold = True; br.font.size = Pt(10)
    br.font.color.rgb = _rgb((0x93, 0xC5, 0xFD))
    doc.add_paragraph()

    # 보고서 유형 배지
    badge_t = doc.add_table(rows=1, cols=1); badge_t.style = "Table Grid"
    badge_c = badge_t.rows[0].cells[0]; cell_bg(badge_c, theme_hex)
    badge_p = badge_c.paragraphs[0]; badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_r = badge_p.add_run(f"  {report_type}  ")
    badge_r.bold = True; badge_r.font.size = Pt(9)
    badge_r.font.color.rgb = _rgb(PALETTE["white"])
    doc.add_paragraph()

    # 제목
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(plan_name_ko)
    tr.bold = True; tr.font.size = Pt(20)
    tr.font.color.rgb = _rgb_hex(theme_hex)

    ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.add_run(plan_name_en).font.size = Pt(11)

    if subtitle:
        sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(subtitle); sr.font.size = Pt(10)
        sr.font.color.rgb = _rgb(PALETTE["gray"])

    doc.add_paragraph()

    # 메타 정보 테이블
    meta = [
        ("플랜 ID",    plan_id),
        ("법적 근거",  legal_basis),
        ("담당 부처",  lead_agency),
        ("투자 규모",  investment),
        ("계획 기간",  period),
        ("관련 기사",  f"{article_count}건 (역사 DB 기준)"),
        ("생성 일시",  f"{now.strftime('%Y-%m-%d %H:%M')}  |  {week_str}  |  ⚠️ CONFIDENTIAL"),
    ]
    info_box(doc, meta, bg_hex="F8FAFC", label_bg=theme_hex)
    doc.add_page_break()

# ─────────────────────────────────────────────
# 기사 처리 유틸
# ─────────────────────────────────────────────
def _is_vi(text):
    """베트남어 텍스트 감지"""
    vi_chars = set('ăâêôơưđáàảãạắằẳẵặấầẩẫậéèẻẽẹếềểễệíìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ')
    return any(c.lower() in vi_chars for c in (text or ''))

def _date_key(art):
    d = str(art.get('published_date', '') or '')
    try:
        if re.match(r'\d{4}-\d{2}-\d{2}', d): return d[:10]
        return datetime.strptime(d.strip(), '%b %d, %Y').strftime('%Y-%m-%d')
    except:
        return '1900-01-01'

def _best_title(art):
    t = art.get('title', '') or ''
    if not _is_vi(t): return t[:90]
    # 베트남어 제목 → 영어/한국어 요약으로 대체
    en = art.get('summary_en', '') or ''
    ko = art.get('summary_ko', '') or ''
    if en and not _is_vi(en): return f"[VI→EN] {en[:80]}"
    if ko and not _is_vi(ko): return f"[VI→KO] {ko[:70]}"
    return t[:90]

def _best_summary(art):
    """(flag, text) 반환 — 베트남어 절대 미노출"""
    ko = art.get('summary_ko', '') or ''
    en = art.get('summary_en', '') or ''
    if ko and not _is_vi(ko): return ('🇰🇷', ko[:220])
    if en and not _is_vi(en): return ('🇺🇸', en[:220])
    return (None, None)

def _parse_plans(art):
    mp = art.get('matched_plans', []) or []
    if isinstance(mp, str):
        try: mp = ast.literal_eval(mp)
        except: mp = []
    return mp

def select_articles(arts, prefer_year="2026", fallback_months=12):
    """연도 우선 선택 로직"""
    if not arts: return [], "-"
    s = sorted(arts, key=_date_key, reverse=True)
    yr = [a for a in s if prefer_year in str(a.get('published_date',''))]
    if yr: return yr, f"{prefer_year}년"
    cutoff = f"{int(prefer_year)-1:04d}-{12-fallback_months+1:02d}-01" if fallback_months <= 12 else "2025-01-01"
    recent = [a for a in s if _date_key(a) >= cutoff]
    if recent: return recent, "최근 1년"
    return s[:20], "전체"

def get_year_trend(arts, plan_id, years=range(2019, 2027)):
    """연도별 기사 건수"""
    counts = {str(y): 0 for y in years}
    for art in arts:
        d = _date_key(art)
        yr = d[:4]
        if yr in counts: counts[yr] += 1
    return counts

def get_province_stats(arts):
    """성별 기사 건수"""
    PROVINCE_ALIASES = {
        "vietnam": "🇻🇳 전국", "national": "🇻🇳 전국", "viet nam": "🇻🇳 전국",
        "ho chi minh": "Ho Chi Minh City", "hcm": "Ho Chi Minh City", "hồ chí minh": "Ho Chi Minh City",
        "hanoi": "Hanoi", "hà nội": "Hanoi",
        "da nang": "Da Nang", "đà nẵng": "Da Nang",
        "hai phong": "Hai Phong", "hải phòng": "Hai Phong",
        "can tho": "Can Tho", "cần thơ": "Can Tho",
        "binh duong": "Binh Duong", "bình dương": "Binh Duong",
        "dong nai": "Dong Nai", "đồng nai": "Dong Nai",
        "ba ria": "Ba Ria-Vung Tau", "vung tau": "Ba Ria-Vung Tau",
        "quang ninh": "Quang Ninh", "quảng ninh": "Quang Ninh",
        "long an": "Long An", "tien giang": "Tien Giang", "ben tre": "Ben Tre",
    }
    counter = Counter()
    for art in arts:
        prov = (art.get('province') or art.get('location') or 'Unknown').lower().strip()
        for alias, norm in PROVINCE_ALIASES.items():
            if alias in prov: prov = norm.lower(); break
        prov = prov.title() if prov not in [v.lower() for v in PROVINCE_ALIASES.values()] else prov
        counter[prov] += 1
    return counter.most_common(15)

# ─────────────────────────────────────────────
# 기사 렌더러
# ─────────────────────────────────────────────
def render_articles(doc, arts, color_key="navy", max_n=6,
                    show_history=True, year_label="2026년",
                    new_article_ids=None):
    """기사 카드 렌더링. new_article_ids에 포함된 기사는 노란 배경(★NEW)으로 마킹"""
    if not arts:
        add_note(doc, "ℹ️ 해당 기간 관련 기사가 없습니다.", color_key="gray")
        return

    _new_ids = set(new_article_ids or [])
    for i, art in enumerate(arts[:max_n], 1):
        title  = _best_title(art)
        flag, sm = _best_summary(art)
        date   = _date_key(art)
        src    = (art.get('source','') or '')[:35]
        url    = art.get('url','') or ''
        is_new = art.get('id','') in _new_ids or art.get('_is_new', False)

        # 번호 + 제목
        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Inches(0.15)
        ap.paragraph_format.space_before = Pt(2)
        # ★ NEW 마킹 (노란 배경)
        if is_new:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'FFFF99')
            shd.set(qn('w:val'), 'clear')
            ap.paragraph_format.element.get_or_add_pPr().append(shd)
            new_r = ap.add_run("★NEW "); new_r.bold = True
            new_r.font.size = Pt(8); new_r.font.color.rgb = _rgb_hex("B45309")
        nr = ap.add_run(f"[{i:02d}] "); nr.bold = True
        nr.font.size = Pt(10); nr.font.color.rgb = _rgb(PALETTE[color_key])
        tr = ap.add_run(title); tr.bold = True; tr.font.size = Pt(10)

        # 메타 (날짜 + 출처)
        mp = doc.add_paragraph()
        mp.paragraph_format.left_indent = Inches(0.3)
        mp.paragraph_format.space_before = Pt(0)
        mp.paragraph_format.space_after  = Pt(1)
        mr = mp.add_run(f"📅 {date}   📰 {src}")
        mr.font.size = Pt(8.5); mr.font.color.rgb = _rgb(PALETTE["gray"])

        # 요약
        if flag and sm:
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Inches(0.3)
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after  = Pt(3)
            sp.add_run(f"{flag} ").font.size = Pt(9.5)
            sp.add_run(sm).font.size = Pt(9.5)
            # EN 병기 (KO가 주언어일 때)
            if flag == '🇰🇷':
                en = art.get('summary_en','') or ''
                if en and not _is_vi(en):
                    ep = doc.add_paragraph()
                    ep.paragraph_format.left_indent = Inches(0.3)
                    ep.paragraph_format.space_before = Pt(0)
                    ep.paragraph_format.space_after  = Pt(3)
                    ep.add_run("🇺🇸 ").font.size = Pt(9)
                    ep.add_run(en[:180]).font.size = Pt(9)
                    ep.runs[-1].font.color.rgb = _rgb(PALETTE["gray"])

    if len(arts) > max_n:
        add_note(doc, f"  ※ 추가 {len(arts)-max_n}건 → Appendix 참조", color_key="gray")

def render_history_timeline(doc, arts, color_key="navy", max_n=5):
    """역사 기사 타임라인 (연도별 대표 기사)"""
    if not arts: return
    by_year = defaultdict(list)
    for art in arts:
        yr = _date_key(art)[:4]
        by_year[yr].append(art)

    for yr in sorted(by_year.keys(), reverse=True)[:5]:
        year_arts = sorted(by_year[yr], key=_date_key, reverse=True)
        yp = doc.add_paragraph()
        yp.paragraph_format.left_indent = Inches(0.1)
        yr_r = yp.add_run(f"📆 {yr}년 ({len(year_arts)}건)")
        yr_r.bold = True; yr_r.font.size = Pt(10)
        yr_r.font.color.rgb = _rgb(PALETTE[color_key])

        for art in year_arts[:2]:
            title = _best_title(art)
            date  = _date_key(art)
            flag, sm = _best_summary(art)
            ap = doc.add_paragraph()
            ap.paragraph_format.left_indent = Inches(0.35)
            ap.paragraph_format.space_before = Pt(1)
            ap.add_run(f"▸ {title}").font.size = Pt(9.5)
            mp = doc.add_paragraph()
            mp.paragraph_format.left_indent = Inches(0.5)
            mp.paragraph_format.space_before = Pt(0)
            mp.paragraph_format.space_after  = Pt(1)
            mr = mp.add_run(f"  {date}")
            mr.font.size = Pt(8.5); mr.font.color.rgb = _rgb(PALETTE["gray"])
            if flag and sm:
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Inches(0.5)
                sp.paragraph_format.space_before = Pt(0)
                sp.paragraph_format.space_after  = Pt(2)
                sp.add_run(f"  {flag} {sm[:150]}").font.size = Pt(9)

# ─────────────────────────────────────────────
# Province 활동 섹션
# ─────────────────────────────────────────────
def render_province_section(doc, arts, plan_id, kpi_db=None, theme_hex="1E40AF"):
    """성(Province)별 활동 현황 섹션"""
    prov_stats = get_province_stats(arts)
    yr_trend   = get_year_trend(arts, plan_id)

    if not prov_stats:
        add_note(doc, "ℹ️ 지역 분류 데이터가 부족합니다.", color_key="gray")
        return

    # 성별 활동 테이블
    max_cnt = prov_stats[0][1] if prov_stats else 1
    rows = []
    for prov, cnt in prov_stats[:12]:
        bar_len = int(cnt / max_cnt * 20)
        bar = "█" * bar_len + "░" * (20 - bar_len)
        if   cnt/max_cnt >= 0.7: grade = "🔴 매우 활발"
        elif cnt/max_cnt >= 0.4: grade = "🟠 활발"
        elif cnt/max_cnt >= 0.2: grade = "🟡 보통"
        else:                    grade = "🟢 적음"
        rows.append([prov, str(cnt), bar, grade])

    styled_table(doc,
        ["성/지역", "기사수", "활동 강도 (Bar)", "활동 등급"],
        rows,
        col_widths=[1.5, 0.7, 2.5, 1.2],
        header_bg="blue", font_size=9)

    # 연도별 트렌드
    trend_row = [yr_trend.get(str(y), 0) for y in range(2019, 2027)]
    styled_table(doc,
        ["2019", "2020", "2021", "2022", "2023", "2024", "2025", "2026"],
        [[str(v) for v in trend_row]],
        header_bg="gray", font_size=9)

    # KPI DB 성별 프로젝트 (있는 경우)
    if kpi_db and plan_id in kpi_db and "provinces" in kpi_db[plan_id]:
        add_note(doc, "📌 마스터플랜 KPI DB — 성별 핵심 프로젝트", color_key="blue")
        prov_data = kpi_db[plan_id]["provinces"]
        kpi_rows = []
        for prov_name, pdata in prov_data.items():
            projects = pdata.get("projects", [])
            kpi_2030 = pdata.get("kpi_2030", {})
            if kpi_2030 and isinstance(kpi_2030, dict):
                kpi_str = " | ".join(f"{k}: {v}" for k,v in list(kpi_2030.items())[:2])
            elif kpi_2030:
                kpi_str = str(kpi_2030)[:60]
            else:
                kpi_str = "-"
            proj_str_parts = []
            for px in (projects or [])[:2]:
                if isinstance(px, dict):   proj_str_parts.append(px.get("name","")[:30])
                elif isinstance(px, str):  proj_str_parts.append(px[:30])
            proj_names = "; ".join(proj_str_parts) if proj_str_parts else "-"
            kpi_rows.append([prov_name, kpi_str[:60], proj_names[:70]])
        styled_table(doc,
            ["성/도시", "KPI 목표 (2030)", "핵심 프로젝트"],
            kpi_rows,
            col_widths=[1.4, 2.5, 3.0],
            header_bg="teal", font_size=9)

# ─────────────────────────────────────────────
# Appendix 빌더
# ─────────────────────────────────────────────
def build_appendix(doc, all_arts, theme_hex="1E40AF", max_rows=80):
    """전체 관련 기사 Appendix 테이블"""
    doc.add_page_break()
    add_heading(doc, "Appendix — 전체 관련 기사 목록", 1, "navy", 13)
    add_hr(doc, theme_hex)
    doc.add_paragraph(f"총 {len(all_arts)}건  (표: 최대 {max_rows}건, 최신순)").runs[0].font.size = Pt(9)
    doc.add_paragraph()

    rows = []
    for i, art in enumerate(all_arts[:max_rows], 1):
        title = _best_title(art)[:55]
        date  = _date_key(art)
        src   = (art.get('source','') or '')[:18]
        pids  = "; ".join(_parse_plans(art)[:2])[:25]
        rows.append([str(i), date, title, src, pids])

    styled_table(doc,
        ["No", "날짜", "제목", "출처", "플랜"],
        rows,
        col_widths=[0.35, 0.9, 3.4, 1.2, 1.5],
        header_bg="navy", font_size=8.5)

# ─────────────────────────────────────────────
# 한국 기업 기회 섹션
# ─────────────────────────────────────────────
def render_korean_opportunity(doc, opportunities, theme_hex="1E40AF"):
    """
    opportunities: list of (category, HIGH/MEDIUM/LOW, description)
    """
    rows = []
    for cat, level, desc in opportunities:
        level_icon = {"HIGH":"🔴 HIGH","MEDIUM":"🟠 MEDIUM","LOW":"🟡 LOW"}.get(level, level)
        rows.append([cat, level_icon, desc[:120]])

    styled_table(doc,
        ["분야", "우선순위", "기회 내용"],
        rows,
        col_widths=[1.5, 1.0, 4.8],
        header_bg="navy", font_size=9.5)

# ─────────────────────────────────────────────
# 지도 삽입 (map_generator 연동)
# ─────────────────────────────────────────────
def insert_map(doc, plan_id, caption=None):
    """
    OSM 지도 이미지 삽입 (map_generator.py 생성 결과)
    없으면 플레이스홀더 텍스트 출력
    """
    map_path = BASE_DIR / f"assets/maps/map_{plan_id}.png"
    if map_path.exists():
        try:
            doc.add_picture(str(map_path), width=Inches(5.5))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cp = doc.add_paragraph(caption)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.runs[0].font.size = Pt(8.5)
                cp.runs[0].font.color.rgb = _rgb(PALETTE["gray"])
        except Exception as e:
            add_note(doc, f"[지도 삽입 오류: {e}]", color_key="gray")
    else:
        note_t = doc.add_table(rows=1, cols=1); note_t.style = "Table Grid"
        nc = note_t.rows[0].cells[0]; cell_bg(nc, "F1F5F9")
        np2 = nc.paragraphs[0]; np2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = np2.add_run(f"🗺️ 지도 준비 중  |  plan_id: {plan_id}\n"
                         f"map_generator.py 실행 후 assets/maps/map_{plan_id}.png 생성")
        nr.font.size = Pt(9); nr.font.color.rgb = _rgb(PALETTE["gray"])
        doc.add_paragraph()

# ─────────────────────────────────────────────
# History DB 로더
# ─────────────────────────────────────────────
def load_history(plan_ids):
    """
    plan_ids: list of str
    return: dict {plan_id: [articles]}
    """
    db_path = BASE_DIR / "config/history_db.json"
    with open(db_path, encoding='utf-8') as f:
        hdb = json.load(f)

    buckets = defaultdict(list)
    for art in hdb['articles'].values():
        plans = _parse_plans(art)
        for pid in plans:
            if pid in plan_ids:
                buckets[pid].append(art)

    return dict(buckets)

def load_kpi_db():
    kpi_path = BASE_DIR / "config/province_project_kpi.json"
    with open(kpi_path, encoding='utf-8') as f:
        return json.load(f)

def save_report(doc, output_dir, filename):
    """보고서 저장 + 파일 크기 반환"""
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    fpath = os.path.join(output_dir, filename)
    doc.save(fpath)
    sz = os.path.getsize(fpath) // 1024
    return fpath, sz

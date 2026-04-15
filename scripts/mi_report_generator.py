"""
Market Intelligence Report Generator — Vietnam Infrastructure Hub v3.0
======================================================================
신규 구조 (per user request 2026-04-14):
  - Province × Project 추적 구조 (스토리텔링 중심)
  - 국가 KPI → Province KPI → 프로젝트별 기사 매핑
  - 베트남어 기사 → 영어/한국어로만 출력
  - 관련성 필터 강화 (강 키워드 매칭)
  - 홍강 델타(VN-RED-RIVER-2030) 플랜 추가
  - 보고서 본문: 2026년 기사 우선, 없으면 최근 1년
"""
import os, sys, json, re, io
from datetime import datetime, timedelta
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 경로 ──────────────────────────────────────────────────────
BASE_DIR    = Path("/home/work/claw")
CONFIG_DIR  = BASE_DIR / "config"
OUTPUT_DIR  = Path(os.getenv("PIPELINE_OUTPUT_DIR", "/home/work/claw/outputs"))
KPI_DB_PATH = CONFIG_DIR / "province_project_kpi.json"
HIST_DB_PATH= CONFIG_DIR / "history_db.json"

# ── 마스터플랜 메타 — PDP8 통합 구조 (21개) ──────────────────
# PDP8 모(母)플랜: VN-PDP8-* (6개 하부 트래킹)
# Decision 768/QD-TTg (Apr 15, 2025) — Revised PDP8
PLAN_META = {
    # ── PDP8 에너지 그룹 (6개) ───────────────────────────────
    "VN-PDP8-RENEWABLE":  {"short":"RE",          "color":(0xF5,0x9E,0x0B), "name_ko":"PDP8 재생에너지 (태양광·풍력·수력)",         "parent":"PDP8"},
    "VN-PDP8-LNG":        {"short":"LNG",         "color":(0xEF,0x44,0x44), "name_ko":"PDP8 LNG 발전 & 인프라",                    "parent":"PDP8"},
    "VN-PDP8-NUCLEAR":    {"short":"Nuclear",     "color":(0x63,0x66,0xF1), "name_ko":"PDP8 원자력 (닌투언 1·2호기)",               "parent":"PDP8"},
    "VN-PDP8-COAL":       {"short":"Coal",        "color":(0x78,0x71,0x6C), "name_ko":"PDP8 석탄 전환·폐지 로드맵",                 "parent":"PDP8"},
    "VN-PDP8-GRID":       {"short":"Grid",        "color":(0x0E,0xA5,0xE9), "name_ko":"PDP8 송전망·스마트그리드",                   "parent":"PDP8"},
    "VN-PDP8-HYDROGEN":   {"short":"H2",          "color":(0x10,0xB9,0x81), "name_ko":"PDP8 수소·그린에너지 신산업",               "parent":"PDP8"},
    # ── 기타 인프라 (15개) ─────────────────────────────────────
    "VN-ENV-IND-1894":    {"short":"D1894",       "color":(0x10,0xB9,0x81), "name_ko":"환경산업 개발프로그램 (Decision 1894)"},
    "VN-TRAN-2055":       {"short":"Transport",   "color":(0x3B,0x82,0xF6), "name_ko":"국가 교통인프라 마스터플랜 2055"},
    "VN-URB-METRO-2030":  {"short":"Metro",       "color":(0x8B,0x5C,0xF6), "name_ko":"도시철도·메트로 개발계획 2030"},
    "VN-WAT-RESOURCES":    {"short":"WatRes",     "color":(0x06,0xB6,0xD4), "name_ko":"국가 수자원 마스터플랜 2021-2030/2050",       "parent":"WATER"},
    "VN-WAT-URBAN":        {"short":"WatUrban",   "color":(0x22,0x78,0xA4), "name_ko":"도시 상수도 인프라 개발 계획 2025/2035",          "parent":"WATER"},
    "VN-WAT-RURAL":        {"short":"WatRural",   "color":(0x16,0xA3,0x4A), "name_ko":"농촌 상수도·위생 국가전략 2030/2045",             "parent":"WATER"},
    "VN-EV-2030":         {"short":"EV",          "color":(0x22,0xC5,0x5E), "name_ko":"전기차·친환경 모빌리티 2030"},
    "VN-CARBON-2050":     {"short":"Carbon",      "color":(0x84,0xCC,0x16), "name_ko":"탄소중립 2050 로드맵"},
    "VN-WW-2030":         {"short":"Wastewater",  "color":(0x14,0xB8,0xA6), "name_ko":"국가 폐수처리 마스터플랜 2021-2030"},
    "VN-SWM-NATIONAL-2030":{"short":"SolidWaste","color":(0x92,0x40,0x0E), "name_ko":"전국 고형폐기물 통합관리 국가전략 2025/2050"},
    "VN-MEKONG-DELTA-2030":{"short":"Mekong",    "color":(0x06,0xB6,0xD4), "name_ko":"메콩델타 지역개발 마스터플랜 2021-2030"},
    "VN-SC-2030":         {"short":"SmartCity",  "color":(0xA7,0x8B,0xFA), "name_ko":"스마트시티 국가전략 2030 (기술·플랫폼)"},
    # 하노이 도시개발 그룹 (Decision 1668/QD-TTg 2024.12.27)
    "HN-URBAN-NORTH":     {"short":"HN-North",   "color":(0xDC,0x26,0x26), "name_ko":"하노이 북부 신도시 (동아인·BRG·노이바이)",    "parent":"HANOI"},
    "HN-URBAN-WEST":      {"short":"HN-West",    "color":(0x7C,0x3A,0xED), "name_ko":"하노이 서부 과학기술도시 (호아락·하이테크파크)", "parent":"HANOI"},
    "HN-URBAN-INFRA":     {"short":"HN-Infra",   "color":(0xF5,0x9E,0x0B), "name_ko":"하노이 도시 인프라 (메트로·링로드·홍강)",      "parent":"HANOI"},
    "VN-IP-NORTH-2030":   {"short":"IndPark",    "color":(0xFB,0x92,0x3C), "name_ko":"북부 산업단지 개발 계획 2030"},
    "VN-OG-2030":         {"short":"OilGas",     "color":(0x94,0xA3,0xB8), "name_ko":"석유가스 개발 계획 2030"},
    "VN-RED-RIVER-2030":  {"short":"RedRiver",   "color":(0xDC,0x26,0x26), "name_ko":"홍강 델타 지역 개발 마스터플랜 2030"},
}

# ── 관련성 강 필터 (플랜별 필수 키워드 — 최소 1개 이상 있어야 통과) ──
RELEVANCE_MUST = {
    # PDP8 에너지 그룹
    "VN-PDP8-RENEWABLE":  ["solar","wind","offshore wind","onshore wind","solar farm","floating solar","rooftop solar","hydropower","pumped storage","battery storage","renewable energy","pdp8","decision 768","điện gió","điện mặt trời","năng lượng tái tạo","thủy điện"],
    "VN-PDP8-LNG":        ["lng","lng terminal","lng power","gas-fired","gas power","natural gas","fsru","regasification","nhon trach","son my","ca mau lng","khí lng","cảng lng","điện khí","nhập khẩu lng"],
    "VN-PDP8-NUCLEAR":    ["nuclear","ninh thuan","nuclear power","nuclear plant","hạt nhân","điện hạt nhân","smr","small modular","resolution 174","resolution 189"],
    "VN-PDP8-COAL":       ["coal phase","jetp","coal retire","coal closure","coal transition","biomass cofiring","ammonia cofiring","coal plant closure","nhiệt điện than","chuyển đổi than","than đóng cửa"],
    "VN-PDP8-GRID":       ["500kv","220kv","transmission line","substation","smart grid","hvdc","power grid upgrade","dppa","direct power purchase","electricity market","lưới điện","đường dây 500kv","trạm biến áp","thị trường điện"],
    "VN-PDP8-HYDROGEN":   ["green hydrogen","hydrogen energy","hydrogen power","hydrogen ready","ammonia fuel","re export","renewable export","energy export singapore","energy export malaysia","new energy","hydro xanh","xuất khẩu điện"],
    # 기타
    "VN-ENV-IND-1894":    ["environmental industry","công nghiệp môi trường","decision 1894","recycling industrial","environmental technology","eco-industrial","environmental equipment","circular economy","epr","kinh tế tuần hoàn"],
    "VN-TRAN-2055":       ["highway","expressway","airport","seaport","railway","bridge","metro","ring road","cao tốc","sân bay","cảng biển","cầu","đường sắt"],
    "VN-URB-METRO-2030":  ["metro","subway","urban rail","tàu điện","đường sắt đô thị","mrt","cat linh","nhon","ben thanh"],
    "VN-WAT-RESOURCES":["water resources","river basin","water security","flood control","drought","saltwater intrusion","dam reservoir","irrigation","groundwater","tài nguyên nước","lưu vực sông","hạn hán","xâm nhập mặn","lũ lụt","đập","hồ chứa"],
    "VN-WAT-URBAN":    ["water supply plant","water treatment plant","clean water supply","drinking water","water pipeline","water network","water loss","nhà máy nước","cấp nước đô thị","nước sạch đô thị","đường ống cấp nước"],
    "VN-WAT-RURAL":    ["rural water supply","rural sanitation","rural clean water","commune water","ethnic minority water","cấp nước nông thôn","vệ sinh nông thôn","nước sạch nông thôn"],
    "VN-EV-2030":         ["electric vehicle","vinfast","ev ","charging station","xe điện","trạm sạc","ô tô điện"],
    "VN-CARBON-2050":     ["carbon","net zero","ndc","greenhouse","carbon credit","tín chỉ carbon","phát thải"],
    "VN-WW-2030":       ["wastewater","sewage","wwtp","sewer","nước thải","xử lý nước thải","thoát nước","yen xa","yên xá"],
    "VN-MEKONG-DELTA-2030":["mekong delta","đồng bằng sông cửu long","can tho","cần thơ","an giang","dong thap","kien giang","ca mau","soc trang","bac lieu","hau giang","vinh long","tien giang","ben tre","tra vinh","long an delta"],
    "VN-SWM-NATIONAL-2030":["solid waste","municipal waste","waste-to-energy","wte","incineration","landfill","composting","hazardous waste","medical waste","recycling","plastic waste","chất thải rắn","rác thải","xử lý rác","đốt rác","bãi chôn lấp","tái chế"],
    "VN-SC-2030":       ["smart city technology","smart city platform","digital city","iot city","e-government","digital twin","thành phố thông minh","chuyển đổi số đô thị"],
    "HN-URBAN-NORTH":  ["dong anh","me linh","soc son","noi bai","brg smart city","brg sumitomo","north hanoi","đông anh","mê linh","sóc sơn","đô thị phía bắc"],
    "HN-URBAN-WEST":   ["hoa lac","hoa lac hi-tech","western hanoi","xuan mai","son tay","thach that","hòa lạc","khu công nghệ cao hòa lạc","xuân mai"],
    "HN-URBAN-INFRA":  ["hanoi metro","hanoi ring road","ring road 4","red river corridor","to lich","hanoi bridge","hanoi master plan","metro hà nội","vành đai 4 hà nội","sông hồng hà nội","quy hoạch hà nội"],
    "VN-IP-NORTH-2030": ["industrial park","industrial zone","vsip","khu công nghiệp","fdi","vinfast","manufacturing"],
    "VN-OG-2030":       ["oil field","crude oil","petroleum","refinery","petrovietnam","dầu khí","lọc dầu","mỏ dầu"],
    "VN-RED-RIVER-2030":["red river","hong ha","sông hồng","hồng hà","홍강","dong anh","hung yen","brg","red river delta","홍강 델타"],
}

# ── 유틸리티 ──────────────────────────────────────────────────
def _rgb(t): return RGBColor(*t)
def _hex(t): return '%02X%02X%02X' % t

def _set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _hr(doc, color="CCCCCC"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'), color)
    pb.append(bot); pPr.append(pb)
    return p

def _h(doc, text, level, color, size=None):
    h = doc.add_heading(text, level)
    for run in h.runs:
        run.font.color.rgb = _rgb(color)
        if size: run.font.size = Pt(size)
    return h

def _extract_year(d):
    m = re.search(r'(20\d{2})', str(d or ''))
    return m.group(1) if m else '미상'

def _date_key(a):
    d = str(a.get('published_date','') or '')
    yr = _extract_year(d)
    if yr == '미상': return '1900-01-01'
    if not d[:4].isdigit():
        try: return datetime.strptime(d.strip(), '%b %d, %Y').strftime('%Y-%m-%d')
        except: return f'{yr}-01-01'
    return d[:10]

def _is_vi(text):
    return any(c in (text or '') for c in ['ă','ơ','ư','đ','ấ','ề','ộ','ừ','ạ','ọ','ổ','ị','ế','ặ','ứ'])

def _best_title(art):
    """영어 또는 한국어 제목 반환 (베트남어면 content 앞부분으로 대체)"""
    t = art.get('title','')
    if _is_vi(t):
        # summary_en에서 첫 문장 추출
        en = art.get('summary_en','')
        if en: return f"[VI→EN] {en[:80]}"
        ko = art.get('summary_ko','')
        if ko: return f"[VI→KO] {ko[:60]}"
        return t  # 번역 없으면 원문
    return t

def _best_summary(art):
    """한국어 요약 우선, 없으면 영어, 없으면 베트남어 내용 표시"""
    ko = art.get('summary_ko','')
    en = art.get('summary_en','')
    content = art.get('content','')

    if ko and not _is_vi(ko): return 'KO', ko
    if en and not _is_vi(en): return 'EN', en
    if ko: return 'KO(VI)', ko  # 베트남어 남은 경우
    if content: return 'CONTENT', content[:300]
    return None, None

def _relevance_check(art, plan_id):
    """강 관련성 필터 — 플랜 필수 키워드 최소 1개"""
    must = RELEVANCE_MUST.get(plan_id, [])
    if not must: return True
    text = (
        (art.get('title','') or '') + ' ' +
        (art.get('summary_en','') or '') + ' ' +
        (art.get('summary_ko','') or '') + ' ' +
        (art.get('content','') or '')[:300]
    ).lower()
    return any(kw.lower() in text for kw in must)

def _select_articles(history: list, plan_id: str):
    """2026년 기사 우선, 없으면 최근 1년, 관련성 필터 적용"""
    # 관련성 필터
    filtered = [a for a in history if _relevance_check(a, plan_id)]
    
    arts_2026 = sorted(
        [a for a in filtered if _extract_year(a.get('published_date',''))=='2026'],
        key=_date_key, reverse=True
    )
    if arts_2026:
        return arts_2026, '2026년'
    
    arts_recent = sorted(
        [a for a in filtered if _extract_year(a.get('published_date','')) in ['2025','2026']],
        key=_date_key, reverse=True
    )
    if arts_recent:
        return arts_recent, '최근 1년'
    
    return sorted(filtered, key=_date_key, reverse=True)[:20], '전체 최신'


# ════════════════════════════════════════════════════════════════
# 핵심 섹션: Province × Project 추적 (스토리텔링)
# ════════════════════════════════════════════════════════════════
def _section_province_project(doc, plan_id, all_history, kpi_data, color):
    hex_c = _hex(color)
    
    plan_kpi = kpi_data.get(plan_id, {})
    provinces_kpi = plan_kpi.get('provinces', {})
    
    _h(doc, "4. Province × Project 진행현황 추적", 1, color, 13)
    _hr(doc, hex_c)
    
    if not provinces_kpi:
        doc.add_paragraph("※ 본 플랜은 특정 Province KPI 데이터가 없습니다.")
        return
    
    # 전체 역사 기사를 Province별로 버킷팅
    art_by_prov = defaultdict(list)
    for a in all_history:
        if not _relevance_check(a, plan_id):
            continue
        prov = a.get('province','') or 'National'
        # normalize
        prov_map = {
            'vietnam':'National','national / unspecified':'National',
            'ho chi minh city':'Ho Chi Minh City','hcmc':'Ho Chi Minh City',
            'hanoi':'Hanoi','ha noi':'Hanoi',
            'da nang':'Da Nang','danang':'Da Nang',
        }
        prov_norm = prov_map.get(prov.lower(), prov)
        art_by_prov[prov_norm].append(a)
    
    # Province KPI가 있는 지역 우선, 그 다음 추가 지역
    provinces_ordered = list(provinces_kpi.keys())
    extra_provs = [p for p in art_by_prov.keys() if p not in provinces_ordered and p != 'National']
    all_provs = provinces_ordered + sorted(extra_provs, key=lambda x: -len(art_by_prov.get(x,[])))
    
    for prov in all_provs:
        prov_kpi = provinces_kpi.get(prov, {})
        arts = sorted(
            [a for a in art_by_prov.get(prov,[]) if _extract_year(a.get('published_date','')) in ['2026','2025','2024']],
            key=_date_key, reverse=True
        )
        if not arts and prov not in provinces_kpi:
            continue  # KPI도 없고 기사도 없으면 생략
        
        # Province 헤딩
        _h(doc, f"📍 {prov}", 2, color)
        
        # KPI 요약
        kpi_2030 = prov_kpi.get('kpi_2030', {})
        focus = prov_kpi.get('focus','')
        investment = prov_kpi.get('investment','')
        
        if kpi_2030 or focus:
            kpi_p = doc.add_paragraph()
            kpi_p.paragraph_format.left_indent = Inches(0.2)
            if focus:
                r = kpi_p.add_run(f"🎯 Focus: {focus}  ")
                r.bold = True; r.font.size = Pt(10); r.font.color.rgb = _rgb(color)
            if investment:
                kpi_p.add_run(f"| 투자: {investment}").font.size = Pt(9)
            if kpi_2030:
                kpi_str = " / ".join(f"{k}: {v}" for k,v in kpi_2030.items())
                kp = doc.add_paragraph()
                kp.paragraph_format.left_indent = Inches(0.3)
                r2 = kp.add_run("2030 KPI: ")
                r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0x1D,0x4E,0xD8)
                kp.add_run(kpi_str).font.size = Pt(9)
        
        # 프로젝트 목록 (KPI DB에서)
        projects = prov_kpi.get('projects', [])
        if projects:
            _h(doc, "주요 프로젝트", 3, color)
            proj_tbl = doc.add_table(rows=1, cols=5)
            proj_tbl.style = "Table Grid"
            for j, h in enumerate(["프로젝트명","규모/용량","투자","현황","기사"]):
                proj_tbl.rows[0].cells[j].text = h
                proj_tbl.rows[0].cells[j].paragraphs[0].runs[0].bold = True
                _set_bg(proj_tbl.rows[0].cells[j], hex_c)
                proj_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                proj_tbl.rows[0].cells[j].paragraphs[0].runs[0].font.size = Pt(9)
            
            for proj in projects:
                row = proj_tbl.add_row()
                pname = proj.get('name','')
                # 관련 기사 매핑 (프로젝트명 키워드로 매칭)
                proj_kws = pname.lower().split()[:3]
                related = [a for a in arts 
                           if any(kw in (a.get('title','') + (a.get('summary_en','') or '')).lower()
                                  for kw in proj_kws if len(kw) > 3)]
                
                status = proj.get('status','')
                status_icon = {"운영중":"🟢","건설중":"🟡","계획":"⚪","시범운영":"🔵","타당성조사":"🔵","MOU":"🟠","진행중":"🟡"}.get(status, "⚪")
                
                row.cells[0].text = pname
                row.cells[1].text = proj.get('capacity', proj.get('area_ha','') and f"{proj.get('area_ha','')}ha" or proj.get('scope','') or '-')
                row.cells[2].text = proj.get('investment', proj.get('developer','') or '-')
                row.cells[3].text = f"{status_icon} {status}"
                row.cells[4].text = f"{len(related)}건" if related else "-"
                
                # 상태별 색상
                if status == '운영중': _set_bg(row.cells[3], "D1FAE5")
                elif status in ['건설중','진행중']: _set_bg(row.cells[3], "FEF9C3")
                elif status == '계획': _set_bg(row.cells[3], "F1F5F9")
                
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs: run.font.size = Pt(9)
            
            doc.add_paragraph()
        
        # 관련 기사 (최신 5건)
        if arts:
            _h(doc, f"📰 관련 기사 ({len(arts)}건)", 3, color)
            for i, art in enumerate(arts[:5], 1):
                display_title = _best_title(art)
                lang, summary = _best_summary(art)
                date = _date_key(art)
                src = art.get('source','') or ''
                
                ap = doc.add_paragraph()
                ap.paragraph_format.left_indent = Inches(0.2)
                nr = ap.add_run(f"[{i}] ")
                nr.bold = True; nr.font.size = Pt(10); nr.font.color.rgb = _rgb(color)
                tr = ap.add_run(display_title[:80])
                tr.bold = True; tr.font.size = Pt(10)
                
                meta_p = doc.add_paragraph()
                meta_p.paragraph_format.left_indent = Inches(0.4)
                meta_p.add_run(f"📅 {date}  |  📰 {src[:25]}").font.size = Pt(8.5)
                
                if summary:
                    sp = doc.add_paragraph()
                    sp.paragraph_format.left_indent = Inches(0.4)
                    flag = "🇰🇷" if lang.startswith('KO') else "🇺🇸" if lang.startswith('EN') else "📄"
                    r1 = sp.add_run(f"{flag} ")
                    r1.font.size = Pt(9.5)
                    sp.add_run(summary[:200]).font.size = Pt(9.5)
                
                url_p = doc.add_paragraph()
                url_p.paragraph_format.left_indent = Inches(0.4)
                ur = url_p.add_run(art.get('url','')[:80])
                ur.font.size = Pt(8); ur.font.color.rgb = RGBColor(0x1D,0x4E,0xD8)
                doc.add_paragraph()
            
            if len(arts) > 5:
                more_p = doc.add_paragraph()
                more_p.paragraph_format.left_indent = Inches(0.2)
                more_p.add_run(f"  ※ 추가 {len(arts)-5}건 — Appendix 참조").font.size = Pt(9)
        else:
            np = doc.add_paragraph()
            np.paragraph_format.left_indent = Inches(0.2)
            np.add_run("ℹ️ 최근 기사 없음 (데이터베이스에서 모니터링 중)").font.size = Pt(9)
        
        _hr(doc, "E2E8F0")
        doc.add_paragraph()


# ════════════════════════════════════════════════════════════════
# 보고서 생성 메인
# ════════════════════════════════════════════════════════════════
def create_mi_report(plan_id: str, all_history: list, kpi_data: dict, output_dir: str) -> str:
    if plan_id not in PLAN_META:
        return None
    
    meta  = PLAN_META[plan_id]
    color = meta["color"]
    hex_c = _hex(color)
    plan_kpi = kpi_data.get(plan_id, {})
    
    report_arts, yr_label = _select_articles(all_history, plan_id)
    
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0); sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.2); sec.right_margin  = Inches(1.2)
    
    now   = datetime.now()
    week  = now.isocalendar()[1]
    year  = now.year
    date_e= now.strftime("%Y-%m-%d")
    date_s= (now - timedelta(days=6)).strftime("%Y-%m-%d")

    # ══ 표지 ══
    doc.add_paragraph()
    tbl_hdr = doc.add_table(rows=1, cols=1); tbl_hdr.style="Table Grid"
    c = tbl_hdr.rows[0].cells[0]; _set_bg(c, "0F172A")
    p = c.paragraphs[0]
    r = p.add_run("VIETNAM INFRASTRUCTURE INTELLIGENCE HUB")
    r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tr2 = tp.add_run(meta["name_ko"])
    tr2.bold=True; tr2.font.size=Pt(20); tr2.font.color.rgb=_rgb(color)
    
    sub_p = doc.add_paragraph(); sub_p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run(f"Plan ID: {plan_id}").font.size=Pt(10)
    
    doc.add_paragraph()
    it = doc.add_table(rows=5, cols=2); it.style="Table Grid"
    nat_kpi = plan_kpi.get('national_kpi',{})
    t30 = nat_kpi.get('2030_target',{})
    kpi_str = " / ".join(f"{k}:{v}" for k,v in list(t30.items())[:3]) if t30 else "-"
    
    for i,(k,v) in enumerate([
        ("보고 기간",f"{date_s} ~ {date_e}  (W{week}/{year})"),
        ("근거 법령", plan_kpi.get('decision','-')),
        ("총 투자 규모", nat_kpi.get('total_investment','-')),
        ("2030 핵심 KPI", kpi_str),
        ("기사 현황", f"역사 전체 {len(all_history)}건 | 보고({yr_label}) {len(report_arts)}건"),
    ]):
        it.rows[i].cells[0].text=k; it.rows[i].cells[1].text=v
        _set_bg(it.rows[i].cells[0], hex_c)
        it.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        it.rows[i].cells[0].paragraphs[0].runs[0].bold=True
        for cell in it.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(10)
    
    doc.add_paragraph()
    gp = doc.add_paragraph(); gp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    gp.add_run(f"생성: {now.strftime('%Y-%m-%d %H:%M')}  |  Pipeline v3.0  |  CONFIDENTIAL").font.size=Pt(9)
    doc.add_page_break()

    # ══ 1. Executive Briefing ══
    _h(doc, "1. Executive Briefing", 1, color, 14); _hr(doc, hex_c)
    
    kt = doc.add_table(rows=2, cols=4); kt.style="Table Grid"
    prov_top = Counter(a.get('province','') for a in all_history
                       if (a.get('province','') or '').lower() not in 
                          ['vietnam','national / unspecified','national','']).most_common(1)
    kpi_vals = [
        ("역사 전체 기사", str(len(all_history))),
        (f"보고 기사 ({yr_label})", str(len(report_arts))),
        ("관련성 통과율", f"{len(report_arts)/max(len(all_history),1)*100:.0f}%"),
        ("최다 활동 지역", prov_top[0][0] if prov_top else "전국"),
    ]
    for i,(h,v) in enumerate(kpi_vals):
        kt.rows[0].cells[i].text=h; kt.rows[1].cells[i].text=v
        _set_bg(kt.rows[0].cells[i], hex_c)
        kt.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        kt.rows[0].cells[i].paragraphs[0].runs[0].bold=True
        kt.rows[0].cells[i].paragraphs[0].runs[0].font.size=Pt(9)
        kt.rows[1].cells[i].paragraphs[0].runs[0].font.size=Pt(14)
        kt.rows[1].cells[i].paragraphs[0].runs[0].bold=True
        for r2 in range(2): kt.rows[r2].cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ══ 2. 국가 KPI 목표 ══
    _h(doc, "2. 국가 KPI 목표 및 플랜 개요", 1, color); _hr(doc, hex_c)
    
    if nat_kpi:
        nt = doc.add_table(rows=1, cols=3); nt.style="Table Grid"
        for j,h in enumerate(["구분","KPI 항목","목표치"]):
            nt.rows[0].cells[j].text=h
            nt.rows[0].cells[j].paragraphs[0].runs[0].bold=True
            _set_bg(nt.rows[0].cells[j], hex_c)
            nt.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        
        for milestone, kpis in [
            ("기준 (2020)",  nat_kpi.get('2020_baseline',{})),
            ("목표 (2025)",  nat_kpi.get('2025_target',{})),
            ("목표 (2030)",  nat_kpi.get('2030_target',{})),
            ("비전 (2050)",  nat_kpi.get('2050_target',{})),
        ]:
            if not kpis: continue
            for ki, (k,v) in enumerate(kpis.items()):
                row = nt.add_row()
                row.cells[0].text = milestone if ki==0 else ""
                row.cells[1].text = k.replace('_',' ')
                row.cells[2].text = str(v)
                if ki==0: _set_bg(row.cells[0], "F1F5F9")
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs: run.font.size=Pt(9.5)
        doc.add_paragraph()
        
        ti_p = doc.add_paragraph()
        ti_p.paragraph_format.left_indent = Inches(0.2)
        ti = ti_p.add_run(f"💰 총 투자 규모: {nat_kpi.get('total_investment','-')}")
        ti.bold=True; ti.font.size=Pt(10); ti.font.color.rgb=_rgb(color)
        if nat_kpi.get('total_projects'):
            ti_p.add_run(f"  |  📋 총 프로젝트 수: {nat_kpi['total_projects']}건").font.size=Pt(10)
        doc.add_paragraph()

    # ══ 3. 전국 단위 주요 뉴스 (Province 미분류) ══
    national_arts = [a for a in report_arts
                     if (a.get('province','') or '').lower() in 
                        ['vietnam','national / unspecified','national','']][:10]
    
    _h(doc, f"3. 전국 단위 주요 뉴스 ({yr_label}, {len(national_arts)}건)", 1, color)
    _hr(doc, hex_c)
    
    if national_arts:
        for i, art in enumerate(national_arts, 1):
            display_title = _best_title(art)
            lang, summary = _best_summary(art)
            date = _date_key(art)
            src = art.get('source','') or ''
            
            ap = doc.add_paragraph()
            nr = ap.add_run(f"[{i:02d}] ")
            nr.bold=True; nr.font.size=Pt(10); nr.font.color.rgb=_rgb(color)
            ap.add_run(display_title[:85]).bold=True; ap.runs[-1].font.size=Pt(11)
            
            mt = doc.add_table(rows=1, cols=3); mt.style="Table Grid"
            for j,txt in enumerate([f"📅 {date}", f"📰 {src[:30]}", f"🏙️ {art.get('province','전국')}"]):
                mt.rows[0].cells[j].text=txt
                _set_bg(mt.rows[0].cells[j], "F1F5F9")
                mt.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(8.5)
            doc.add_paragraph()
            
            if summary and lang:
                flag = "🇰🇷" if lang.startswith('KO') else "🇺🇸"
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Inches(0.3)
                sp.add_run(f"{flag} ").font.size=Pt(9.5)
                sp.add_run(summary[:250]).font.size=Pt(9.5)
            
            # 영어 요약도 추가 (한국어만 있을 때)
            if lang and lang.startswith('KO') and art.get('summary_en'):
                ep = doc.add_paragraph()
                ep.paragraph_format.left_indent = Inches(0.3)
                ep.add_run("🇺🇸 ").font.size=Pt(9.5)
                ep.add_run(art.get('summary_en','')[:200]).font.size=Pt(9.5)
            
            up = doc.add_paragraph()
            up.paragraph_format.left_indent = Inches(0.3)
            ur = up.add_run(art.get('url','')[:80])
            ur.font.size=Pt(8); ur.font.color.rgb=RGBColor(0x1D,0x4E,0xD8)
            doc.add_paragraph(); _hr(doc,"E2E8F0"); doc.add_paragraph()
    else:
        doc.add_paragraph(f"※ {yr_label} 전국 단위 기사 없음")
    
    doc.add_page_break()

    # ══ 4. Province × Project 추적 ══
    _section_province_project(doc, plan_id, all_history, kpi_data, color)
    
    doc.add_page_break()

    # ══ 5. Province 활동도 요약 ══
    _h(doc, "5. Province 활동도 종합", 1, color); _hr(doc, hex_c)
    
    # 관련성 통과한 기사만
    relevant = [a for a in all_history if _relevance_check(a, plan_id)]
    prov_cnt = Counter()
    for a in relevant:
        pv = a.get('province','') or 'National'
        if pv.lower() in ['vietnam','national / unspecified']: pv='🇻🇳 전국'
        prov_cnt[pv] += 1
    
    if prov_cnt:
        at = doc.add_table(rows=1, cols=4); at.style="Table Grid"
        for j,h in enumerate(["지방성","건수","비중","활동도"]):
            at.rows[0].cells[j].text=h
            at.rows[0].cells[j].paragraphs[0].runs[0].bold=True
            _set_bg(at.rows[0].cells[j], hex_c)
            at.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        total = sum(prov_cnt.values())
        max_c = prov_cnt.most_common(1)[0][1]
        for pv, cnt in prov_cnt.most_common(12):
            pct = cnt/total*100
            ratio = cnt/max_c
            grade = "🔴 매우 활발" if ratio>=0.6 else "🟠 활발" if ratio>=0.3 else "🟡 보통" if ratio>=0.15 else "🟢 적음"
            row = at.add_row()
            row.cells[0].text=pv; row.cells[1].text=f"{cnt}건"
            row.cells[2].text=f"{pct:.1f}%"; row.cells[3].text=grade
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.font.size=Pt(9)
    doc.add_paragraph()

    # ══ 6. 리스크 & 기회 ══
    _h(doc, "6. 리스크 및 비즈니스 기회", 1, color); _hr(doc, hex_c)
    
    plan_full = {
        "VN-WW-2030":   ("처리율 15→70% 갭이 최대 리스크 (ODA 의존)","WWTP EPC, MBR 처리기술, 하수관, 슬러지 처리","JICA·ADB ODA 자금 확정. 2026-2028년 입찰 집중."),
        "VN-WAT-URBAN":("노후관 교체 재정 부담, 지하수 오염","정수 설비, PE 상수관, 수질 모니터링","농촌 상수도 연결 사업 지속 확대."),
        "VN-ENV-IND-1894":("지방 재정 한계, 주민 반대(NIMBY)","WTE EPC, 환경 모니터링, 매립가스 활용","EPR 시행으로 민간 투자 가속 예상."),
        "VN-PDP8-RENEWABLE":  ("FIT 불확실, 계통 연계 지연","해상풍력 EPC, ESS, 송전 자재","2026년 해상풍력 경쟁입찰 시작."),
        "VN-TRAN-2055": ("토지보상 지연, 예산 집행률 65-70%","교량·터널 공법, 스마트 톨링","Ring Road 입찰 2026년 집중."),
        "VN-PDP8-LNG":  ("LNG 현물가 변동성","LNG 저장탱크 EPC, FSRU, 가스 파이프라인","2026-2027 LNG터미널 EPC 입찰 집중."),
        "VN-RED-RIVER-2030":("토지보상 지연, 홍강 수위 변동","신도시 인프라 EPC, 스마트시티 솔루션","BRG 스마트시티 Phase2 입찰 진행중."),
        "VN-SC-2030":   ("데이터 주권법, 공공예산 집행 지연","도시통합관제, 스마트파킹, 전자정부","2026년 Thu Duc 1단계 입찰."),
        "VN-IP-NORTH-2030":("미중 무역전쟁 완화시 FDI 분산 가능","산업단지 인프라 EPC, 스마트팩토리","반도체·EV 서플라이체인 이전 수혜."),
    }
    
    risk, opp, outlook = plan_full.get(plan_id, ("-", "-", "-"))
    
    rb = doc.add_table(rows=4, cols=2); rb.style="Table Grid"
    for i,(k,v) in enumerate([
        ("구분","내용"),
        ("🚨 주요 리스크", risk),
        ("💡 한국 기업 기회", opp),
        ("📈 시장 전망", outlook),
    ]):
        rb.rows[i].cells[0].text=k; rb.rows[i].cells[1].text=v
        if i==0:
            for cell in rb.rows[i].cells:
                _set_bg(cell, hex_c)
                cell.paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
                cell.paragraphs[0].runs[0].bold=True
        else:
            _set_bg(rb.rows[i].cells[0], "F1F5F9")
            rb.rows[i].cells[0].paragraphs[0].runs[0].bold=True
        for cell in rb.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.size=Pt(10)
    doc.add_paragraph()

    # ══ 7. Appendix ══
    _h(doc, f"7. Appendix — 전체 관련 기사 ({len(relevant)}건, 관련성 통과)", 1, color)
    _hr(doc, hex_c)
    
    note_p = doc.add_paragraph(
        f"※ 관련성 필터 통과 기사 (총 역사 {len(all_history)}건 중 {len(relevant)}건)\n"
        "   전체 이력은 Database Excel 파일의 플랜별 시트에서 확인하세요."
    )
    note_p.runs[0].font.size=Pt(9); note_p.runs[0].font.color.rgb=RGBColor(0x64,0x74,0x8B)
    doc.add_paragraph()
    
    recent50 = sorted(relevant, key=_date_key, reverse=True)[:50]
    if recent50:
        app_t = doc.add_table(rows=1, cols=5); app_t.style="Table Grid"
        for j,h in enumerate(["No","날짜","제목 (EN/KO)","출처","지역"]):
            app_t.rows[0].cells[j].text=h
            app_t.rows[0].cells[j].paragraphs[0].runs[0].bold=True
            _set_bg(app_t.rows[0].cells[j], "1E40AF")
            app_t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            app_t.rows[0].cells[j].paragraphs[0].runs[0].font.size=Pt(9)
        for ri, art in enumerate(recent50, 1):
            row = app_t.add_row()
            row.cells[0].text=str(ri)
            row.cells[1].text=_date_key(art)
            row.cells[2].text=_best_title(art)[:60]
            row.cells[3].text=(art.get('source','') or '')[:20]
            row.cells[4].text=(art.get('province','') or '')[:20]
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs: run.font.size=Pt(8.5)

    # ── 저장 ──
    fname = f"MI_REPORT_{meta['short']}_W{week:02d}_{now.strftime('%Y%m%d')}.docx"
    fpath = os.path.join(output_dir, fname)
    doc.save(fpath)
    return fpath


# ════════════════════════════════════════════════════════════════
# main
# ════════════════════════════════════════════════════════════════
def main():
    """
    v4.0 오케스트레이터 — report_regional.py + report_kpi.py + 통합 보고서 3종
    15개 보고서 일괄 생성
    """
    out_dir = OUTPUT_DIR / "reports" / "MI_Reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    print(f"📁 출력: {out_dir}")
    print()

    generated = []

    def _run(label, fn, *args):
        print(f"  {label}", end=' ', flush=True)
        try:
            fpath = fn(*args)
            if fpath and os.path.exists(fpath):
                sz = os.path.getsize(fpath) // 1024
                print(f"→ ✅ {sz}KB  [{os.path.basename(fpath)}]")
                generated.append(fpath); return fpath
            else:
                print("→ ❌ (파일 없음)"); return None
        except Exception as e:
            print(f"→ ❌ {e}"); return None

    # ── 1. PDP8 통합 (에너지 6 sub)
    from scripts.mi_pdp8_report_v4 import generate_pdp8_report
    _run("🔋 PDP8 통합 (KPI형)        ", generate_pdp8_report, str(out_dir))

    # ── 2. Water 통합 (수자원 3 sub)
    from scripts.mi_water_report_v4 import generate_water_report
    _run("💧 Water 통합 (KPI형)       ", generate_water_report, str(out_dir))

    # ── 3. Hanoi 통합 (도시개발 3 sub)
    from scripts.mi_hanoi_report import generate_hanoi_report
    _run("🏙️ Hanoi 통합 (지역성)      ", generate_hanoi_report, str(out_dir))

    # ── 4-10. 지역성 보고서
    from scripts.report_regional import generate_regional_report
    regional_plans = [
        ("🛣️ VN-TRAN-2055           ", "VN-TRAN-2055"),
        ("🚇 VN-URB-METRO-2030      ", "VN-URB-METRO-2030"),
        ("🌊 VN-MEKONG-DELTA-2030   ", "VN-MEKONG-DELTA-2030"),
        ("🏔️ VN-RED-RIVER-2030      ", "VN-RED-RIVER-2030"),
        ("🏭 VN-IP-NORTH-2030       ", "VN-IP-NORTH-2030"),
        ("💧 VN-WW-2030             ", "VN-WW-2030"),
        ("♻️ VN-SWM-NATIONAL-2030   ", "VN-SWM-NATIONAL-2030"),
    ]
    for label, pid in regional_plans:
        _run(label, generate_regional_report, pid, str(out_dir))

    # ── 11-15. KPI 보고서
    from scripts.report_kpi import generate_kpi_report
    kpi_plans = [
        ("🌿 VN-ENV-IND-1894        ", "VN-ENV-IND-1894"),
        ("🏙️ VN-SC-2030             ", "VN-SC-2030"),
        ("⛽ VN-OG-2030             ", "VN-OG-2030"),
        ("🚗 VN-EV-2030             ", "VN-EV-2030"),
        ("🌍 VN-CARBON-2050         ", "VN-CARBON-2050"),
    ]
    for label, pid in kpi_plans:
        _run(label, generate_kpi_report, pid, str(out_dir))

    total_kb = sum(os.path.getsize(p)//1024 for p in generated if p and os.path.exists(p))
    print(f"\n✅ MI 보고서 v4.0 생성 완료: {len(generated)}/15개  |  총 {total_kb}KB")
    return generated

if __name__ == "__main__":
    sys.path.insert(0, str(BASE_DIR))
    main()
